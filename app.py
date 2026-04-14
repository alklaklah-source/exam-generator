import streamlit as st
import pandas as pd
import xlrd
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment
import random
import io
import zipfile
import gc
import json
import re
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time

# ===== مكتبات اختيارية =====
try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    from google import genai as _google_genai
    GEMINI_OK = True
except ImportError:
    try:
        import google.generativeai as _google_genai_old
        GEMINI_OK = True
    except ImportError:
        GEMINI_OK = False

try:
    import anthropic as anthropic_lib
    CLAUDE_OK = True
except ImportError:
    CLAUDE_OK = False

OPENROUTER_OK = True  # يستخدم requests المدمجة - لا يحتاج تثبيتاً إضافياً


st.set_page_config(
    page_title="منصة المهندس أحمد سيداحمد - معهد طيبة",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;700;900&display=swap');

    * {font-family: 'Tajawal', sans-serif;}

    .main {background-color: #f8f9fa;}

    /* الهيدر الرئيسي */
    .header-container {
        background: linear-gradient(90deg, #004d40 0%, #00695c 100%);
        padding: 16px 30px;
        border-radius: 15px;
        margin-bottom: 25px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.15);
        display: flex;
        align-items: center;
        justify-content: space-between;
    }
    .logo-box {
        flex-shrink: 0;
    }
    .header-logo {
        width: 120px;
        display: block;
    }
    .header-text {
        flex: 1;
        text-align: center;
    }
    .header-title {
        font-size: 28px;
        font-weight: 900;
        color: #ffffff;
        margin: 0;
    }
    .header-subtitle {
        font-size: 17px;
        color: rgba(255,255,255,0.88);
        margin-top: 6px;
    }

    /* تحسين البطاقات */
    div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlock"] {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        border: 1px solid #eaeaea;
    }

    /* تحسين الكونسول */
    .log-container {
        background-color: #1e1e1e;
        color: #00e676;
        font-family: 'Courier New', monospace;
        padding: 15px;
        border-radius: 8px;
        height: 250px;
        overflow-y: auto;
        font-size: 13px;
        direction: ltr;
        text-align: left;
        border: 2px solid #333;
    }

    /* تحسين الأزرار */
    .stButton button {
        width: 100%;
        background-color: #00695c;
        color: white;
        font-weight: bold;
        font-size: 18px;
        padding: 12px;
        border-radius: 8px;
        border: none;
        transition: all 0.3s;
    }
    .stButton button:hover {
        background-color: #004d40;
        box-shadow: 0 4px 10px rgba(0,0,0,0.2);
    }

    h1, h2, h3, h4, p, label {text-align: right; direction: rtl;}
    .stFileUploader label {font-weight: bold; color: #333;}
</style>
""", unsafe_allow_html=True)


def is_cell_highlighted(cell):
    if not cell.fill or not cell.fill.start_color: return False
    color = cell.fill.start_color
    if color.type == 'rgb':
        if color.rgb in [None, '00000000', 'FFFFFFFF']: return False
        return True
    elif color.type == 'theme': return True
    elif color.type == 'indexed':
        if color.indexed != 64: return True
    return False


def normalize_text(text):
    if text is None: return ""
    text = str(text).strip()
    return text.replace('ة', 'ه').replace('ى', 'ي').replace('أ', 'ا').replace('إ', 'ا').replace('آ', 'ا')

def clean_for_comp(text): return normalize_text(text).replace(" ", "")

def normalize_difficulty(val):
    """تحويل قيمة الصعوبة إلى إحدى القيم الثلاث: سهل / متوسط / صعب"""
    if val is None: return "غير محدد"
    v = normalize_text(str(val)).replace(" ", "").lower()
    if any(x in v for x in ['سهل', 'easy', 'سهله']): return "سهل"
    if any(x in v for x in ['متوسط', 'medium', 'medium']): return "متوسط"
    if any(x in v for x in ['صعب', 'hard', 'صعبه']): return "صعب"
    return "غير محدد"

def force_align_options(options, correct_text, target_idx):
    final_opts = options[:]
    while len(final_opts) < 4: final_opts.append("---")
    current_idx = -1
    clean_corr = clean_for_comp(correct_text)
    for i, opt in enumerate(final_opts):
        if clean_for_comp(str(opt)) == clean_corr: current_idx = i; break
    if current_idx != -1:
        if current_idx != target_idx:
            temp = final_opts[target_idx]; final_opts[target_idx] = final_opts[current_idx]; final_opts[current_idx] = temp
    else: final_opts[target_idx] = correct_text
    return final_opts

# --- Word Formatting ---
def set_section_rtl_and_margins(section):
    sectPr = section._sectPr
    if not sectPr.find(qn('w:bidi')):
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); sectPr.append(bidi)
    section.left_margin = Inches(0.4); section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5)

def fix_paragraph_alignment(paragraph):
    p_fmt = paragraph.paragraph_format
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_fmt.bidi = True
    p_fmt.left_indent = Pt(0); p_fmt.right_indent = Pt(0); p_fmt.first_line_indent = Pt(0)
    p_fmt.space_before = Pt(4); p_fmt.space_after = Pt(4)

def configure_table_layout(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None: tbl_pr = table._tbl.add_tblPr()
    bidi = OxmlElement('w:bidiVisual'); tbl_pr.append(bidi)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'right'); tbl_pr.append(jc)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}'); border.set(qn('w:val'), 'nil'); tbl_borders.append(border)
    tbl_pr.append(tbl_borders)
    tbl_cell_mar = OxmlElement('w:tblCellMar')
    for m in ['top', 'start', 'bottom', 'end']:
        width = OxmlElement(f'w:{m}'); width.set(qn('w:w'), '0'); width.set(qn('w:type'), 'dxa'); tbl_cell_mar.append(width)
    tbl_pr.append(tbl_cell_mar)

def force_font(run, size=14, is_bold=False):
    run.font.name = 'Times New Roman'; run.font.size = Pt(size); run.bold = is_bold
    rPr = run._element.get_or_add_rPr()
    fonts = OxmlElement('w:rFonts'); fonts.set(qn('w:ascii'), 'Times New Roman'); fonts.set(qn('w:hAnsi'), 'Times New Roman'); fonts.set(qn('w:eastAsia'), 'Times New Roman'); fonts.set(qn('w:cs'), 'Times New Roman'); rPr.append(fonts)
    sz = OxmlElement('w:szCs'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    b = OxmlElement('w:bCs'); b.set(qn('w:val'), '1' if is_bold else '0'); rPr.append(b)
    if not rPr.find(qn('w:rtl')): rPr.append(OxmlElement('w:rtl'))

def add_question_block(doc, q_num, q_text, options):
    table = doc.add_table(rows=2, cols=4)
    configure_table_layout(table)
    q_cell = table.cell(0, 0).merge(table.cell(0, 3))
    q_p = q_cell.paragraphs[0]
    q_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; q_p.paragraph_format.bidi = True
    run = q_p.add_run(f"{q_num}) {q_text}"); force_font(run, size=14, is_bold=True)
    chars = ['أ', 'ب', 'ج', 'د']
    for i, opt_text in enumerate(options):
        opt_cell = table.cell(1, i); opt_p = opt_cell.paragraphs[0]
        opt_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; opt_p.paragraph_format.bidi = True
        lbl = chars[i] if i < len(chars) else "-"
        run_opt = opt_p.add_run(f"{lbl}. {opt_text}"); force_font(run_opt, size=14, is_bold=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# --- File Readers ---

def _detect_format_and_cols(header_values):
    """
    يكشف شكل البنك ويرجع نوعه والأعمدة:
    - النوع B (الجديد): عمود "الإجابة الصحيحة" + أعمدة "خطأ1/2/3"
    - النوع A (القديم): الخيارات بين عمود الوحدة والسؤال
    """
    cols = {'u': -1, 'q': -1, 'obj': -1, 'diff': -1, 'correct': -1, 'wrongs': []}
    for c, v in enumerate(header_values):
        nv = v  # already normalized
        if 'وحده' in nv or 'الوحده' in nv: cols['u'] = c
        elif 'سؤال' in nv: cols['q'] = c
        elif 'هدف' in nv: cols['obj'] = c
        elif 'صعوب' in nv or 'مستو' in nv: cols['diff'] = c
        elif 'اجاب' in nv and 'صحيح' in nv: cols['correct'] = c   # الإجابة الصحيحة
        elif 'خطا' in nv or 'خطأ' in nv: cols['wrongs'].append(c)  # خطأ1 / خطأ2 / خطأ3

    # النوع B: يوجد عمود الإجابة الصحيحة وعمود خطأ على الأقل
    fmt = 'B' if (cols['correct'] != -1 and len(cols['wrongs']) > 0) else 'A'
    return fmt, cols

def _read_xls_questions(file_obj):
    try:
        content = file_obj.read()
        book = xlrd.open_workbook(file_contents=content, formatting_info=True)
    except: return pd.DataFrame()
    sh = None
    for n in book.sheet_names():
        if 'بنك' in n or 'اسئله' in normalize_text(n): sh = book.sheet_by_name(n); break
    if not sh: sh = book.sheet_by_index(0)

    hr = -1; fmt = 'A'; cols = {}
    for r in range(min(20, sh.nrows)):
        vs = [normalize_text(sh.cell_value(r, c)) for c in range(sh.ncols)]
        if any('سؤال' in x for x in vs) and (any('وحده' in x for x in vs) or any('اجاب' in x for x in vs)):
            hr = r
            fmt, cols = _detect_format_and_cols(vs)
            break
    if hr == -1: return pd.DataFrame()

    data = []
    for r in range(hr+1, sh.nrows):
        u = normalize_text(sh.cell_value(r, cols['u'])) if cols['u'] != -1 else ""
        q = str(sh.cell_value(r, cols['q'])).strip() if cols['q'] != -1 else ""
        if not q: continue
        cat = u
        if cols['obj'] != -1: cat += str(sh.cell_value(r, cols['obj']))

        diff_val = "غير محدد"
        if cols['diff'] != -1 and cols['diff'] < sh.ncols:
            diff_val = normalize_difficulty(sh.cell_value(r, cols['diff']))

        if fmt == 'B':
            # الشكل الجديد: عمود الإجابة الصحيحة + أعمدة الخطأ
            corr = str(sh.cell_value(r, cols['correct'])).strip() if cols['correct'] < sh.ncols else ""
            wrongs = []
            for wc in cols['wrongs']:
                if wc < sh.ncols:
                    w = str(sh.cell_value(r, wc)).strip()
                    if w: wrongs.append(w)
            if not corr: continue
            all_opts = [corr] + wrongs
            all_opts = [x for x in all_opts if x]
            if len(all_opts) >= 2:
                data.append({'category': cat, 'unit': u, 'question': q,
                             'options': all_opts[:4], 'correct_text': corr, 'difficulty': diff_val})
        else:
            # الشكل القديم: الخيارات بين عمود الوحدة والسؤال
            start_opt = min(cols['u'], cols['q']) + 1
            end_opt = max(cols['u'], cols['q'])
            opt_cols = list(range(start_opt, end_opt))
            o_txt = []; o_clr = []
            for ci in opt_cols:
                if ci < sh.ncols:
                    val = str(sh.cell_value(r, ci)).strip()
                    clr = book.xf_list[sh.cell_xf_index(r, ci)].background.pattern_colour_index
                    o_txt.append(val); o_clr.append(clr)
            corr = ""
            for i, x in enumerate(o_txt):
                if x and o_clr[i] != 64: corr = x; break
            real_opts = [x for x in o_txt if x]
            if real_opts and corr:
                data.append({'category': cat, 'unit': u, 'question': q,
                             'options': real_opts[:4], 'correct_text': corr, 'difficulty': diff_val})
    del book
    return pd.DataFrame(data)

def _read_xlsx_questions(file_obj):
    try: wb = openpyxl.load_workbook(file_obj, data_only=False)
    except: return pd.DataFrame()
    sh = wb.active
    rows = list(sh.iter_rows())

    hr = -1; fmt = 'A'; cols = {}
    for r_idx, row in enumerate(rows[:20]):
        vs = [normalize_text(cell.value) for cell in row]
        if any('سؤال' in x for x in vs) and (any('وحده' in x for x in vs) or any('اجاب' in x for x in vs)):
            hr = r_idx
            fmt, cols = _detect_format_and_cols(vs)
            break
    if hr == -1: return pd.DataFrame()

    data = []
    for row in rows[hr+1:]:
        try:
            u = normalize_text(row[cols['u']].value) if cols['u'] != -1 and cols['u'] < len(row) else ""
            q = str(row[cols['q']].value if row[cols['q']].value else "").strip() if cols['q'] != -1 and cols['q'] < len(row) else ""
        except IndexError: continue
        if not q: continue
        cat = u
        if cols['obj'] != -1 and cols['obj'] < len(row): cat += str(row[cols['obj']].value)

        diff_val = "غير محدد"
        if cols['diff'] != -1 and cols['diff'] < len(row):
            diff_val = normalize_difficulty(row[cols['diff']].value)

        if fmt == 'B':
            # الشكل الجديد: عمود الإجابة الصحيحة + أعمدة الخطأ
            corr = str(row[cols['correct']].value).strip() if cols['correct'] < len(row) and row[cols['correct']].value else ""
            wrongs = []
            for wc in cols['wrongs']:
                if wc < len(row) and row[wc].value:
                    w = str(row[wc].value).strip()
                    if w: wrongs.append(w)
            if not corr: continue
            all_opts = [corr] + wrongs
            all_opts = [x for x in all_opts if x]
            if len(all_opts) >= 2:
                data.append({'category': cat, 'unit': u, 'question': q,
                             'options': all_opts[:4], 'correct_text': corr, 'difficulty': diff_val})
        else:
            # الشكل القديم: الخيارات بين عمود الوحدة والسؤال
            start_opt = min(cols['u'], cols['q']) + 1
            end_opt = max(cols['u'], cols['q'])
            opt_cols = list(range(start_opt, end_opt))
            o_txt = []; o_is_colored = []
            for ci in opt_cols:
                if ci < len(row):
                    cell = row[ci]
                    val = str(cell.value if cell.value else "").strip()
                    o_txt.append(val); o_is_colored.append(is_cell_highlighted(cell))
            corr = ""
            for i, txt in enumerate(o_txt):
                if txt and o_is_colored[i]: corr = txt; break
            real_opts = [x for x in o_txt if x]
            if real_opts and corr:
                data.append({'category': cat, 'unit': u, 'question': q,
                             'options': real_opts[:4], 'correct_text': corr, 'difficulty': diff_val})
    del wb; del rows
    return pd.DataFrame(data)

def fetch_smart_questions(uploaded_file):
    uploaded_file.seek(0)
    if uploaded_file.name.lower().endswith('.xls'): return _read_xls_questions(uploaded_file)
    elif uploaded_file.name.lower().endswith('.xlsx'): return _read_xlsx_questions(uploaded_file)
    return pd.DataFrame()

def get_master_pattern_from_file(uploaded_file, limit=30):
    if uploaded_file is None: return [random.randint(0,3) for _ in range(limit)]
    try:
        uploaded_file.seek(0)
        wb = openpyxl.load_workbook(uploaded_file, data_only=False)
        sheet = wb.active
        pattern = []
        for row in sheet.iter_rows():
            found_in_row = -1
            for cell in row:
                if is_cell_highlighted(cell) and cell.value:
                    # Strip the cell value and remove common decorators
                    txt = str(cell.value).strip()
                    for ch in [')', '(', '.', '-', '،', ',', ' ']:
                        txt = txt.replace(ch, '')
                    # Normalize Arabic hamza/madda variants so أ/إ/آ → ا
                    txt_norm = txt.replace('أ', 'ا').replace('إ', 'ا').replace('آ', 'ا')
                    # Only match EXACT single-letter cells (not full Arabic words)
                    if txt_norm == 'ا': found_in_row = 0
                    elif txt_norm == 'ب': found_in_row = 1
                    elif txt_norm == 'ج': found_in_row = 2
                    elif txt_norm == 'د': found_in_row = 3
                    if found_in_row != -1: pattern.append(found_in_row); break
            # *** FIX: do NOT add a random entry for rows with no highlighted letter ***
            # (previously this line caused header/empty rows to push a random index,
            #  shifting every question to use the wrong model answer position)
        while len(pattern) < limit: pattern.append(random.randint(0,3))
        del wb
        return pattern
    except: return [random.randint(0,3) for _ in range(limit)]

def _sample_balanced(df, n):
    """سحب n سؤال بشكل متوازن من الفئات"""
    if df.empty or n == 0: return pd.DataFrame(columns=df.columns)
    grp = df.groupby('category'); cats = list(grp.groups.keys())
    random.shuffle(cats)
    base = n // len(cats); rem = n % len(cats)
    sel = []
    for i, cat in enumerate(cats):
        q = base + (1 if i < rem else 0)
        g = grp.get_group(cat)
        sel.append(g.sample(n=q) if len(g) >= q else g)
    res = pd.concat(sel)
    if len(res) < n:
        left = df[~df.index.isin(res.index)]
        need = n - len(res)
        if not left.empty: res = pd.concat([res, left.sample(n=need) if len(left) >= need else left])
    return res.sample(frac=1)

def generate_balanced_exam(all_data_df, total):
    """توليد نموذج واحد متوازن (للاستخدام عند السماح بالتكرار)"""
    if all_data_df.empty: return pd.DataFrame()
    res = _sample_balanced(all_data_df, total)
    return res.reset_index(drop=True)

def generate_balanced_exam_with_mix(all_data_df, total, easy_pct, mid_pct, hard_pct):
    """
    توليد نموذج متوازن مع توزيع حسب نسب الصعوبة المحددة.
    easy_pct + mid_pct + hard_pct = 100
    """
    if all_data_df.empty: return pd.DataFrame()

    easy_n = round(total * easy_pct / 100)
    mid_n  = round(total * mid_pct  / 100)
    hard_n = total - easy_n - mid_n  # الباقي للصعب لضمان مجموع صحيح

    df_easy = all_data_df[all_data_df['difficulty'] == 'سهل']
    df_mid  = all_data_df[all_data_df['difficulty'] == 'متوسط']
    df_hard = all_data_df[all_data_df['difficulty'] == 'صعب']
    df_other = all_data_df[all_data_df['difficulty'] == 'غير محدد']

    parts = []
    shortfalls = 0  # النقص بسبب قلة الأسئلة

    def take(df_pool, needed):
        nonlocal shortfalls
        if needed <= 0 or df_pool.empty:
            shortfalls += max(0, needed)
            return pd.DataFrame(columns=all_data_df.columns)
        got = _sample_balanced(df_pool, min(needed, len(df_pool)))
        shortfalls += needed - len(got)
        return got

    parts.append(take(df_easy, easy_n))
    parts.append(take(df_mid,  mid_n))
    parts.append(take(df_hard, hard_n))

    # تعويض النقص من "غير محدد" ثم من الكل
    if shortfalls > 0:
        used = set(pd.concat([p for p in parts if not p.empty]).index) if any(not p.empty for p in parts) else set()
        backup = df_other[~df_other.index.isin(used)]
        if backup.empty:
            backup = all_data_df[~all_data_df.index.isin(used)]
        parts.append(take(backup, shortfalls))

    non_empty = [p for p in parts if not p.empty]
    if not non_empty: return pd.DataFrame()
    res = pd.concat(non_empty).iloc[:total]
    return res.sample(frac=1).reset_index(drop=True)

def _distribute_no_repeat(pool_df, n_per_model, num_models):
    """توزيع n_per_model سؤال لكل نموذج من pool_df بدون تكرار قدر الإمكان"""
    pool = pool_df.sample(frac=1).reset_index(drop=True)
    buckets = []
    for m in range(num_models):
        start = m * n_per_model
        end   = start + n_per_model
        buckets.append(pool.iloc[start:end] if start < len(pool) else pd.DataFrame(columns=pool_df.columns))
    return buckets

def generate_all_unique_exams(all_data_df, total_per_model, num_models=6,
                               easy_pct=0, mid_pct=0, hard_pct=0):
    """
    توليد num_models نماذج بدون تكرار.
    إذا كانت نسب الصعوبة محددة (مجموعها=100) تُطبَّق مع ضمان عدم التكرار.
    """
    if all_data_df.empty:
        return [pd.DataFrame()] * num_models

    use_mix = (easy_pct + mid_pct + hard_pct == 100 and easy_pct > 0)

    if use_mix:
        # حساب عدد أسئلة كل مستوى لكل نموذج
        easy_n = round(total_per_model * easy_pct / 100)
        mid_n  = round(total_per_model * mid_pct  / 100)
        hard_n = total_per_model - easy_n - mid_n

        df_easy  = all_data_df[all_data_df['difficulty'] == 'سهل'].reset_index(drop=True)
        df_mid   = all_data_df[all_data_df['difficulty'] == 'متوسط'].reset_index(drop=True)
        df_hard  = all_data_df[all_data_df['difficulty'] == 'صعب'].reset_index(drop=True)
        df_other = all_data_df[all_data_df['difficulty'] == 'غير محدد'].reset_index(drop=True)

        easy_bkts  = _distribute_no_repeat(df_easy,  easy_n, num_models)
        mid_bkts   = _distribute_no_repeat(df_mid,   mid_n,  num_models)
        hard_bkts  = _distribute_no_repeat(df_hard,  hard_n, num_models)

        results = []
        used = set()
        for m in range(num_models):
            parts = [b for b in [easy_bkts[m], mid_bkts[m], hard_bkts[m]] if not b.empty]
            df_m = pd.concat(parts) if parts else pd.DataFrame(columns=all_data_df.columns)
            used.update(df_m.index.tolist())
            results.append(df_m)

        # تكملة النقص من "غير محدد" ثم من الكل
        unused = all_data_df[~all_data_df.index.isin(used)].sample(frac=1).reset_index(drop=True)
        if unused.empty:
            unused = all_data_df.sample(frac=1).reset_index(drop=True)
        ptr = 0
        for m in range(num_models):
            if len(results[m]) < total_per_model:
                need = total_per_model - len(results[m])
                extra = unused.iloc[ptr:ptr+need]
                ptr += need
                results[m] = pd.concat([results[m], extra])
            results[m] = results[m].iloc[:total_per_model].sample(frac=1).reset_index(drop=True)
        return results

    else:
        # التوزيع حسب الفئات (بدون نسب صعوبة)
        grp = all_data_df.groupby('category')
        cats = list(grp.groups.keys())
        random.shuffle(cats)
        base_per_cat = total_per_model // max(len(cats), 1)
        rem = total_per_model % max(len(cats), 1)

        model_buckets = [[] for _ in range(num_models)]
        used_indices = set()

        for i, cat in enumerate(cats):
            q_per_model = base_per_cat + (1 if i < rem else 0)
            df = grp.get_group(cat).sample(frac=1)
            needed_total = q_per_model * num_models
            if len(df) >= needed_total:
                for m in range(num_models):
                    chunk = df.iloc[m * q_per_model : (m+1) * q_per_model]
                    model_buckets[m].append(chunk)
                    used_indices.update(chunk.index)
            else:
                per_m = max(1, len(df) // num_models)
                for m in range(num_models):
                    start = m * per_m
                    if start < len(df):
                        chunk = df.iloc[start:min(start+per_m, len(df))]
                        model_buckets[m].append(chunk)
                        used_indices.update(chunk.index)

        unused_pool = all_data_df[~all_data_df.index.isin(used_indices)].sample(frac=1)
        results = []
        for m in range(num_models):
            result_df = pd.concat(model_buckets[m]) if model_buckets[m] else pd.DataFrame(columns=all_data_df.columns)
            if len(result_df) < total_per_model and not unused_pool.empty:
                need = total_per_model - len(result_df)
                extra = unused_pool.iloc[:need]
                unused_pool = unused_pool.iloc[need:]
                result_df = pd.concat([result_df, extra])
            result_df = result_df.iloc[:total_per_model].sample(frac=1).reset_index(drop=True)
            results.append(result_df)
        return results


# ===================== دوال بناء بنك الأسئلة =====================

def extract_pdf_content(pdf_bytes):
    """استخراج محتوى PDF - نص أو صور حسب النوع"""
    if not PYMUPDF_OK:
        return None, None, "يجب تثبيت PyMuPDF: py -m pip install PyMuPDF"
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        if len(full_text.strip()) > 200:
            return full_text, None, None  # PDF نصي
        else:
            # PDF مصوّر - استخرج الصفحات كصور
            images = []
            for i, page in enumerate(doc):
                if i >= 15: break  # أول 15 صفحة فقط
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                images.append(pix.tobytes("png"))
            return None, images, None  # PDF مصوّر
    except Exception as e:
        return None, None, str(e)

def build_ai_prompt(unit, num_q, difficulty):
    """بناء الـ prompt لتوليد الأسئلة"""
    diff_map = {"سهل": "سهلة ومباشرة", "متوسط": "متوسطة المستوى", "صعب": "صعبة وتحليلية", "مختلطة": "متنوعة المستويات"}
    diff_text = diff_map.get(difficulty, "متنوعة")
    return f"""أنت متخصص في إنشاء أسئلة اختيار من متعدد باللغة العربية للمقررات الجامعية الهندسية.

المطلوب: أنشئ {num_q} سؤالاً {diff_text} بناءً على المحتوى التعليمي المرفق.

شروط يجب الالتزام بها:
- الأسئلة باللغة العربية الفصحى فقط
- كل سؤال له 4 خيارات: إجابة صحيحة واحدة + 3 إجابات خاطئة معقولة ومقنعة
- الإجابات الخاطئة يجب أن تكون قريبة من الصحيحة لتصعيب التخمين
- الوحدة: {unit}
- المستوى: {difficulty}
- لا تكرر أي سؤال

أخرج النتيجة بصيغة JSON فقط، بدون أي نص قبلها أو بعدها:
[
  {{
    "السؤال": "نص السؤال",
    "الإجابة الصحيحة": "الإجابة الصحيحة",
    "خطأ1": "إجابة خاطئة 1",
    "خطأ2": "إجابة خاطئة 2",
    "خطأ3": "إجابة خاطئة 3",
    "الوحدة": "{unit}",
    "الصعوبة": "{difficulty}"
  }}
]"""

def parse_ai_questions(text):
    """استخراج الأسئلة من استجابة الذكاء الاصطناعي"""
    try:
        match = re.search(r'\[.*?\]', text, re.DOTALL)
        if match:
            return json.loads(match.group()), None
        return [], "لم يتم العثور على أسئلة في الاستجابة"
    except Exception as e:
        return [], f"خطأ في تحليل الاستجابة: {str(e)}"

def call_gemini(api_key, prompt, text_content=None, images=None):
    """استدعاء Gemini API - يستخدم المكتبة الجديدة google-genai"""
    if not GEMINI_OK:
        return None, "يجب تثبيت المكتبة: py -m pip install google-genai"
    try:
        from google import genai as _genai
        # استخدام v1alpha لأوسع نطاق من النماذج المجانية
        client = _genai.Client(
            api_key=api_key,
            http_options={"api_version": "v1alpha"}
        )
        _model = "gemini-2.0-flash-lite"

        if text_content:
            full_prompt = f"{prompt}\n\nالمحتوى التعليمي:\n{text_content[:30000]}"
            response = client.models.generate_content(model=_model, contents=full_prompt)
        elif images:
            import PIL.Image
            parts = [prompt]
            for img_bytes in images[:10]:
                parts.append(PIL.Image.open(io.BytesIO(img_bytes)))
            response = client.models.generate_content(model=_model, contents=parts)
        else:
            response = client.models.generate_content(model=_model, contents=prompt)
        return response.text, None
    except Exception as e:
        err_str = str(e)
        if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
            return None, "تجاوزت الحصة المجانية اليومية لـ Gemini. انتظر حتى الغد أو فعّل الدفع من aistudio.google.com"
        return None, f"خطأ Gemini: {err_str}"

def call_claude(api_key, prompt, text_content=None, images=None):
    """استدعاء Claude API"""
    if not CLAUDE_OK:
        return None, "يجب تثبيت المكتبة: py -m pip install anthropic"
    try:
        client = anthropic_lib.Anthropic(api_key=api_key)
        if text_content:
            full_prompt = f"{prompt}\n\nالمحتوى التعليمي:\n{text_content[:15000]}"
            msgs = [{"role": "user", "content": full_prompt}]
        elif images:
            parts = [{"type": "text", "text": prompt}]
            for img_bytes in images[:5]:
                b64 = base64.standard_b64encode(img_bytes).decode()
                parts.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}})
            msgs = [{"role": "user", "content": parts}]
        else:
            msgs = [{"role": "user", "content": prompt}]
        resp = client.messages.create(model="claude-opus-4-6", max_tokens=4096, messages=msgs)
        return resp.content[0].text, None
    except Exception as e:
        return None, f"خطأ Claude: {str(e)}"

def call_openrouter(api_key, prompt, text_content=None, images=None,
                    model="qwen/qwen3.6-plus:free"):
    """استدعاء OpenRouter API - يدعم عشرات النماذج المجانية"""
    import urllib.request, json as _json
    try:
        if text_content:
            full_prompt = f"{prompt}\n\nالمحتوى التعليمي:\n{text_content[:30000]}"
        else:
            full_prompt = prompt

        payload = _json.dumps({
            "model": model,
            "messages": [{"role": "user", "content": full_prompt}],
            "max_tokens": 4096,
        }).encode("utf-8")

        req = urllib.request.Request(
            "https://openrouter.ai/api/v1/chat/completions",
            data=payload,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://localhost",
                "X-Title": "Exam Generator",
            },
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        text = data["choices"][0]["message"]["content"]
        return text, None
    except Exception as e:
        err = str(e)
        if "401" in err: return None, "مفتاح OpenRouter غير صحيح"
        if "429" in err: return None, "تجاوزت الحد المسموح — انتظر قليلاً ثم أعد المحاولة"
        return None, f"خطأ OpenRouter: {err}"

def questions_to_excel(questions_list):
    """تحويل قائمة الأسئلة إلى Excel بتنسيق متوافق مع التطبيق"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "بنك الاسئلة"
    headers = ["السؤال", "الإجابة الصحيحة", "خطأ1", "خطأ2", "خطأ3", "الوحدة", "الصعوبة"]
    header_fill = PatternFill(start_color="004d40", end_color="004d40", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    correct_fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='right')
    for r, q in enumerate(questions_list, 2):
        ws.cell(row=r, column=1, value=q.get("السؤال", ""))
        correct_cell = ws.cell(row=r, column=2, value=q.get("الإجابة الصحيحة", ""))
        correct_cell.fill = correct_fill  # تظليل أخضر للإجابة الصحيحة
        ws.cell(row=r, column=3, value=q.get("خطأ1", ""))
        ws.cell(row=r, column=4, value=q.get("خطأ2", ""))
        ws.cell(row=r, column=5, value=q.get("خطأ3", ""))
        ws.cell(row=r, column=6, value=q.get("الوحدة", ""))
        ws.cell(row=r, column=7, value=q.get("الصعوبة", ""))
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = 35
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

# ===================== توليد مسائل رياضية برمجياً =====================

def _wrong_answers(correct, count=3, spread=None):
    """توليد إجابات خاطئة قريبة من الإجابة الصحيحة"""
    import math as _math
    correct = round(correct, 4)
    if spread is None:
        spread = max(1, abs(correct) * 0.3)
    wrongs = set()
    attempts = 0
    while len(wrongs) < count and attempts < 200:
        attempts += 1
        delta = random.uniform(-spread, spread)
        if abs(delta) < 0.001 * max(1, abs(correct)):
            delta = spread * random.choice([-1, 1])
        candidate = round(correct + delta, 4)
        if candidate != correct and candidate not in wrongs:
            # تقريب جميل
            if correct == int(correct):
                candidate = int(round(candidate))
                if candidate == int(correct): continue
            wrongs.add(candidate)
    while len(wrongs) < count:
        wrongs.add(correct + (len(wrongs)+1) * random.choice([-1, 1]))
    return [str(w) for w in list(wrongs)[:count]]

def generate_math_questions(topic, num_q, difficulty, unit_name):
    """
    يولّد أسئلة اختيار من متعدد في الرياضيات برمجياً بدون ذكاء اصطناعي.
    المواضيع المدعومة: جمع/طرح/ضرب/قسمة، كسور، مساحات، قوى وجذور، نسب مئوية، إحصاء
    """
    import math as _math
    questions = []

    # قاموس قوالب الأسئلة حسب الموضوع
    def q_arith():
        if difficulty == "سهل":
            a, b = random.randint(10,99), random.randint(10,99)
            op = random.choice(['+', '-', '×'])
        elif difficulty == "متوسط":
            a, b = random.randint(100,999), random.randint(10,99)
            op = random.choice(['+', '-', '×', '÷'])
        else:
            a, b = random.randint(100,9999), random.randint(11,99)
            op = random.choice(['×', '÷'])
        if op == '+':  ans = a + b
        elif op == '-': a, b = max(a,b), min(a,b); ans = a - b
        elif op == '×': ans = a * b
        else:
            b = random.randint(2,12)
            a = b * random.randint(2,50)
            ans = a // b
        text = f"ما حاصل العملية الحسابية التالية: {a} {op} {b} ؟"
        return text, str(ans), _wrong_answers(ans, 3)

    def q_fraction():
        d1, d2 = random.randint(2,9), random.randint(2,9)
        n1, n2 = random.randint(1,d1-1) if d1>1 else 1, random.randint(1,d2-1) if d2>1 else 1
        if difficulty == "سهل":
            # جمع كسور بمقام مشترك
            d2 = d1
            n2 = random.randint(1, max(1, d1-n1-1)) if d1 > n1 else 1
            ans_n = n1 + n2; ans_d = d1
        else:
            ans_n = n1*d2 + n2*d1; ans_d = d1*d2
        from math import gcd
        g = gcd(abs(ans_n), abs(ans_d))
        ans_n //= g; ans_d //= g
        ans_str = f"{ans_n}/{ans_d}" if ans_d != 1 else str(ans_n)
        wrongs = [f"{ans_n+random.randint(1,3)}/{ans_d}", f"{ans_n}/{ans_d+1}", f"{n1+n2}/{d1+d2}"]
        text = f"ما ناتج جمع الكسرين: {n1}/{d1} + {n2}/{d2} ؟"
        return text, ans_str, wrongs

    def q_area():
        shapes = {
            "مستطيل": lambda: (random.randint(3,20), random.randint(3,20)),
            "مثلث":   lambda: (random.randint(4,20), random.randint(4,20)),
            "دائرة":  lambda: (random.randint(2,10), None),
        }
        shape = random.choice(list(shapes.keys()))
        if shape == "مستطيل":
            l, w = shapes[shape]()
            ans = l * w
            text = f"ما مساحة المستطيل الذي طوله {l} وحدة وعرضه {w} وحدة؟"
        elif shape == "مثلث":
            b, h = shapes[shape]()
            ans = 0.5 * b * h
            text = f"ما مساحة المثلث الذي قاعدته {b} وحدة وارتفاعه {h} وحدة؟"
        else:
            r, _ = shapes[shape]()
            ans = round(_math.pi * r * r, 2)
            text = f"ما مساحة الدائرة التي نصف قطرها {r} وحدة؟ (π ≈ 3.14)"
            ans = round(3.14 * r * r, 2)
        return text, str(ans), _wrong_answers(ans, 3)

    def q_power():
        if difficulty == "سهل":
            base, exp = random.randint(2,9), random.randint(2,3)
            ans = base ** exp
            text = f"ما قيمة {base}^{exp} ؟"
        elif difficulty == "متوسط":
            n = random.choice([4,9,16,25,36,49,64,81,100,121,144])
            ans = int(_math.sqrt(n))
            text = f"ما الجذر التربيعي للعدد {n} ؟"
        else:
            base, exp = random.randint(2,5), random.randint(3,4)
            ans = base ** exp
            text = f"ما قيمة {base}^{exp} ؟"
        return text, str(ans), _wrong_answers(ans, 3)

    def q_percent():
        total = random.randint(50, 500)
        pct   = random.choice([10, 20, 25, 30, 40, 50, 75])
        ans   = total * pct // 100
        text  = f"ما هي نسبة {pct}% من العدد {total}؟"
        return text, str(ans), _wrong_answers(ans, 3)

    def q_stats():
        n = random.randint(4, 7)
        nums = [random.randint(10, 99) for _ in range(n)]
        op = random.choice(["المتوسط", "المدى"])
        if op == "المتوسط":
            ans = round(sum(nums)/len(nums), 2)
            text = f"ما متوسط الأعداد التالية: {', '.join(map(str,nums))} ؟"
        else:
            ans = max(nums) - min(nums)
            text = f"ما مدى مجموعة الأعداد: {', '.join(map(str,nums))} ؟"
        return text, str(ans), _wrong_answers(ans, 3)

    # خريطة الموضوع للدالة
    topic_map = {
        "عمليات حسابية (+−×÷)": q_arith,
        "كسور":                   q_fraction,
        "مساحات أشكال هندسية":   q_area,
        "قوى وجذور":             q_power,
        "نسب مئوية":              q_percent,
        "إحصاء (متوسط ومدى)":    q_stats,
    }
    gen_func = topic_map.get(topic, q_arith)

    for _ in range(num_q):
        try:
            q_text, correct, wrongs = gen_func()
            questions.append({
                "السؤال":          q_text,
                "الإجابة الصحيحة": correct,
                "خطأ1":            wrongs[0] if len(wrongs) > 0 else "—",
                "خطأ2":            wrongs[1] if len(wrongs) > 1 else "—",
                "خطأ3":            wrongs[2] if len(wrongs) > 2 else "—",
                "الوحدة":          unit_name,
                "الصعوبة":         difficulty,
            })
        except Exception:
            continue
    return questions

# ===================== واجهة المستخدم =====================

st.markdown("""<div class="header-container"><div class="logo-box"><img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAARMAAAEqCAYAAAA23LTdAAAACXBIWXMAAA7EAAAOxAGVKw4bAAAF8WlUWHRYTUw6Y29tLmFkb2JlLnhtcAAAAAAAPD94cGFja2V0IGJlZ2luPSLvu78iIGlkPSJXNU0wTXBDZWhpSHpyZVN6TlRjemtjOWQiPz4gPHg6eG1wbWV0YSB4bWxuczp4PSJhZG9iZTpuczptZXRhLyIgeDp4bXB0az0iQWRvYmUgWE1QIENvcmUgNi4wLWMwMDIgNzkuMTY0NDYwLCAyMDIwLzA1LzEyLTE2OjA0OjE3ICAgICAgICAiPiA8cmRmOlJERiB4bWxuczpyZGY9Imh0dHA6Ly93d3cudzMub3JnLzE5OTkvMDIvMjItcmRmLXN5bnRheC1ucyMiPiA8cmRmOkRlc2NyaXB0aW9uIHJkZjphYm91dD0iIiB4bWxuczp4bXA9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC8iIHhtbG5zOmRjPSJodHRwOi8vcHVybC5vcmcvZGMvZWxlbWVudHMvMS4xLyIgeG1sbnM6cGhvdG9zaG9wPSJodHRwOi8vbnMuYWRvYmUuY29tL3Bob3Rvc2hvcC8xLjAvIiB4bWxuczp4bXBNTT0iaHR0cDovL25zLmFkb2JlLmNvbS94YXAvMS4wL21tLyIgeG1sbnM6c3RFdnQ9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC9zVHlwZS9SZXNvdXJjZUV2ZW50IyIgeG1wOkNyZWF0b3JUb29sPSJBZG9iZSBQaG90b3Nob3AgMjEuMiAoV2luZG93cykiIHhtcDpDcmVhdGVEYXRlPSIyMDIxLTAyLTA4VDExOjI2OjE2KzAyOjAwIiB4bXA6TW9kaWZ5RGF0ZT0iMjAyMS0wMi0wOVQwMDozOTo0NyswMjowMCIgeG1wOk1ldGFkYXRhRGF0ZT0iMjAyMS0wMi0wOVQwMDozOTo0NyswMjowMCIgZGM6Zm9ybWF0PSJpbWFnZS9wbmciIHBob3Rvc2hvcDpDb2xvck1vZGU9IjMiIHBob3Rvc2hvcDpJQ0NQcm9maWxlPSJzUkdCIElFQzYxOTY2LTIuMSIgeG1wTU06SW5zdGFuY2VJRD0ieG1wLmlpZDozMWVjODIwMS1jODdlLTJmNGQtYTliYi0wNWJhMzAxZjlhZmYiIHhtcE1NOkRvY3VtZW50SUQ9ImFkb2JlOmRvY2lkOnBob3Rvc2hvcDpmYTE5YzI5Zi0wMDQyLTMwNGEtOGQ4ZC01MjkzZmI1ZmQzOWIiIHhtcE1NOk9yaWdpbmFsRG9jdW1lbnRJRD0ieG1wLmRpZDowNjBiNDU0Yi0yMDQ3LTIyNDgtYWMxZC1mMjE5MjcwYmVkNDkiPiA8eG1wTU06SGlzdG9yeT4gPHJkZjpTZXE+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJjcmVhdGVkIiBzdEV2dDppbnN0YW5jZUlEPSJ4bXAuaWlkOjA2MGI0NTRiLTIwNDctMjI0OC1hYzFkLWYyMTkyNzBiZWQ0OSIgc3RFdnQ6d2hlbj0iMjAyMS0wMi0wOFQxMToyNjoxNiswMjowMCIgc3RFdnQ6c29mdHdhcmVBZ2VudD0iQWRvYmUgUGhvdG9zaG9wIDIxLjIgKFdpbmRvd3MpIi8+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJzYXZlZCIgc3RFdnQ6aW5zdGFuY2VJRD0ieG1wLmlpZDozMWVjODIwMS1jODdlLTJmNGQtYTliYi0wNWJhMzAxZjlhZmYiIHN0RXZ0OndoZW49IjIwMjEtMDItMDlUMDA6Mzk6NDcrMDI6MDAiIHN0RXZ0OnNvZnR3YXJlQWdlbnQ9IkFkb2JlIFBob3Rvc2hvcCAyMS4yIChXaW5kb3dzKSIgc3RFdnQ6Y2hhbmdlZD0iLyIvPiA8L3JkZjpTZXE+IDwveG1wTU06SGlzdG9yeT4gPC9yZGY6RGVzY3JpcHRpb24+IDwvcmRmOlJERj4gPC94OnhtcG1ldGE+IDw/eHBhY2tldCBlbmQ9InIiPz725RFTAABaG0lEQVR4nO2dd5gkV3W331tVnSfvzO5szlFhhQQIZLAwGPMRjMEYgzBBAmxjAzbYGNsSDtjCmGiSkBASQQEUEEEgWGlBMkYJ5QQoorRabd6d3N0VzvfHrdtd09OzO7PTsz2zc9/n6Wemu6urbqVfnXPuuecqEcFisVimitPsBlgslqMDKyaWWlqA3wdWNbshltmFN9UVKKUa0Q7LDEBEWkAG/GIfXqoF5XqrlVK/bXa7LNNHI8Mc1jKxACAiDjAQlvbyxM2fYO+vbwZ4V5ObZZlFWDGxGDpA2PfQefTf/Tn23XgZ+P6ZInJcsxtmmR1YMbEgIjlgL8XfEv32HLL7irj3/Jjgf68FuE9ErC9rOSRWTCwAzyMaQJ45iw4O0FmG/K5tRF/9EFEYAKxocvssswArJnMcEcmD/LzU/y2UXEsmDOj0PVpLIengMfp3PwtwcrPbaZn5WDGxnCLiU+z/OkQHQClkQNHiODAS0PfsEwSlkW+LyEua3VDLzMaKyRxGRBZBuJV9X6Fl6B5w00Qpn6DXZyQdQeCy/EO/S+narwHcICJuk5tsmcFYMZnbnCDBfkYOXIEQQJRCIgcRQADS0AbPXP1Vyv17AV7W1NZaZjRWTOYoItIBcs1I/3X4/p2IEwJZkFRiqQzFMswb2c7wfTeAyLUi0t6kJltmOFZM5iCxu7I/CnczvOd/cJwR3EoutIcCIgFEUSy00+L3Mfz5dzD82/sA3t6kZltmOFZM5iZrkDLF/ZeTlgdJu6BHRYwAAUo5IA44IUMiqGKZBcPDyK1XI1H4BRFZ2OT2W2YgVkzmGCKSAh6Myr/B3/UfpKWMF2WIQiAaBqeEIxEiEThl8mmHdAbcMvC9zzHyxEMAr2/qTlhmJFZM5h4rkTL+vvPIO/tJhw4qbEUkDQ44HuCAKMARvFIfhEAGckP7Gdn6dSK/dI6IHNPk/bDMMKyYzCHiWMlDkf8k/p6rSDkKJwJHPJTkQWVRLkRuHDNRLnkRSmXAyeAoIXvdFxm492cAD8SDAy0WwIrJXOMYpMjQngvJOSGMOKBAomGiyIEoRxRRtVCiDG4A6WyBkBQ4UCiVSN9+FcFwP8Apzd0dy0zCiskcQUQKIPeODF8PwxchfhnIASHKGQRVhigDEVRKXEQOvq//9cUHBYQwsuVrDD90O8Av7CBAi8GKydzhOImG6N/9GTxnL64KtPkhIUIEDIESIkkjERABEqEyWRTgUIIU0AJdaYf9F/0bft9egMgKigWsmMwJYqvkFn/4l2S4C8IyKhNpa0RpVwdHwBlCSEEIEgLKp0yWUmmItAsIFAcB5dK7/T52X32u2USuWftmmTlYMTnKia2GQQn3MrL7HAqUcCKInDJIEXBBwFGAO4w4QiRZVARIgHJc/AAdS/EcigLRiE8mHCF3+w8o7XgSbFexBSsmc4EOEEr7f0i6dB2OP4LnQhBB6DmgvHgcDkAEygcnipPYHFx88rkUUdxd3NGZx3E9/IEA98lfceDHF0AYXCIiK5u2h5YZgRWTo5/jJOzHP/Adcu4QTgQIRMoDN02oFKgUCkC54PrglLXrE2UIRvoh5SKpNCOlCH9kmCiKSOXStKVdMjdcSP+dWwF+a7uK5zb25B/FiMgHIfr5YP89eDwGaI0IfIAMQRTHTSU2O8QDpTUFBUiafBZGimVK4uKlPVQIIJDKMTI4hDuwm/KdP0aCMsCaI7+XlpmCmmqpezvVxcxFRB5m92lr6bsBPxgiUoKL4IYhSiJwhYCAyNPdwZmoAx7zKf5yiNK90F726M+nyAcBXimKM2FD+gpQ9qCzBN4AkG9j4EOX0HrKHwI4Sik7TeQswU51YTkkIpIG1tJ/H2F5hDAMEQlBRfoCiiIIw0Oux5VI/+MAngAOmQByPoQKwjSQchi+6SeEw4NgSzzOWayYHL2sjaISFJ8higKUAsfRlqRSisoHh8CLdNJJmBJd6kRBtgS5MgRAMQ9DwwNkbr6W8sP3AtxiYydzE3vSj0JEZCkSPFAauAEcwXUdrRsqqlgnuAkxqXiq0Zh1uYSIEnw3ouzG34cObuCAQOAAUUiHGsb//lco7dsB8Lrp3kfLzMOKyVFGnFfyVFh6gvKud0DkI/igAkAQU14giiAKJrDGiNCJiJKhMaUAFy9I4Q1Doa2FoH8H0V1XEzx0G8BVIpKdht2zzGCsmBx9dENEce/3KKR2EUY+IiHasxEcV/+FKB4afHCCuGfHjSAdOoAi9BwC1yETOhTCNOHeQcppKBSK9J33YYLtjwO8aXp30zLTsGJyFBEXPtoVlbcR7v88XgiOIyjj4hACke76dZjQ2Q8cUAKZ0EFFitCBYioicOLendDBLTik28Evl+jZ/xTBz64CkW/oOrOWuYIVk6OLY5ARhnd/lqy3A8qgUg5KCVEUEoZCEERINDY2Mh5RbJUQKIg8fAdCR3DxQQIQn6AcEQSQEUipkKHvfInhx38D8Nbp2lHLzMOKyVGCiHjA3cP7fwojnyedCiFohdBHCCudN5UOnDi/wHF0+KTyeRhSSR1SSvcIR+jENuUQOg6uRPozJ4R0hGQgq8AtAVGA4xUp3rEVyiNfFJHfO4KHwdJErJgcPWySqIj0fZG0MrHVjuq3B8ktrOYtOZX35jMvTr9HOVpPBNKh/osHpbRQVkpnsYUQeQ6BGuHA9z5HsOUigOttiYK5gRWTowARyQH3lvb9nEx4EyqOrYYcOsBqRENPvCVjvswEaCFytcuTDmOBccB3IXSAUCBSRMolctO051P0+H143/4YoV8GWN/A3bXMUKyYHB38buTvQw78B54M4/rgOS34auDgv5Kav3Eym1K6zAlEqEhnuvpuAER4YYQS8GMxSQMpgcjzKKZciiUfd3CYQjAEpZ0ceOQegD+elr22zCismMxydOEjtpT2XUPWu0vnnUUpBBfxEmIyxkjRp97ER5RyKmIyKjFWIHLAd0Mc/LgCm0Oo4nX6+m/oCiqXJo2LWyzj+GXKUZnwlu8R7NvxMRGxNU+OcqyYzH5eHpV3Eg1+iSgoEoWgHI8gLOF64/Ta1Ml2rwz4MpZJJQir3RtREDqRDpZEOmHNiXSVtggI/CJuGJJ2XfA8yIJKQ27rBQzeeBXAd+Oua8tRihWTWYyI9EL0veH9V+CphwmDOFbqFkEVJ7CGQ5/+MCEqgQu4QOTh+SncMEWYdSEHqRBSQyNQKlGOfEYiKPvQemAPg1d8ksFnHgN42+Huq2XmY8VkdvOGyN9DOHwRod+Pq9IoFJEjeGmI/EOvwPTcSKILp/KvCJFycCQuTO/E4iIKAg8lLv3lkMgDzwUlDrgeUS6FZCClE2bpjIoMPvEgwIUisnjajoalqdh6JrMUEVmGhE/u3f52MuG3yPlxnkfQEncDO0CKA7kRMl5IUCqRVaCiHEo6GPIVTk4RZfsplgZo9yBTbIVtbYQPtvPkvQEdmSV0PXU95BVDRUGcNDkvi1v2IShBMdIRWAeKafDTHpFSpEPIhIJTCtjjddFdHoFUjqF/u4rCiS/5mFLqI009eJYKjaxn4jVsTZYjRjwz35NB8TG8oevJOOAGgOTiKvMlQjfQxdMA10mRUlm8MAVOgBSfpS0LZFyGilkysgpfXozb+UroXIr3e89hVRCy947/Zc+lz1De/ihdmRRZAb+vHxHwCh5kMuA6lEaGUGVoiQLCEMoRlFIKJ58l6w6ACghDBU8+gGx+0VkicoFS6ommHkRLw7GWySxERE4hGrhp6PH3kC1/D1d8EIU4EDo+oVONsboeyAikfT2FBU7s+7RBeQBG1PHkl34BJ3scbqbLbOIEtM1xG8N99N9+LSPXfY3Cb/6PQlhCZT3KQchIpAjFIaMcCgoIyzqdNpMmUNBXLJPNQ6oMKvLY7fTQ+i/fotVaJzOGRlomVkxmGfHQ/pHy3stI73wPhAOAA15AkNID8yIBJ/JwIxenWMJNoQOncTcubhtBtJyg673QeirZ1vXEvlEBKCulgnhbHnoK0J9TGmbgtusY2PpN8g9fT8ofIKUER3S2red4OCpNUCoTRgGplIOT8iBVpjgC2VQLhFn8416M988XoXItpyilbmnGMbRUsWIyhxGR5xEN3bb3V69mXuY2KJZ0hlkqoJTSAVI38PCCjC65KBF4vp7bwvEoqeNg/sdwCs8lle0xq20DhpRSdfuS4xKQxwJ3SmmYvluvZviK/6Rn54OkoggCLWC+SiGOS9Z1tZVS9KEVigGUBj3axUMyGQ584Et0vvytYOvFNh0rJnOUOG1+eGjfpUQ7/5xWZ0SPiVEuoRvhewKSIhU4uDKkfxN1Eqos5fQx0HEGbsvvkMkvB3gtsBUoTfSGjmM1xwN3URxgcOtFjFz+Cdr7d+gZ/4iIghARcF0ARRQJYUsHSrJ4fQMgEUOnvJr0351PqqXzdUqpHzT+SFkmihWTOYqI/JmE2y7Z8dgJ9Ob3ovrRwREHfMchJAtRjlTk43IAFAz2/Bchx5JvewGpTA/AO4ArlVIjU2hHFng+8HPKIxRvvY6+r/4T3f2P46ZCKAUMR+BkIVvupE+VcLIemZJP2o8YaWtn52v+lhVvPRPAHc8iskw/VkzmICIyD2TPnqc/RC74LAVABhXKE8TTCWW+gARdqHA+XvqFON4iosUfIp3qMKvJKqVKDWxTGtgA3Eu5SOmXP6bvik+Te+YespmIoFQiV5wH6SJ7gyFy7RmUD+VSxMiyzbR9+Fvkl659j1LqK41qk2VyWDGZg4jIa8Xf9YPh3/4+Ofd+ZAjcdI5IjRCmoCSQyTsMDm8m2/Up0i3Pw023gbZELgOC6bIA4mr0xwD3URph8OYf0n/5x2k/8ASFfsDth1TEkAe+DzlXEeTa6X/NWSx804cAWpRSQ9PRNsvBsWIyxxCR5UjwhL/3uzi7/xrX3Uvo64zUwMvgploYKS9CFd5Kpv21ZFvWA+olwK2NtEQm0E4XOA64m+IQw7+6lYGrP036gf+lMyrqivhlIfTBz8GejafS9hdfpm35pvcrpb50pNppqWLFZA4RP/VDSo8y/NgfkePXiAt+CspuB1H6fWRaTiOTX4VyswDvB77ezCd9LCrPAW6X8iDFh29j+Nufpf3u6/EcXeoRF/blXPzTPsb8P/w7lJtap5R6pFltnqs0UkwQkSm9LNOLiGyUqCRDT7xVBu9Hosc8KT22QPY88ceyb88V4vt9EnN8HMOYMYiIKyLPFRGR0rCEN/9Q+v/uJVL8oxaRP8pK8Gpk35tbpP+2n4iIfLbZ7Z2LTPX+H6UFVkxmLiKSExEp7vtfOfDwPCk/XZC+R18jA7uvl8AfNCJyqszwGfREJCsim0VEpDggw3dfK/s/8XaRP2wXeXVBhs5+u4QjgyIiL2x2W+caVkzmCCLymijYJwe2nSv7n/5L2f/s58UvPmtE5AWiM1RnDbE4Pr9iqdz4Q+l76wul/9/fJfvuu1VE5KPNbuNcw4rJHEFEPlws/kYG+h6QMCxKghnlzkyW2FJ5jhaVEQkP7JHBHU9LFIYitkTBEaWRYmIDsDMU0RXdI2EERQ7gROBBoHi0pKBLMk9Fs0Qp9UwTmzTnaKRBYMVkBhMLynOBXymlhpvdnuki3k9HKRU2uy1zDSsmFoulITRSTGZ0L4DFYpk9WDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyYWi6UhWDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyYWi6UhWDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyaWOYGIuCKyWUQ2NLstRytWTCxHNSLSISIvAYJisXiP75d/IyLvFpFUs9t2tGGnB7UcdcRCsQH4HRE5d2BggN/85jd885sXUihk+chZH6W9oxMgp5QqNre1zaWR04N6DVuTxdJERMQBFgAnA98rlUo89dRTXHrJZfzwRz+ir283mbTL0OBeOtpaOPMjH0MpNSIic15QGoW1TCyzGhFpBVYA90VRxJ49e7jxxhv57ne/y6233krKy1Mul3DdgCAcobWQxi8P87nPn88f/L/XARSUUsPN3Idm0kjLxIqJZdYhIi6wHNggIteUy2UeeughLr74Yq677jqGh4dRShEEAX7ZIZdLUy4P4KoA5QRk0g5LFy/k0suvZV73ApjDgmLFxDInEZEO4ATghiAI6Ovr48c//jFbtmzhpptuwnEc0uk0vu8TRRGpVArHd8EtE6gA33fxUu2U/SL5wiBv+tOX869nnYvnZWCOCoqNmVjmDCKSBpaig6nfHBgY4JFHHuHiiy9my5YtAJTLZQBc1yWKIkQEESEMQ5AIoghRDkoplCM4TsTQwDA/+tGPeMHzT+VVr3wHQBaYc2LSSKxlYplxxMHUhcBaYivkmWee4brrrmPLli38+te/JggCPM9jcHAQz/NwXRcRIYoiQF+XSilcCYjwiMQliCKclIvnRJRLg6RSIae++GS+fO5lZDJ5lFLzlFL7mrnvRxrr5liOSuJg6jrgjjAMGRoa4u677+aiiy7ipptuolQq4bouQRAQhiGZTAbH0RZHGIYEQQBoC0UpRRRFeKpIJDlQWUIJiGQEz3NwVRpPXHy/jzPP+gCnv/tDKOUAOEqpxt1hMxwrJpajBhHx0G7McSLyg2KxyPbt2/nud7/L97//ffbv38/Q0BDZbBbHcfB9H4AoiirXnrmGlVI4jlP5PooiHKcIksFx8jiegx8MxCJTQAUpstmIQn6IL5z7ZZ538itRSi1QSu1qxrFoBjZmYpn1iMh8YD3wf0EQsH//fm666SbOP/98nn76aUZGRgAdD3EcBxGhWCwiIuRyOaIoIgxDHRcBHMfBcZyKiCil8DwPRYYgBD8okvZSuG4KCUMQvWypGFAcGeS7V13K5ue8mHS6sFNE5pSgNAorJpYjhojkgDXASVEUfX1oaIjHHnuMc889lwceeIC9e/dSKpUq4qGUIpPJVH5vLJFSqUQYhqTTaTzPM+smDMOKkJiXSArXjQjCEYJyhON4KFJ4noMgIIpCfh7fu+qHPP/kU3j9G/4KYB5gxWSSWDfHMq3EwdSlaCvk2lKpxLPPPstVV13Fddddx+7du9m1axe5XI4wDFFKkUqlGBkZIZPJEARBLApaLFKpFI7jUCqV8DyvIjqmB8d13Uqvju/7SOSSzTkEYQnHcZEwQ+iHeCkFUQCRwlEejhph3aYu/ueLF7Jy5Qko5a5QSj3Z5MM37diYiWXGIyIFYDVwbxiG9Pf3c9ttt3Heeefx1FNP0dfXVwmSTms78EAFKAIQDyUeEIEq4yAoyRKFKaKwSK51hLe/4w186MP/g+N4MAfG7lgxscxI4mDqCuC5wLdLpRJPPvkkn/nMZ7j55pspl8sVC6NYLJLJZKb9+pF4YLyS+C8BKB+UFhckQ8prwZGIUrCXdGaIz3zmHF756tMAtUwp9fS0NrDJ2ACsZUYhIguA44Ctvu+za9cubrjhBr7xjW+wZ88e+vr6yGQylEoloiiipaUFz/MIw7ChF3M9IqeME2VBUqBii0SVibNREFyGRgaBiHyugOdm2bplK6f+3mvI5VqeEpGjXlAahRUTy2EhIllgFXBsFEWX9/f388ADD3DxxRdzyy23MDQ0BEAYhpXuWvO3WCwSBAGp1BEoKaLC+B8XJWhBQQBXWyZAJpcCFRCEwsC+Mtdccz3Hbb6At53+13he9ikRySqlStPf2NmNdXMsEyYeYNcLbAauKZfL7Nmzh+9973tceeWVbNu2rZKZajJSTbDUjJVJpVL4vk8QBBQKBYrF6Q1JhG4JJ8yhogyKCEcNgfKJSCGSJZQIvACIIPDIplvwCHCdfZx3wVc4+YWvxnHcTUqp30xrQ5uEdXMsRxQRaQFWAveFYciBAwe45557uPDCC3nooYfYuXMnuVyOVCpFEASUy+VK12wYhhQKBUBbJL7vV8RmeHi4Yq1MX+NTRCrCdUraIMGDKA0q3q4XoFxBBPwoxIt8JAwYCX2uvOwyjjv+hRRaun8tIouUUs9Ob2NnN9YysdQlDqb2AhtF5LpyucyTTz7JJZdcUunS9X2fTCZDuVzG87w449Qhm80ShiG+71cExXTrJgfjmeWnk1B5oIq4+KjIw4nyEGWIlBC5I6h0GZ+AMFCknQIqiEi5gishrgp5z3vfxF+9/+OmdyejlCpPa4OPMNYysUwb8TD/VcCdYRgyMDDAT3/6Uy666CK2bdvGvn37EBEymUxlnIzrunieVxkzUywWKzkjrutWxs84jkMQBBWXx3XdSgbr9O1QCtQIqBIodJwkyoHjgxoipAQoosghnU1zYGg3rZkMCkUUeFz9g0s58aTf4fkvfDWumzoWuGt6Gzx7sZaJxQzzXwCcLCJXFotFnnjiCc477zxuu+02du/ePe29LjOWIM+r/+gYPvofZ9PRthLluOuUUo80u1mNwuaZWBqCiHShM1NvDsOQ3bt38+1vf5urrrqKffv2VSyNuYwSj5bWQd77vtM54/SP4KVyL1dK/bTZ7WoUVkwsh008PmYJcKqIfHV4eJi77rqLCy64gIceeoj9+/dX3BHf93Fdt9lNbirZLOw/sIuTTtjIF79wHkuXH4vrpZ+rlLqz2W1rBFZMLJMinvphHrAJ+Jnv+zz11FNccMEF3H333Wzbtq3SRWuCpa7rkslkKkP+5yrKOYCr2lCh8La3/RF//48fJZNtQynnqOjdsQFYy4SIx8csB34VRREDAwNs3bqVr3/962zfvp2hoSGGh4fJZrOVcTLJgXKm2NBcJoyKRFELaWnhh1f/hOc8dx1/8Ioz8FL5DcCsF5NGYi2To4w4mLoYeB5weblc5oEHHuCrX/0q99xzDzt37sTzvEqhoXQ6XbFGHMchlUoRhiGlUgkRqQzxn6ukc4OUBvNkvU5K5Z087+QFfOGcS5g3by2umzpeKXV/s9s4FaybYxlFPMy/C51YdlsQBOzatYvvfOc7XHvttTzzzDOVwkJJa8PERGrrpoJOfT8iXbcznCA4QNrrJJ/u5MD+XbR0DHPGu/6QD3zwS3heHiCvlBppdjsPFysmFgBEJA/0AE+ICCMjI9x22218+ctf5tlnn2Xbtm2jksSSYmH+mvhIskqZsUbK5fKcD8A6DkRhSOA75NMFlDNEvq2fL593Pied+BpcN32CUureZrfzcLFiMoeJx8csAI4BrgvDkO3bt3PllVeydetWnn76aYaGhhARUqnUtNcLOdpRpIEijvKIwhQpx0XcXZz60vV86Ytb8LzcrK4bawOwc5B4fMwK4P4oiujr6+P666/nwgsv5JFHHqFYLFaqkAGV6mOWqeIQiQICgkCB6xGUPO69+zGuuPJznPbmD6OU14Et82gtk5lMPMy/HXiZiFxqgqnnnnsuDz74IM888wzpdLpSMzVZYNmIiT0/U0RSoEJcVxEGCs/LEAU+jnuAVWtauPzy/6NQ6MF1vfVKqYeb3dzJYt2co5g4mNqOHh9zRxRF7N69my9/+cts3bqV/fv3V+qiJiu3m3hHKpUinU4DzPkckUYgIroAdcqhXC4ionAkC9EIuZY+3vKWV3LmRy5CKRdm4RSjVkyOQmI3ZiHwsJnWYevWrZx33nkcOHCAvXv3VoKhw8P6ek2n0wRBUOl5MbkiyddcD6BOFeWERKGH47pEUiSKwFN5PHFxvH247k6+esEVPP+Fr8Nx3I1KqQeb3ebJYMXkKCERTN0kIluDIODxxx/nU5/6FHfccQelki7uFQQBxWKxYnnEv62IiOmtMZ+b70xpRMvho5yQwI/AUbhefNwlC6ED0SDtbT7Hn7CSL5xzNflCJ66bmlWCYsVkFiMiCmhDT/9wfxRF9Pf385Of/IRLL72UJ598knK5XLE4isUirutWXBfTO2OmdUjOF5O0UJJp8ZYpoAKiKATHJPA5EIEjKQiErvYMytnJBz/8ft7wp/+E46QBPKXUrFBx25szC4kzU3uBJ0WEcrlcqVZ2++23MzAwMGr6SyPSuVxOz/8SWxymTggkp8CsJpiZTNZkTonl8HGdFC2tGdo62li5ah3zexbS2d7GymWr6GzvZumSxXgpn+6eVqLIR6kUSqk5OYmXtUymkTiY2gqcAPxvGIbs3LmTK6+8ku9///vs2LGDcvnghbsUU8sTmepzR6Q6VscIlQn61iN5PTiOg6hq/CY5F7CxpkyFNvNbY0lVqrM51fmEk3MKm0Q7k9Fr4kPmszAM8TwPL572Jrl9UxXfJOd1dHTQ1dVNz/z59PTMp3fhQpYvX8m8eT0sWrqM9rZ2CoUW0uk0rutVjgHwFeApIAUcAG4DHptNOSfWzZnhxFbIQuLMVDPA7jvf+Q4PPvggAwMDBEEwobEvaornWg7z9FStHyqWTvJmNj1I9c5/8sYXFSSm6pRRmbgm6zZpaZkYj+d5+tiE0ajtmiCzKUqdy+Uqv03WnXUcR88SGPTT1dXFokVL6Oqax/wFvfT0zGfp0mUsWrSEnp755PIF8rk8mUwGL5XCcdzkfr0fbWXsRAvGADAI9B0NFeutmMxAYitkHrpy+9YwDHn88cc599xzufvuu3niiScqN1Ly5pxJiWX1zqV5iid7i0z7jfVQW54x2VWtnDD5JK98njwWtSSDyUpGX/DmN6Y9ruvS0dFBd3cPvQsXsmD+Anp6eljQ20tv70IWL1pIoaWFfK5AOpOpDHKM9/U/gG1AH7AP2Avsj9+PAP5siX0cLlZMZhBxl24P8FsTTL3yyiu47LJvs3//fgYG9PwxxlUYGRnBcZwxAdVx18/UCi47MjU3yYwiNpgcF+MupNPpygBCIxJJ8XCUP0o0zXfme1N0OnGDV5YXEVra2pg3bx7zurvp6ppH97xuli1fzooVK5nfM5/Ork5yuTy5XI50OoPnuShVOWb/DTyOFoo9aLE4APQDQ0qpOV9jwYpJk4m7dBcDJ4jID8Iw5OGHH+a///vj3H//fZTLZYaGhuIbxhmTC2K6d4eHhw/p5jRDTJLXhBCOyaRNWhRmMKCJdSRdmjAMcfBHxVmMpWPWkcvn6Z43j+7u+WSzWVpaW1m3dh2bn3MiCxcuIpvNkS8UKBQKpFJpXNdJisU/AjvQArEnfu0HhoCyFYtDY3tzmkRcuf1Y4BdhGLJv3z6+9a1LuOyyy3nqqSfjAF11BK4WD4dMJlN5kpsnePLpfDCiKWq1c4hrJRkLqXUn9F9wnGoui8ldCYJy3GVdnXDLBD5NjovrKua1z6Ozs4t587pJZ9N0dc5j46Zj2HTMsfT0zCeby2o3JF/A81K1vVDvQ8cn9qFFYy86ZjEAlJSaakTJ0kismByCuGZqN3p8zNeHhga5/fbbueyKy7jhhhvwPI+BgQHyLQVSqZQesRsI6ayelDv0o0qXbzLeYG7MQ7k5zhR7cw71e4VC4j4foRocVUoLShAInutQDgMkivBSaZQSlBI62ztIpzN0d8+jUGihpa2VVavWsPmEzaxYsZKWllbS6RT5QoFsNlexzGL+E3gGLRa74//3ElsV6HiFHfI8i7BiUodEsaFjiLt0n376aS782oVct3UrO3fuBLT/PzIyRC5XoFwuE4aC56XjXBGHIAiJEvVCkrGGiU5ApaYY8+CQYqSodiALoFDKwXUd8vk8bW29LF66lCiKWLRoMb/zohfR0zOflpYWuru7yWZz5PJ5Muk0KhH3AP4GbVEU0T0h29HuyDBzILA5F7ExkwQikkHXTH3IdOn+6Ec/YsuWLfziF78AqMQ4TI+FoTamkPw/uWxtMlltGjxUu2KjKMKJ77lkV6yJT5jPjJthgp/J/6PAH9VmI2DpdJqurnnMm9dNW3sbhUKBZcuXs379BpYsWcq8ed0UWlrIZXOk0ulR5Q2As9FBTNMTshPthgwAJaxVMWuwAdgGIrpyewvwfGCL7/vccccdXHTRRWzduhXPS1MsFlFKkU6nK4WWTczDdI3W64kAIBrdNZrMtajtSgXGrMd1goTrMTrzNbmeyvKxK1EotJDP51i8ZInuFk2nWdDby6ZjjmHVqtX0zF9Aa0sLmWw27gXxkufyLKoxin1UA5z9wIgViqMHG4CdIvH4mAKwGrhHROjv7+fSSy/lsssuY/v27eRyucr0l0DFUjC9EeZ/YyXUWimVvIwwQt+j+uYXRHsWSscqRKgksDmOg6Ociki4roMSVXGLKj0guRydXV309Mwnk8nQ3tZOS2sry5cvZ+3a9SxeuoSuzi7tfmQycS9IJV7xTeBR4Gl0rGI/VdEYQPeC2MCmZdLMGcskFhAzf8x2EaFUKvHTn/6Ua665hi1btlQm3R4eHq4UW86k0qMEBRg1HYSxEGp7QioWiDPaEqktZGSqw5suY8/zSKVSdHZ20tHZyfzuLgqFFlasXMHatetZsKCX9o52WgotZLM5srHoJVyQzwNPoKdh2MdosbC5FZZRWDdnEoiIB2TQvQZEUcRTTz3F+eefz3XXXcfAwAAigu/7lZs6nU5X3Y1Iuxme52EG6CVjE8lxKvXEIptNV6wYY110dXXR29tLd3cPruuSzxdYvXo1a9aupXdBLx2dnXR0dOhErFQapyoWX0AHMndTFYpn0RZFP1C0gU3LZLBiMgFEpBU9PuYhgKGhIX70ox/x1a9+laeffpowDPF9P86H0K6K67qUy+WKtZFKpYj84VEjc5MZnskkLGN9ZDIZ2ts76O7upqW1FeV6rFi+nA0bN7F0yRI6u7ro7Oiita2VTCZDJpNNZn+eQ1UstlPN2hwC9isVj1qzWBqEFZNxiIOpC4EXiMjlQRBw22238bWvfY3rr7++knlqxMGkiZtkMjNozIiK4zjkM1AqlSojXD3PI18oML+nh9bWdnp6eli6dBmr16xl4aJFdHV1097eTj6fJ5VOx1ZORSw+iRaInVQHjm1Hd5cOHg0DxyyzCxuATRDnhHQCa4FbzDD/K664gksuuYT9+/eTSqVwXbcSBzFWSSaTIQgCMpnMqIFs8+fPp62tjc7OTlpyDvN7F3LcscezYuUq2ts7aG1to6W1lXTcZZpI7/4PtEjsotr7sZN4pKntBbEczcxKyySRVLYUuEtEGB4e5totP+ZLX/oi27dvo1Qq6e5O8SsioeMTedra2mlta6N7Xje9CxexYeMmlixZSkdHJx0dnbS3d5AvFEin0riju0z/De127EbHKsz/g+j0bisWllnFnHZz4lohJdMbc+89d/PFL36em266kSAI6OrqoqenB8dRrF6zltZ8jo3HHMuKFSvp6pqnLYu2VgqFFlKpdDJe8T9ocdiDzrHYgQ5y7kZncdqxIJajjjkrJiLSDhzQ88fcz8UXXUQmk6FQaOFFLzqFxYuX0N7REVfFSpFKpZPt+zRaGPag3ZBn4vcmEcvOC2GZc8xJMRGRdhE5UCqVKJfLZLPZSqGbmE+ha1ckg5s70UVuSja/wmIZy5wUE6hYJivQk1SZ2hV9aLGw+RUWyySZs2JisVgaSyPFZGplvCwWiyXGionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDaIqYiIgtHGuxHGUcUTEREVdE1gCRiBREJB1/ZsXFYpnlHLHq9CLiAoHv+4gI6XS6dpGXAo+ip6/wgdDOdWOxTC+zdeLytwwMDPC2t72Nxx57jJNPPplVq1Zx7LHHsmbNGhYvXnx9S0tLcrpOROSDwN3AU+hZ+EqAAJGdqtNimVkcEctERDpEZP9ZZ53FD3/4Q8rlMkEQ4LounucRRXq+70WLFrF48WKWL1/OSSedxLp161i6dCn5fB7PG6N7L0BP8bkfCIDATsRlsUyOWTUJVxwPie677z5e/epXV6b0HB4epq2tDd/3UUqhlMJ1XRzHoVwuIyLk83kymQzLli1j4cKFLF++nA0bNrBp0yYWL15MJpPBdd3k5t4NPIy2ZPaiLZkQEGvJWCxjmW1i8vxisfjLd77zndx5552ICEEQkEqlKJfLFTEIwxARqbyPogilFJ7nUS6XAUilUiiliKKI9vZ2li9fzvz581m3bh2bN29m5cqV9PT0kM/nR7lLMacAT6InKi+jrRkrMpY5zawRExFJA6XLL7+cj370o4yMjOB5HmEYViYcb9TOmHa0t7ezYsUKNmzYwLJly1i/fj1Llixh4cKF5HK5ihWU4IVod2kf2pKJlFJRQxplscxwZpOYvPmZZ5759pve9Ca2bdsGaOvCiElSVA6XKIpwHAfHcRCRioUD4DgOmUwGx3FoaWlh/vz5rF69ms2bN7Np0yZ6enpYsGAB2Wy21l16J3AfWmQG0FaMb0XGcrQxK8RERBZFUfTM2WefzcUXX4zrupTL5YplYt7X3MSTxoiJsThq9ycMQ5RSFcGJoqiyfcdxKBQKrFy5kmXLlrFs2TKOOeYY1q1bR29vb0WIErwDeAh4Gm3J+CboKyLKukyW2caMFxMTdP3lL3/JGWecge/7RFFEqVQil8sRRVFFVKa6fdd1iaKIKIpGWSRGXEyMRUQqL9NuIxTJZYyF09HRwcKFC1m7di1r1qzh2GOPZcWKFXR3d5PL5ZIi8zF00PebQDfakgmtyFhmA7NBTF4zMjLyw3e9613ceOONpFIpfN8HoFAoUCwW68UuDgvTfhEZtT4jDkk3yHXdcQXGCJwRuXqCl0qlWLRoEWvWrOG4445jyZIl9PT0sGzZMhYsWEAmk0m24cNoC+bbwEKqXdhi3SXLTGFGi4mI5IDhz3/+85xzzjmVp73JK3Fdl+HhYQDS6XQlx+RwSbY/aZGY71zXrcRSjDCY5ZK9R0mrJQx1ukomkyEIglEulHGTTLvb29srsZmWlhaWLl3KqaeeykknnURPTw/z5s0jnU7XukunAfcDO4E+IMIm4lmawIwVE+PePPTQQ7zxjW+kWCxi0udre28cxyGdTle6fQ8XIxa1ro4RABPkNe5LUmyMaJiEOPNbsx7TfW3aW+s+AZUeqmS3dRAElXWtWLGCRYsWsXz5clauXMmaNWtYt24dPT09ld8keD3wGLCdOPBrrRjLdDKTxWSt7/sPn3nmmVxzzTUMDg5WrA/Ti+P7PplMpmIpTDUAGwRB5Uav192cjIMY8UhaJuZ9Ms/FiBBUxap2nUnXJ2mp1Oa3mCQ80JZYKpUiiiLmz5/P4sWL2bBhA6tXr64k4rW3t5NOp5Pr+Dt0fsxv0JbMIDoRz1oylikzI8XEDOT7xS9+wemnn04YhhU3wTI+xlpyXZdMJkNvby/r16/nmGOOYdGiRaxatYr58+fT2dlZz5J5A3pw5E5Gx2SsyFgmxEwVk9fv3Lnzu+9///u56667RvWOWA5O0lIyVo7J/k2n03ieR3t7Ox0dHWzcuJGTTz6ZjRs30tXVRUdHB6lUqvY4vx74NXpw5DCJ3iWLJcmMExMR6RKRveeffz6f/OQnK+6L53nWMjkEtXEc85nB5OYYwTABYHOMN27cyPLlyyuWzNKlS1m6dCmdnZ14npdc798AvwUeB3ZQjclYK2YOM6PEBFBAdM899/ChD32IRx55pNJr4/v+lGMiRzvJXiLzvjYJL9nNDVSOrznGyW5xk+3b29vL4sWLOeWUU1i0aBHLli1j/vz5tLa21rpL70MLzENod6mIHVIwZ5hpYnJyuVy+9b3vfS8/+clPKrkWQRDgeV5DG3s0kxSEZO5LKpWqfJbECIzneRVBMQKTDAabwLJSis7OTnp7e9mwYQPHHXdcpZepp6eH9vb2SvA55t3owO824Fl04NcGfY8yZpKYpIDyZZddxr/9279RLpdJpVKICL7vk06nK70ilvrUdmXX5skk81ySWbtGcExinVkHMKrXysRfzO+Mm2S263ke3d3dlVoya9euZcOGDaxZs4auri4KhUKtu/RB4EHgEWA3MIRNxJu1zBQxcYBwx44dnHbaaWzbtq3S9VsoFBAR+vv7yefzDWvs0Ui9PJmkeJjiUcmMXagKh+kaN1aFWVcyzyYpRLVJhmZZM1YpCALK5XLFktmwYQPr169n7dq1zJ8/n0WLFrFw4ULa2tpqLZm/QbtKj6IDv0V04NdaMjOYmSImC6Mo2v7pT3+ac845p1L0yPM8isUiURSRz+ennJR2tFMrHrWCkRyomMyNMcslM3hrxcisozZPJvmZ67qVJDtzDmt74oyYOY5DKpWiUCiwZMkSVqxYwcqVK1m9ejVr166lt7eX1tbWeu7S42iRqRSssiIzM5gJYuIA4d1338Vpb/lT/HJY81Xyr7V+j2aMKCmlaG9vZ9myZRx//PFs3LiR3t5eVqxYwbx582hpaakNxp+OLvHwOLAH7S5ZkTnCzAQxOXloaOjW97zn3dx8y80gyRwHKyZzjdo4T9I6EhEWLlxYcZGOOeYYNm3axPLly2lvb6+ITMKSeT86JvMg2pIpY2My00azxSQDFK+44nL++Z//qc7XVkzmEsn4TD2XyoxxMssm3bG2tjaOP/54Vq9ezcqVK1myZAlLliwZ5S4lOAMdk9lGtSqetWSmSDPFRAHRo48+wgc/+EF+9asH6pQRqLVSrJgczSQFJDnWybxMzCcIAqIoqgR6TdC3VCpV1mECzoVCoZKAd8IJJ7B27VpWr15NT09Pvap4f4lOxnsYHfgtWytm4jRTTNYHQfDge97zF9xwww2VoNxobPr8XKQ2g9e8koFjIyCmyxuqAyVNr5QZCBmG4agMatd1aWtrY/HixZx44okce+yx9Pb20tvbS3d3d60l83HgV+gR2M+iYzIjVmTG0iwxcYBwy5Yt/NM/fZgDBw7UySOxQjKXqXV3kiQT64Ax/5sHk+u6Y4pT1ebJJMlkMvT09NDR0cGKFSvYuHEjxxxzDCtXrqSzs7PWkvkYulfpYXTgdz9z3JJplpi8bGBg4KdnnHEGt9/+S9LpdOVpUsWKyVzDZNwmE+HqFaiqt7yp/5L83nSHJ2v7HgyTaW3cpyiKKukInZ2dlZq+69atq8Rkuru7K9OhxPwX8ADakjExmTkhMs0QkzYR6Tv//PP5r//6r0o+QhCU4xNiRWSuUpu/UhuINWJiclWSVkuyuLgRkGTy3XiDH817s24z2tpYNkkXyrhUYRiSTqfJZrOkUimWLVvGmjVr2LRpE+vWrWPlypV0dXWRzWaTIvMv6FkKdqJFZg96HuyjpsxDM8TkHx5++OFP/tmf/Rm7du2q1HSt6dKrQ4QVGstMJjmOrL29nXXr1rFixQo2b97MqlWrWLlyJfl8vl7g90/QlswO9MRuZs6lWSUyR1pMThoZGbnjHe94G3fddU/FDE2OWh1L0jq0YmKZuSTHLxkLykxZayoCtre3s379ejZs2FDJ+F25ciUdHR3j1ff9LXqQ5AFmuCVzJMXEA/xrr72W973vr1FKm5CmOHQQBKTT6To/s2JimR3Uxnlqy0GEYUgqlRoVLE5aMqtXr+aEE07g2GOPZfHixZWqeDXuEuiqeKZ3qQ8dk2m6wBxJMTlt27Zt33rrW9/C9u3bGRkZwXVTZLPZSiPGFj+qjVlZMbHMXIybnhzXZIK55vsgCCr5MKY+r4nRJDshUil9b2QyGVasWMHatWt5znOew4oVK1i6dGmlvm9CZD6I7lV6Ch2TOcARLlh1pMRkQRRFO84++z+56KJvAsQTjmcAGB4eJpXKxIqdFBArJpbZRVJQDMlcl2SsxCTi1VYT9H2/kidTrzxEOp2mpaWF9evXV8puLliwgPnz59PW1lYrMh9CDyd4Ai0yw0yTyBwJMXGA8NZbb+Uv//LP6e/vRykVJxQlF0megNq/o5ezWGYiyeEAMLpb21grtd3ZJhnPdEQkq9yZOGKyNESypoyp9SMilelnu7u7Wbt2LZs2bWLt2rWsWrWKxYsX09raWjtTwd+gBcZMhzLIFIcUHAkxeWOxWLziTW96E/c/cDdI9aCUy0E86VSOUqlUmXOmrpioCMSrXbfFMmMwAdjaZLjk3NTJIK2pL2PKNhiXKFnRLjk0wCxXW/gKqi6UWYeJ15jpYBYvXsyaNWs48cQTWbduHQsXLqxUxaspvfkBtCXzJLp3aVApNaHiy9MtJh7gn3322VxyySUEYZEoJBFwzcbdwqmapDUrJpbZR/Lmrze/krFCkol0wKjPk+VJa9dhPoOqFW/cJh02SFU6NTzPq/QgmXsrGRROpVK0tLRUpkA57rjjWLx4McuXL2fhwoW0tLQkRebf0bGYh9ECs4M6QwqmU0z0QL5HHuaNf/wqhoaGQNobtrHDQur1Fk0CNbXiTKqBB/swW9Dk7c9uZIphhilHKeTgBdWTmcKjfmZ6mUiP+SyJSfhLpVIsXLSQDRvXs/mEYznm+E309i6gZ0E32Vw2Lr0JwJfQbtKj6J6lp9C9Sz4wpb2tFZNVQRA89pfvfDs33/J/Wl1Vk6vLH+JkHJJDTBdzqHTt6Y6rH+rJcKj2HZqDx6wa+WSaHqY4F/UUD59zCDGf6vE7WJavXiCou6whWQXPZAObhNKWlhZGSiVWrlzF8hUr2LB+E2vXb2D58pX0LlhIvtBCNptFqco18j70zJGmiPjApPYlsQMOEN515+38+Z+fTj6TZWhoiHLYj0g0yq80ZpvxJWt3rpapHfCpDo+YWgBYyZEbnlFfOKrtP5wJzeqP7E6sPfFd4qI65LLNxnEmphKhjH6Y1F6fIge/ficrJko5iESVY+m4VfdnPEw8ph6lcv+YWMzo7alRMZmki6WUwg+DSj0ZPaQgQz5XoKOzk57uBaxYuZr58xeydu0mNmzYzIL5i8jlWvC8FEqpNiYhKEkxaRGRgcD3EYxQhHWFYNTuTPnJaTC1TRu0ukOQbPVEN1n3hE+1wQ07frOIOodMpmZhj0vtOdM33jgL1zMMmiyg9duqRn1fvQQFUHWv0+p+j3arlFKEUQgCrufhVH5bWU8eGJlQW2vEogVYjH4c/ho4FnAZe5hlnP8nw3intPbzid5tk21HcvmJbqOemTJVNTiU6XMk1OZwt3GoYz6Z9U70epgstcdXMblrZapTUh7O/THe/SV1PgupHiM3fp/cZ7O/5j42y3vx3xDIxf/3oYXDiT8fYYJCAgdPWpvQQZ+6T2+xWJpFI2NmB7PhZnpkzmKxzCBmTkTNYrHMaqyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA1hyrOKz/zpJS0WSz0aPU2NtUwsFktDsGJisVgaghUTi8XSEKYcM7HUJQOsBrqBAnoO12eBp9BzuFomxmTnBbboYzYfWIK+/hxgN/AQMMg0Hs/JiIkDrALS8fsw/v2hJt52x/lcxS8/8ZnEnzlULySFnjw6jNflxC/znfk8ojppcyr+zR5gP0fugvSA/we8Enhn3M4UUAT6gW8BVwM/P4JtOhxcYCX64tvJ9LY1AywHsuhzF1E95+Y6aBSTmRw9ituwG/0wONR1PhPYCLwUOAM4geok5T7wK+B/gJ8AB6Zj4webuDxJHngXcCb6SZsGSkA5/m7Memv+1n6eJOlqCVWRUDWfm8+CeLugL3o3/sx8P4y+gV3gMeBs4DuH2L9G0Is+Pu9lfPcxRLf1I8BnmJmC0os+138Xv/8S8B9Mj0W1Gng/8Jfo8zWCPnfmgRLEy413PY3HZI9r7fUG1QfYNvQNeCvwXWamqCjgT4Avou/PloMsezPwPuDuhrdCRCbyerOIDMrsZK+IrJKJ7efhvtpE5EuTaFMgIu8QETXN7ZrsS4nIP9a0tSQifzsN23JE5JsiEk3iuDWTsoh8Q0R6pPnnqfb1ChEZmMS+PCoiXdLg62+iAdg/RyvebKQAvGwa1+8Ab0NbJBPFBb4OnDItLTp8cow9Ving92l8sL4DOI7GujHTSQp4B3A92nprNua4bQQu4+DWSC2rgW/T4JjpRC+Qn1J1LWYbPnDbNK1bAa9CuyzjMXSQ334OmNfgNk0FQcd2kijgURp/06eZmS7DoTgW+ATNF0FBi//fooW5HgdzTf8AeDcNfEhMdEXfAH6EjpME6Bu0HL83sZPalx+/TJwg+TLBrWEm5t+GiXWFiW2E6ODmMDpYWELfvAPx/7vQfuQDE9zPyaDQCv/v6CBiPT6P9k8vo/6NcyJwHjOnV03QgdAkEdA6DdsqMf5NMET1fPs1r+R1FDbolVxfkNjWeA/QtwCvPox9bhQmxnMm2muox2PouNf70Z0Q9fgUDdyPiV7EzwJvRfdSLELfuAH6IjM3idT8hdG9LgZV8/p7dDS/nrCNANcC9wBPxsvkqJ5oN94HP/7roC/SFrTI3AHcR+ODhwrtPp0PnFTnewE+iBayCC0mQ+genuSxcNAn8yx0kLPZAVnTCzZE1a01x7TRT2KF7iWp5RHgv+Pv2mq+M8fnUG2ZzHE0gd5kkN9BX08Z4M1odzS5TQ9tkf6Y5lhXgr4f/576980DwCuA7fH7u4GvAetqliugraxHgAen3qrmBPnM3x4RuadOgCgQkbtF5EQRSdX8RtWsq/ZVb7lGt98Rkf8QHZysZUhEPhsvk2xPVkQuqbO8iMhOEfnjaWrvZF4FEfmBiIQ17fuqiHgN3tY8GXvufdFBzuRxqz32jX7Vbqf2+wUi8hMZyw0i0joN52Air2NEZHudNomI/Db+vvZeOFiQ9koRycjY4z2pVzMyYJNPjVcAm+ss8yDwe8BdaKsjafUkfy91XvWWaySC7oY7k2rOTZJbgX9CP7GS7Smhu4TruVzz0TkAixvd2MOg1pKc7m0lcai6WfUS1uqd76m+attR+/1O4Lo6be+q074jQQfacltY57sB9LX3a0afR0Hvw+eo77r9MfDPU21YM9PpU8Bf1fl8CB0YOnBEWzMxFDoA9wXqJ+M9CpzG6ES8JE8C/4DOXailCx3IzU29mVOinpgcqZsmmaTWbJfPoNDnurY9JcZPyJyudij09fP/6nxfBj7N6Jyq2ofqv6Nd7lq330GL0OuZwoOkmWKyBlhf5/P/Q/e+NDtaXo9FwCeBnjrf7QH+Gp0xWQ9zUq8FPsbYXp4W4A3ooFm9JKq5gou2+JzEy028nAa/3HG248XtOA3dY1KbRLkHLShHArPt16IDqrWxzhI6fvNpxrfKBS0i7wXur/N9Bm25bOYwr71m9iJson7uyneYeV2G5uZ+G/ByxopwgO6V+Wn8frynqnnqX4DuyTkdbaEZPLS5+SRwSQPaPRv5A/S+e+jsatNVPZFs6lqS11FtJwDo85EMMJuUfh8tJK3A7zC2h0uhM0mPlJgIOj3+Yur3rO1CP8hGOLRFN4juAdrC2LSEpWgX6k8ZmyJwSJopJkn/2BCi4yQzkT9Bxzxqj1mIHnMz0d4YQYvPvwDHAyfXfF+It3Mb8PAU2jtbaUUf65nMr9E9eUeKFnSSYz0OAH8B7GDiruFdaHfpPMbG/V4K/CP6+pzUQ72Zbk69HTc5KzONFwAXUrWkkj7ng2jTcbw4yXjsRJ/Q7XW+Ww98k7Fi2wymM5g9Gymh8652HaHtOegUg03Ut0o+hXadJ3OOIvT1dQFj77cUWkzeczgNbRYdjD0AeaqDu2YCJjHtE4x2yUzgbRf6KTpeluuhuBH9BBip893z0YMUbc2ZmYVCuwpHojyCQndSvJmxFoSgr0sTJ5ksEfBh9Aj22gehi+6tfBmTiJ80+0KtbagpJzBTSAEfQN/YtW0dQSebPcThX1SCfkJ8EZ3Fm8RBJ7m9+zDXfbgcyZ6c8S7UEqNNbGMdRYlXoyymeus2mbD1zPw02g3d0IBtH4oXAP/KWAtV0CkIn2DyFnGSYXRy5WN1vlscr3/Cwz2aGTOpdyHNpB4MhR4afzr10+W/g84qnOoFHQL/BZzK2PhJJ/oJcRc6m3cusAO4lLF1a6aLMqNr9BgxMQHY16LzgJIsRJdp+AemT2x7gI/X2TZo1/jv0fGSqWxf0PGf/wa+wtjrfDO6h+ftTCB+MlPGhMw0FPAStBlYbzTm7WiLoVG9Tv1oc/YqdFGiJMuBL6N7OQ40aHszhXpW0DVon73ZFelMD945wBXA2prvVqNvvuI0bNtFX3un1vmujLaMbm3QtgS4CN0Z8F5GC4qHTle4F+1OmeXr0kw3p97TZiYE+xT6yfNFdOm7Wrahn1ZTMS9rEfT4iTPR+Qu1PAf99EjV+W46qM0IPVIkB+81G+P63IN+ate2aQ31M6CnioO2BP62zndl4Kvom9+0sREIWsC/X+e7LNrV+r1DraTZMZOZSAqdibqozne70N1w01XK8Ap0dm1tQNZDD+x6P81xBY+UyLvoJ+NMcnehvqvl1vlsKhhL6Lnoa6Deg8N06SaHajSKAB0/+XWd71oY/56oMBMtk2ai0CUDXouOVyQZRvfLb2H62hmhM2xvZazlU0CnPL9gmrY9HkdyrI6D3u+Z8JAzo4ePRV8TtW16ksamMQjaIv4a9V3rHegHynS4VcltvJ/6o7lPQMf2xu0gaXbM5EheqAfDtOFU4D+pX9f2f9Fp8NNNiWqG4pqa73rQT61XMX7a/lQZLxV7OqgXc3o5OgvYFEM29X0P9hs4eBsPdY0lb5Aw8VkLOpbVVec3T9DYDNg0+vo6puZzYxW+C/gt0/vAFeAGtGh8os73b0G7fZ9LLF/BiolG0F19X6C+kNyGDrgeqYS6x9A+7HmMHQd0LNp6eTfTE1uoN5J3OoY3mJhELcvRuTfpxHJJasfINBqz/ojxLaQD6DhKo46LQqcBvKPOd2V0zGI6LeIkgh7BvhCdFpHEi9vyCDpQPopmmpPZOtv3aE4GrAL+kPoDD/egR1s+eyQbBHwP/QSoffpl0WMnnjNN280xVuDTdT6bCsaFqJesZ7aXXDb5ql1Po1+Gg90bl6B7OBpFOzpBsZ4FtgV9cx/J8Wo+2kq6r853HcAb6/2omWIyRH3/rxlJa4K+iWqj8xF6qoefcOQtKEFbIJcxVmAjdDJRo9sUoHMYkgJWZPwErsNF0EJSb/T1TCZCd99/gMZZCQpdUa72HDvoGiR/RXOywvfE295X57vaim1Ac8XkbsY29Kdov7AZ3MXYcTJfpxonaUZwOEB3Fz9K9YIqo4/RTdPQphI6o3cX1Ys7QD8dG/1kNAHt8ayTmcYedCD2T2m8sG5DlxBIivgBdHrCZAbwNZpb0NefOUcRuo1XUydEMdFJuKaL16GTc1ai4xLvoHmJWQod2PwY2t36OTrLcDqj5xNlGdpK2YweSfxBpk90PXRE/y3oIO/XqRbcafTFkkcf41ehe6uSs/ll0A+bZOW1qVKv8FK9GEyy0tpT6N61b6FTAmp/3ygK6FjRm+PtnAX8bJq2NRkcdCb438VtOQdtrY+J1zVbTED3p2fQbo9Ru2Y2yswGWGpyO2px0W0rU73pprN9xuWb7hiWqSNiAr3mb3KmxmZiKuIfCcyxMNucruD3ZFHoe1TQ10Pd+3QmiInF0uwHiKUBzITkIIvFCslRQLPzTCwNQKnmeALWqrUksZaJxWJpCFZMLBZLQ5ismGSYfYlGeaZnrlyLxZJgMmLioosq/+40teVwUejYT73AgYuuGl9bwYyD/KaR7ZpJJSgtlmllomJiuu4WUk2lbXb/v6EX+CzVYdu17Xoe9S2T1wJnTGO7lqDH1tQbOGixHHVMVExM2H4XcH3NZ80mB7yGas9UbbtuRGcx1rIZWDF9zSIDvJj69WMtlqOOyXQNR+gZ7Uz2JcwMQUkBj1PfpYjQI37rDdU3maTThYOeSHomZDBOmUN0Azvoka+ghx8U0eelHZ3ZXFt533IUMtkArBk9OhNqtRqG0UJSm/ZtXDOf+jf0bhpbxzW5XdCDo4qJbc8Ut3C6eB26TMO16Cp1G4AH0AMH147/M8vRwmQsExNQNHOLGFJokTH5+gpd4OYU4Gm0YN2PHrRVO8CqVpAK8W8Xx9saQF+gT1KtvGW20xq3YxN6OoDl6OpXMHpmwBTVKQwy6DEnETqWMhSvR6haEhIfl7BO++phpmQw1k8ubvuKeD8WJto9zGgrSSX+LkPfdMPoSZ6eRJfPmymifTBM+yP0PpuJxwvomJGXWO5I7o9N0z+CTEZMHHQ5t7uBb8eftaPdiO+gh8QfR3VCqRvQ85YuQpeiM/Px1jN5s+hSfaehh9s/jK429rvx7wN0sPTxePkl6NKGHegbcA26FP/d6Jv4G+hh8w56tOMv0WUXj0eXZdwd/78QLTYL0CUIPosWotPQN8alEzguvegiv6Y62FlAd9yuXnT1tp/Fx+bb6NooUL3QN6FHYXroId/70dMovBj4UdzeejU5x0VEnHj7B4BdSqnpvqFS6GJO29FFg/aj9+fP0Bba06Zp09wO4u1uQM9Gl0Zfm08coW3PbURkoi9XRG4QkffE75WIdIjIAyLyXBF5jYjcKCJvEJFOEUmLSFZEuuPXjSJyl4ikEutUIuKIyNUicrOIbBaRrvg78/v5IvLnIrJdRNbVbLtDRF4uIo+LyMki0i4ivfE2zLpviNuGiBREpEdEWkTkoyKyJd5eV9xGFb++LCL/NcHjslxEdsS/N9toF5E/EJHfisgLRaQ1/ixds+9/KiLb4n2YH7fbidexSkT+VUQeEpGlB2tDzfl0RETCsCwiocR0iUhOtMg0+jpRE/zsSL1SIvL2eL99EXlfE9sy2dfzRd9f5lpsdnsm9ZqoZZI000uJ9yG6t+QitDvyWsYWPCqizf6Xop9QL2Z0j9CbgFXoXBAzZ6+i6qYU0XOF9KELyKyPt3sg/v5BtHvyULxMX2IdUHVdVNz+4fh9P/CDRHvNPpnpFkztikORibdttmfW/xt00Zv7qT837fFoS+g16CK9yTYMo+uVnI12d65GH5+JlAPwAIrlvaCG8aSddLZrr2meiOSVUo0sSGT2aT5wItq18dDWyiDa6gzQ18mOBm73YO0x15FLtbDPbHB5/gQ9X86n0PfUv1F/HqUZyUTFxNxkPYyeAiKFvrk3outC7qP+STMxlveiJ5N6fuK7AvqGS07+bX6fXNetaLdqCfoGM98V0OZssjcnGVvpi9uZXCfovJSemuVNWyfTCzMSt8vEQsx6WtGlFd3E52Y7GbQr9j7q1xJNtuUS4IXook3/MIH2+AD57AL6B39Ded9/E+Seh5d7AenCBlCpYRFZBOxWSjWiTkcGeBF6KstNVOeT2Y12UYvoB9AbaZyYHEwYTC0aUxvl8XGWm2mk0OcuHb/ejQ5qvwHtpte7J2pJxuBqr+lpZ7KWyTaqQUFTEcsF7kRP3jPejppK5Hegpw3IUC1RN0I1SJlcPikIoIXqPVQL5kji90bsqPO7Lqpxmtr2DSa2l/x+MgWBnHgbtcdS0DdPquYz0MWgDwBba9pT79hFaL//TPRT/lCV35T5k8ssQ/yvEgWXURr8XcKR03DzzyNdWLMdHERkGfDsFETFRVuVpuSgQvfmPIiOZ/0O+mGTi1+mfbX7mQGWoo9lP+OLTh4tUFl0ztOuOssI1cJCMHbS7xQ6SNyBjvEcQN/Ek7Va2tHX7V709TLIxGcLSF6jZrs+Oqb2DFp4T0HH3G4AXgH83yHW5wInAS9BP3AzcRvvRedZXcPU5yY+KJMJwAraKklmk0bop7spL3eo7s8ALUZpqmLyC/S0AX8NnEtVRZMl9BT6ZH2XsQejAx3wq70hTA/NENULK3kChxh//9PUnwhpPLYx+qlg2tiWWCZ5Ab0N7WIlrbHxEHRJS0EHdrcdbGGlVCQiLiCpVEs0MtRLLu/hyAOUD/wnbnQy/f3HI6kTae/+3afARUQ6gX6l1GSfYDm0KW6ehP+NDoSPoG/aD6CnRhCq7ka9Hr0XoV05B7gY/dCo15aXoS21FrQr8M+MvR7M8AqzLeOu9qBvytPRPXlZ9DVYRFt9F1G/ePJ4/D56wqwSWjyH0J0MVzDx+XSy6DmkF8ftvho9ReeF6Hlr/hp93K5BW//m3CevsTS6zOa70dnpxXi9plrdn1ItCv4NdPnPZ5gGa2UyGbBmB5IXhYt+ktye+OxQ6zBdiMTvt6HjKP+Odl9OQz+lMjW/Ta4jKVpD6CdErQiZ9uYZ/cQw33tUn5bJ7xz0gZ9oopUb71OtNVNGn2iTy5JslzH/TXwmjb4ATAnLbPxdPv6ugL5Aj51Ig5RSUdyD88nc6isohy8l7WQgeIjBfReRis7E738XA7svxS/uAGQ/EIoOdE8mUJtFn7thtG//RarXR0BVlE1Kgdl/gzkm7VSnBa2dhCpJctDmUupfbxHVY96HPm4r0VXlv4Z+4reiz3Me/TD6OPpGbuPQJBM281QznV+E7pW7Ad0bl1y2Hm9Fu+7fQddV/Rw6TnYO+nh9IP6sk+oEXeaaTt4Hr0EL+AaqD88fo++nT8br34k+D29Gx/fedoj2eWjr5qVoS6ebCWjFZPNMyoy2AMyNN1FFN9Wtay+Ce9FzrP4+Wkw+gw7o7kOX+9+CdqNM7kfyICh0jKN2btakMNRLThth7LgZFbcxZKx5XIs5sUZ4arOCjQ9ce8LM7/4ZfQG68baM9ZVDX+C/Roukj77IFzD5TNJPUXjxh+nNM7DzUpxO6Auhz8uSDw6Q2fcOUgegmIG92dNZOP+c8x3y5wOIiHsIS8WN2zoSt+076AeLwUGf6yDer/Em1SJebijez0cZ/6kZxetqZfyq9g76ug7Q5+BD6BjdRvR5/SbaZdgfb+8jaLF5IXr61X/l4DVfTfvvQc+1lEZfu29Fn6+T0Dfzixg766I597+Htsbz6Hvq6bjNC9FxkqvQIvEVdJH1tvjzjrjdScukI7H+Mjrw/1ZGpxO8Hd2JYe73R2r2Jdm+56CF6AXo82sSVa9Fp1mMG/earJtjXsnPJlr0VsXLHmBs6rug/boL45dJQjsWfXK2oBX13YyeItEkjNW2y2wvRB/UjprPBX2CapOpjFClOLQLkhSNEmPFzIhvKvHerL8HPWPflsT3JfTJMxaOOZEmhuNS7cGaUM+EUmqPiKxMt21+PD30Wdj+d6xsa4FiDrwCRAUYOkB2OMPi1gFCriIovBCvsApwwkMIiomD5dCCV3tzC/q4++ibPyms9QL0xkpMJifWMkzVMhlPTCReh0ma+zP0ef41+ub/LaPPxU/RlvVK9EwJlwK/GmfdJH77ONUZAn6IFtPz0KK0Gi1Kf8voe8NYM38c728J/eD8d3Re0MfR06O+Am2h3Y92e94Z/34NVS/AtOVOtPv/YvRD6TjgZrT79Ey8X29FH68M2i29uWZfzP8b0Q9v0yNn7oERdE/Ty9FiU9fVnkzegRGDer85VKwEqievg7GD8motjV3og/Z1dI/HC9BTT9yFDkqZ5SKq2aq1wS/jhmVq2pwM3Na2zfzfzcQrknvxNkxg1KzHQZuWtf5ziLa4IrQ47Ea7CP1oc3R//NqOPg674++SGcSTCaI9jfIgeC7cthC+1wc3DjJ43TNw4xDcloefRvCDn+Fe9lcUL/kd9t15PmF5GPQTdzwEbSXsQz+dX4/2/Q154JVU3Y6DCWBI9fh1U9+ag6rrGzH++XEY3XsH+oY7mdFCYvahH52MORB/dso46zXUOweCnv3uzVTHib2O6nil5P50om9MQd/sZ6OPz2/Q8SJBXzMnxf9vpfrA3lSzLrPdP0Ef/++jhWAdugfoTHQaxsvRx+5/0LGm2t4e4jZ/kursDB9HC+Mp6PuwjD43H2Qc3ZhsElO9k1wbwzjYb81Tw7gdncCrqVoXZhmn5ndPo7NnL0MHm5KUqB//AH0CTL5DLR76wBuSgrKT6oVwKPJoM7/2qWuesNmaz0ErfvJGrSeqyfcO+jjVWj8TQQE8eNU32PaTbQz/EvZfO0jppiJ9W/fT/6NdDP98P/03HuDJO4UdpReR7n4+bjoP9aeHTLZ5GB1ANwP9Po32tVein4Cb0cfYQd/M4wnKYqpP8HsZa+ma3yxCn2+H8SduN93CZh2/Rfv9tT13hhAdL8nH353C4ddG3oO2UiJ0D5+JwSSvrX70k91Yc8nrI4i37aKvkRVotw/0uU9T3zsQdAb6O9CxljL6mP5nvI5d6IfyWTW/SZJGWzVtaPH9JDqw/Jv4//vjNv8/RqeHVJism1PP5A2YWJdY0oUwYpFBB+3+ED0ozJCcFyYpMleiI/210ewRxndLTFSbmt956G692jYSr6/uAavDqVS7p5MYcavtyhW0SfwvVH375L7Wa8+p6MmqfjbBNukfi3iAP7LzMVrLu8njkC+3kPcVYXmIMAxId8KB9HzKv/c+lrzur3E75oF2wc5VSh2qGzpCP+lejhbUP0ALUAF94RWpPmz+Ce2q/l/8uXmALELfBEa8W9HWq3kam/O/CO37m/lbcugnpU81tyVE38QnxJ+FwOXxejqoXsNpquO18mjTXcW/WUz1vNSi4naawlo+1bFfxoUxLlaEDmCaCb6N9dyKflhtRrsV/4wOtK5FWylptEgHaCFYR/XaWIbOXTLHxNyPXvxahg5MG8vNi9v4LNql7oiPn3GbzfGI0CLRFn93J9V5rEAL9B1oD2E5OgxRO9Zu0gpcTzj8Op8dDJNMBPqgXoU2nd5V2zjGuj9tVN0GcwP66MBV7TzB5iIOqXbzJtefYWzRJLPO+9GBuX9FH9zaJ6p5vwhtSu5hbN2SAfRNVUA/jZLr+DH6wnkLukvyUG7LK9GR/8nOLtgKwrPXnEf6lu/TKcBQP6Q6Yf4m9ndmKC/fwII3foSORetAm8o/mMRYnhDtjv0h2hTejL4R+tDn9RfomEBvvOy30E/lEH0eu9HXw6J4fRF6DNbL0WLRinb5TK9QR7ycQl8z70GfH4eqeARUxQh0ouRrE+vbE7dne7zOfegnslk+z/hB8854v1ZQvReMCJmg/bz4Ny3oB+W/xL/ri7ffj77pzfb+AZ1X0klVUFNUp05ti9fvxvt8RrwvIdWH5Eh8PBdTvQ6HEssch7b4clRjliPxcRqIt9sfvzfBcrNPRsiWxOstM96sm5PIvXdF5FoR+YvEZ52ix9S8egK/V6LHsTwpemyMGQdzoojsEz1+xjnIb5eKHqfy3MRniB6f88u4DfXGM1wtIq+v8/k7ReTBeL+oWWeriDwqIm8eZ52IHgN0nYh8QkR+LiILa77vEj3u5hXj/P5lIrJbRH7/EMfsH0Xk+viY1V1unPOqRET8+26VZ995ksifZURei0SvQA584FTZduUXxO/bIzF/ITo3ZVwmcG57RI+PepXocUauiGRE5PR4PysDhepQFpFIRALR42nGw3wXHGQZQ1TzfiK/ERG5RUTyUv+8LxSRoQmsoziBZXypHpPksUkeg6DOZxPZ9k4Rea+I/L2IDCfWc7BzIFI9ZoOix9iZe/SMxDI/EX1eD3tsDlQVOM1oszzDxHx5c9WbPmtjqt2NHgF8PnAB1VGeYbzuAjqK/Vl0t95dNevrQ/vn56EDaTei81WG4+0ku1ST1sG1wN/E6zwX3eX1ePzdIPAX6IDYAvSI2L64TW1os/gstPXyBNVRxkn2o586X0P7rr9Ej4Y27tjP0ElJX0O7ClvjdnvoJ9wadD7AGrR5Xy/uUxcRqZjA3jfPpmvbI5BKMdTSSekNf07LK06nfcEqs3haKTXVui6CjmHUxjFC9LF7FG29LKWaMm7KNmTRZniGqsuSzE3ZTzXPpoC2Zj30U95kUrdQtdrM9JpDVC1T0NdoG/o8L40/T3bpl9FWyS2MnxFbBj6PNvXTcVuy6Cd+hA6oLoh/b/bPuD3tcVtNMNq4R6YL27RVUXVRJN6PLqqJZ2ZdEn9mev364tf96GvJjOm5g6oLsyhuc4h2VczYqWG0pbIcPVK/DZ3g9tJ4P98Ur2t//HndDPHJTA/qoLPpdqCH8yv0wf8TdE/LExNYRzvad/53Rg8YVGhX5e/RJu59aFPqOKo30UfQwgNj3Q6AP0J3jxXQ2aVb0Bfk69En+ZbEb4l/vw7tahyD7i77EqN7HjbE212PNov7qBb9+Ti6v74H3fX2FUbngZjtvAYd5e9H31jX1bR7Sbzfp6CP4R50ALMFHSO6kEPkl9SbhEtE3rL/2W2Xdn5gKfvownnz2bS85E/wOiuTC2QAf6IuzSSuk7pNTPyt7Y0wn+eodoObWEmKar0c426a7M6JxuoO1qZkYDQZm6sXvzI4Nb9pNrW9U8m/yWXq9eAkOzxM+OEf0K57iL6/HfQ9sR+dlfuFOuvXK5vkRXKw7r2J/h5Gn7haMmjrxfRzm4Smiab/JgNLE8WY+OYiTuKglbqDaldiX836D5VrkzyZ490ABfQTKIv2Y3cfZNnRK68RExHJAiPbH3yA/d96Nxs/cAVO1zLQwnYdEEy2xskUxeRQHIuOB+TRovzz+PN3oQe6PYxOmOoBPoq+Pj5G1Uq1TJ3kRfR8dNDf9HANo8/J7Rzk/j/cicunKiqWBlJHTF45sHfPj/MtLbiZSiJvbgK9M5OiQQKjgFehzeccumvzwvi7C9EmdoS24JajE7Ra0IHILdjrcDo4mNU2/o+m+YljaR6vRLt3yWEIM5XFaBfVRQ9qNLkVz0NbLTvQMa4ser8yVGNjluljUu6cFRPLTKbeE9IyQzncTD+L5UgwXkDRMgM57JqgFovFksSKicViaQj/H0ecN7mC/VekAAAAAElFTkSuQmCC" class="header-logo"/></div><div class="header-text"><div class="header-title">منصة المهندس أحمد سيداحمد</div><div class="header-subtitle">معهد طيبة العالي للتدريب</div></div><div class="logo-box"><img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAARMAAAEqCAYAAAA23LTdAAAACXBIWXMAAA7EAAAOxAGVKw4bAAAF8WlUWHRYTUw6Y29tLmFkb2JlLnhtcAAAAAAAPD94cGFja2V0IGJlZ2luPSLvu78iIGlkPSJXNU0wTXBDZWhpSHpyZVN6TlRjemtjOWQiPz4gPHg6eG1wbWV0YSB4bWxuczp4PSJhZG9iZTpuczptZXRhLyIgeDp4bXB0az0iQWRvYmUgWE1QIENvcmUgNi4wLWMwMDIgNzkuMTY0NDYwLCAyMDIwLzA1LzEyLTE2OjA0OjE3ICAgICAgICAiPiA8cmRmOlJERiB4bWxuczpyZGY9Imh0dHA6Ly93d3cudzMub3JnLzE5OTkvMDIvMjItcmRmLXN5bnRheC1ucyMiPiA8cmRmOkRlc2NyaXB0aW9uIHJkZjphYm91dD0iIiB4bWxuczp4bXA9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC8iIHhtbG5zOmRjPSJodHRwOi8vcHVybC5vcmcvZGMvZWxlbWVudHMvMS4xLyIgeG1sbnM6cGhvdG9zaG9wPSJodHRwOi8vbnMuYWRvYmUuY29tL3Bob3Rvc2hvcC8xLjAvIiB4bWxuczp4bXBNTT0iaHR0cDovL25zLmFkb2JlLmNvbS94YXAvMS4wL21tLyIgeG1sbnM6c3RFdnQ9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC9zVHlwZS9SZXNvdXJjZUV2ZW50IyIgeG1wOkNyZWF0b3JUb29sPSJBZG9iZSBQaG90b3Nob3AgMjEuMiAoV2luZG93cykiIHhtcDpDcmVhdGVEYXRlPSIyMDIxLTAyLTA4VDExOjI2OjE2KzAyOjAwIiB4bXA6TW9kaWZ5RGF0ZT0iMjAyMS0wMi0wOVQwMDozOTo0NyswMjowMCIgeG1wOk1ldGFkYXRhRGF0ZT0iMjAyMS0wMi0wOVQwMDozOTo0NyswMjowMCIgZGM6Zm9ybWF0PSJpbWFnZS9wbmciIHBob3Rvc2hvcDpDb2xvck1vZGU9IjMiIHBob3Rvc2hvcDpJQ0NQcm9maWxlPSJzUkdCIElFQzYxOTY2LTIuMSIgeG1wTU06SW5zdGFuY2VJRD0ieG1wLmlpZDozMWVjODIwMS1jODdlLTJmNGQtYTliYi0wNWJhMzAxZjlhZmYiIHhtcE1NOkRvY3VtZW50SUQ9ImFkb2JlOmRvY2lkOnBob3Rvc2hvcDpmYTE5YzI5Zi0wMDQyLTMwNGEtOGQ4ZC01MjkzZmI1ZmQzOWIiIHhtcE1NOk9yaWdpbmFsRG9jdW1lbnRJRD0ieG1wLmRpZDowNjBiNDU0Yi0yMDQ3LTIyNDgtYWMxZC1mMjE5MjcwYmVkNDkiPiA8eG1wTU06SGlzdG9yeT4gPHJkZjpTZXE+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJjcmVhdGVkIiBzdEV2dDppbnN0YW5jZUlEPSJ4bXAuaWlkOjA2MGI0NTRiLTIwNDctMjI0OC1hYzFkLWYyMTkyNzBiZWQ0OSIgc3RFdnQ6d2hlbj0iMjAyMS0wMi0wOFQxMToyNjoxNiswMjowMCIgc3RFdnQ6c29mdHdhcmVBZ2VudD0iQWRvYmUgUGhvdG9zaG9wIDIxLjIgKFdpbmRvd3MpIi8+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJzYXZlZCIgc3RFdnQ6aW5zdGFuY2VJRD0ieG1wLmlpZDozMWVjODIwMS1jODdlLTJmNGQtYTliYi0wNWJhMzAxZjlhZmYiIHN0RXZ0OndoZW49IjIwMjEtMDItMDlUMDA6Mzk6NDcrMDI6MDAiIHN0RXZ0OnNvZnR3YXJlQWdlbnQ9IkFkb2JlIFBob3Rvc2hvcCAyMS4yIChXaW5kb3dzKSIgc3RFdnQ6Y2hhbmdlZD0iLyIvPiA8L3JkZjpTZXE+IDwveG1wTU06SGlzdG9yeT4gPC9yZGY6RGVzY3JpcHRpb24+IDwvcmRmOlJERj4gPC94OnhtcG1ldGE+IDw/eHBhY2tldCBlbmQ9InIiPz725RFTAABaG0lEQVR4nO2dd5gkV3W331tVnSfvzO5szlFhhQQIZLAwGPMRjMEYgzBBAmxjAzbYGNsSDtjCmGiSkBASQQEUEEEgWGlBMkYJ5QQoorRabd6d3N0VzvfHrdtd09OzO7PTsz2zc9/n6Wemu6urbqVfnXPuuecqEcFisVimitPsBlgslqMDKyaWWlqA3wdWNbshltmFN9UVKKUa0Q7LDEBEWkAG/GIfXqoF5XqrlVK/bXa7LNNHI8Mc1jKxACAiDjAQlvbyxM2fYO+vbwZ4V5ObZZlFWDGxGDpA2PfQefTf/Tn23XgZ+P6ZInJcsxtmmR1YMbEgIjlgL8XfEv32HLL7irj3/Jjgf68FuE9ErC9rOSRWTCwAzyMaQJ45iw4O0FmG/K5tRF/9EFEYAKxocvssswArJnMcEcmD/LzU/y2UXEsmDOj0PVpLIengMfp3PwtwcrPbaZn5WDGxnCLiU+z/OkQHQClkQNHiODAS0PfsEwSlkW+LyEua3VDLzMaKyRxGRBZBuJV9X6Fl6B5w00Qpn6DXZyQdQeCy/EO/S+narwHcICJuk5tsmcFYMZnbnCDBfkYOXIEQQJRCIgcRQADS0AbPXP1Vyv17AV7W1NZaZjRWTOYoItIBcs1I/3X4/p2IEwJZkFRiqQzFMswb2c7wfTeAyLUi0t6kJltmOFZM5iCxu7I/CnczvOd/cJwR3EoutIcCIgFEUSy00+L3Mfz5dzD82/sA3t6kZltmOFZM5iZrkDLF/ZeTlgdJu6BHRYwAAUo5IA44IUMiqGKZBcPDyK1XI1H4BRFZ2OT2W2YgVkzmGCKSAh6Myr/B3/UfpKWMF2WIQiAaBqeEIxEiEThl8mmHdAbcMvC9zzHyxEMAr2/qTlhmJFZM5h4rkTL+vvPIO/tJhw4qbEUkDQ44HuCAKMARvFIfhEAGckP7Gdn6dSK/dI6IHNPk/bDMMKyYzCHiWMlDkf8k/p6rSDkKJwJHPJTkQWVRLkRuHDNRLnkRSmXAyeAoIXvdFxm492cAD8SDAy0WwIrJXOMYpMjQngvJOSGMOKBAomGiyIEoRxRRtVCiDG4A6WyBkBQ4UCiVSN9+FcFwP8Apzd0dy0zCiskcQUQKIPeODF8PwxchfhnIASHKGQRVhigDEVRKXEQOvq//9cUHBYQwsuVrDD90O8Av7CBAi8GKydzhOImG6N/9GTxnL64KtPkhIUIEDIESIkkjERABEqEyWRTgUIIU0AJdaYf9F/0bft9egMgKigWsmMwJYqvkFn/4l2S4C8IyKhNpa0RpVwdHwBlCSEEIEgLKp0yWUmmItAsIFAcB5dK7/T52X32u2USuWftmmTlYMTnKia2GQQn3MrL7HAqUcCKInDJIEXBBwFGAO4w4QiRZVARIgHJc/AAdS/EcigLRiE8mHCF3+w8o7XgSbFexBSsmc4EOEEr7f0i6dB2OP4LnQhBB6DmgvHgcDkAEygcnipPYHFx88rkUUdxd3NGZx3E9/IEA98lfceDHF0AYXCIiK5u2h5YZgRWTo5/jJOzHP/Adcu4QTgQIRMoDN02oFKgUCkC54PrglLXrE2UIRvoh5SKpNCOlCH9kmCiKSOXStKVdMjdcSP+dWwF+a7uK5zb25B/FiMgHIfr5YP89eDwGaI0IfIAMQRTHTSU2O8QDpTUFBUiafBZGimVK4uKlPVQIIJDKMTI4hDuwm/KdP0aCMsCaI7+XlpmCmmqpezvVxcxFRB5m92lr6bsBPxgiUoKL4IYhSiJwhYCAyNPdwZmoAx7zKf5yiNK90F726M+nyAcBXimKM2FD+gpQ9qCzBN4AkG9j4EOX0HrKHwI4Sik7TeQswU51YTkkIpIG1tJ/H2F5hDAMEQlBRfoCiiIIw0Oux5VI/+MAngAOmQByPoQKwjSQchi+6SeEw4NgSzzOWayYHL2sjaISFJ8higKUAsfRlqRSisoHh8CLdNJJmBJd6kRBtgS5MgRAMQ9DwwNkbr6W8sP3AtxiYydzE3vSj0JEZCkSPFAauAEcwXUdrRsqqlgnuAkxqXiq0Zh1uYSIEnw3ouzG34cObuCAQOAAUUiHGsb//lco7dsB8Lrp3kfLzMOKyVFGnFfyVFh6gvKud0DkI/igAkAQU14giiAKJrDGiNCJiJKhMaUAFy9I4Q1Doa2FoH8H0V1XEzx0G8BVIpKdht2zzGCsmBx9dENEce/3KKR2EUY+IiHasxEcV/+FKB4afHCCuGfHjSAdOoAi9BwC1yETOhTCNOHeQcppKBSK9J33YYLtjwO8aXp30zLTsGJyFBEXPtoVlbcR7v88XgiOIyjj4hACke76dZjQ2Q8cUAKZ0EFFitCBYioicOLendDBLTik28Evl+jZ/xTBz64CkW/oOrOWuYIVk6OLY5ARhnd/lqy3A8qgUg5KCVEUEoZCEERINDY2Mh5RbJUQKIg8fAdCR3DxQQIQn6AcEQSQEUipkKHvfInhx38D8Nbp2lHLzMOKyVGCiHjA3cP7fwojnyedCiFohdBHCCudN5UOnDi/wHF0+KTyeRhSSR1SSvcIR+jENuUQOg6uRPozJ4R0hGQgq8AtAVGA4xUp3rEVyiNfFJHfO4KHwdJErJgcPWySqIj0fZG0MrHVjuq3B8ktrOYtOZX35jMvTr9HOVpPBNKh/osHpbRQVkpnsYUQeQ6BGuHA9z5HsOUigOttiYK5gRWTowARyQH3lvb9nEx4EyqOrYYcOsBqRENPvCVjvswEaCFytcuTDmOBccB3IXSAUCBSRMolctO051P0+H143/4YoV8GWN/A3bXMUKyYHB38buTvQw78B54M4/rgOS34auDgv5Kav3Eym1K6zAlEqEhnuvpuAER4YYQS8GMxSQMpgcjzKKZciiUfd3CYQjAEpZ0ceOQegD+elr22zCismMxydOEjtpT2XUPWu0vnnUUpBBfxEmIyxkjRp97ER5RyKmIyKjFWIHLAd0Mc/LgCm0Oo4nX6+m/oCiqXJo2LWyzj+GXKUZnwlu8R7NvxMRGxNU+OcqyYzH5eHpV3Eg1+iSgoEoWgHI8gLOF64/Ta1Ml2rwz4MpZJJQir3RtREDqRDpZEOmHNiXSVtggI/CJuGJJ2XfA8yIJKQ27rBQzeeBXAd+Oua8tRihWTWYyI9EL0veH9V+CphwmDOFbqFkEVJ7CGQ5/+MCEqgQu4QOTh+SncMEWYdSEHqRBSQyNQKlGOfEYiKPvQemAPg1d8ksFnHgN42+Huq2XmY8VkdvOGyN9DOHwRod+Pq9IoFJEjeGmI/EOvwPTcSKILp/KvCJFycCQuTO/E4iIKAg8lLv3lkMgDzwUlDrgeUS6FZCClE2bpjIoMPvEgwIUisnjajoalqdh6JrMUEVmGhE/u3f52MuG3yPlxnkfQEncDO0CKA7kRMl5IUCqRVaCiHEo6GPIVTk4RZfsplgZo9yBTbIVtbYQPtvPkvQEdmSV0PXU95BVDRUGcNDkvi1v2IShBMdIRWAeKafDTHpFSpEPIhIJTCtjjddFdHoFUjqF/u4rCiS/5mFLqI009eJYKjaxn4jVsTZYjRjwz35NB8TG8oevJOOAGgOTiKvMlQjfQxdMA10mRUlm8MAVOgBSfpS0LZFyGilkysgpfXozb+UroXIr3e89hVRCy947/Zc+lz1De/ihdmRRZAb+vHxHwCh5kMuA6lEaGUGVoiQLCEMoRlFIKJ58l6w6ACghDBU8+gGx+0VkicoFS6ommHkRLw7GWySxERE4hGrhp6PH3kC1/D1d8EIU4EDo+oVONsboeyAikfT2FBU7s+7RBeQBG1PHkl34BJ3scbqbLbOIEtM1xG8N99N9+LSPXfY3Cb/6PQlhCZT3KQchIpAjFIaMcCgoIyzqdNpMmUNBXLJPNQ6oMKvLY7fTQ+i/fotVaJzOGRlomVkxmGfHQ/pHy3stI73wPhAOAA15AkNID8yIBJ/JwIxenWMJNoQOncTcubhtBtJyg673QeirZ1vXEvlEBKCulgnhbHnoK0J9TGmbgtusY2PpN8g9fT8ofIKUER3S2red4OCpNUCoTRgGplIOT8iBVpjgC2VQLhFn8416M988XoXItpyilbmnGMbRUsWIyhxGR5xEN3bb3V69mXuY2KJZ0hlkqoJTSAVI38PCCjC65KBF4vp7bwvEoqeNg/sdwCs8lle0xq20DhpRSdfuS4xKQxwJ3SmmYvluvZviK/6Rn54OkoggCLWC+SiGOS9Z1tZVS9KEVigGUBj3axUMyGQ584Et0vvytYOvFNh0rJnOUOG1+eGjfpUQ7/5xWZ0SPiVEuoRvhewKSIhU4uDKkfxN1Eqos5fQx0HEGbsvvkMkvB3gtsBUoTfSGjmM1xwN3URxgcOtFjFz+Cdr7d+gZ/4iIghARcF0ARRQJYUsHSrJ4fQMgEUOnvJr0351PqqXzdUqpHzT+SFkmihWTOYqI/JmE2y7Z8dgJ9Ob3ovrRwREHfMchJAtRjlTk43IAFAz2/Bchx5JvewGpTA/AO4ArlVIjU2hHFng+8HPKIxRvvY6+r/4T3f2P46ZCKAUMR+BkIVvupE+VcLIemZJP2o8YaWtn52v+lhVvPRPAHc8iskw/VkzmICIyD2TPnqc/RC74LAVABhXKE8TTCWW+gARdqHA+XvqFON4iosUfIp3qMKvJKqVKDWxTGtgA3Eu5SOmXP6bvik+Te+YespmIoFQiV5wH6SJ7gyFy7RmUD+VSxMiyzbR9+Fvkl659j1LqK41qk2VyWDGZg4jIa8Xf9YPh3/4+Ofd+ZAjcdI5IjRCmoCSQyTsMDm8m2/Up0i3Pw023gbZELgOC6bIA4mr0xwD3URph8OYf0n/5x2k/8ASFfsDth1TEkAe+DzlXEeTa6X/NWSx804cAWpRSQ9PRNsvBsWIyxxCR5UjwhL/3uzi7/xrX3Uvo64zUwMvgploYKS9CFd5Kpv21ZFvWA+olwK2NtEQm0E4XOA64m+IQw7+6lYGrP036gf+lMyrqivhlIfTBz8GejafS9hdfpm35pvcrpb50pNppqWLFZA4RP/VDSo8y/NgfkePXiAt+CspuB1H6fWRaTiOTX4VyswDvB77ezCd9LCrPAW6X8iDFh29j+Nufpf3u6/EcXeoRF/blXPzTPsb8P/w7lJtap5R6pFltnqs0UkwQkSm9LNOLiGyUqCRDT7xVBu9Hosc8KT22QPY88ceyb88V4vt9EnN8HMOYMYiIKyLPFRGR0rCEN/9Q+v/uJVL8oxaRP8pK8Gpk35tbpP+2n4iIfLbZ7Z2LTPX+H6UFVkxmLiKSExEp7vtfOfDwPCk/XZC+R18jA7uvl8AfNCJyqszwGfREJCsim0VEpDggw3dfK/s/8XaRP2wXeXVBhs5+u4QjgyIiL2x2W+caVkzmCCLymijYJwe2nSv7n/5L2f/s58UvPmtE5AWiM1RnDbE4Pr9iqdz4Q+l76wul/9/fJfvuu1VE5KPNbuNcw4rJHEFEPlws/kYG+h6QMCxKghnlzkyW2FJ5jhaVEQkP7JHBHU9LFIYitkTBEaWRYmIDsDMU0RXdI2EERQ7gROBBoHi0pKBLMk9Fs0Qp9UwTmzTnaKRBYMVkBhMLynOBXymlhpvdnuki3k9HKRU2uy1zDSsmFoulITRSTGZ0L4DFYpk9WDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyYWi6UhWDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyYWi6UhWDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyaWOYGIuCKyWUQ2NLstRytWTCxHNSLSISIvAYJisXiP75d/IyLvFpFUs9t2tGGnB7UcdcRCsQH4HRE5d2BggN/85jd885sXUihk+chZH6W9oxMgp5QqNre1zaWR04N6DVuTxdJERMQBFgAnA98rlUo89dRTXHrJZfzwRz+ir283mbTL0OBeOtpaOPMjH0MpNSIic15QGoW1TCyzGhFpBVYA90VRxJ49e7jxxhv57ne/y6233krKy1Mul3DdgCAcobWQxi8P87nPn88f/L/XARSUUsPN3Idm0kjLxIqJZdYhIi6wHNggIteUy2UeeughLr74Yq677jqGh4dRShEEAX7ZIZdLUy4P4KoA5QRk0g5LFy/k0suvZV73ApjDgmLFxDInEZEO4ATghiAI6Ovr48c//jFbtmzhpptuwnEc0uk0vu8TRRGpVArHd8EtE6gA33fxUu2U/SL5wiBv+tOX869nnYvnZWCOCoqNmVjmDCKSBpaig6nfHBgY4JFHHuHiiy9my5YtAJTLZQBc1yWKIkQEESEMQ5AIoghRDkoplCM4TsTQwDA/+tGPeMHzT+VVr3wHQBaYc2LSSKxlYplxxMHUhcBaYivkmWee4brrrmPLli38+te/JggCPM9jcHAQz/NwXRcRIYoiQF+XSilcCYjwiMQliCKclIvnRJRLg6RSIae++GS+fO5lZDJ5lFLzlFL7mrnvRxrr5liOSuJg6jrgjjAMGRoa4u677+aiiy7ipptuolQq4bouQRAQhiGZTAbH0RZHGIYEQQBoC0UpRRRFeKpIJDlQWUIJiGQEz3NwVRpPXHy/jzPP+gCnv/tDKOUAOEqpxt1hMxwrJpajBhHx0G7McSLyg2KxyPbt2/nud7/L97//ffbv38/Q0BDZbBbHcfB9H4AoiirXnrmGlVI4jlP5PooiHKcIksFx8jiegx8MxCJTQAUpstmIQn6IL5z7ZZ538itRSi1QSu1qxrFoBjZmYpn1iMh8YD3wf0EQsH//fm666SbOP/98nn76aUZGRgAdD3EcBxGhWCwiIuRyOaIoIgxDHRcBHMfBcZyKiCil8DwPRYYgBD8okvZSuG4KCUMQvWypGFAcGeS7V13K5ue8mHS6sFNE5pSgNAorJpYjhojkgDXASVEUfX1oaIjHHnuMc889lwceeIC9e/dSKpUq4qGUIpPJVH5vLJFSqUQYhqTTaTzPM+smDMOKkJiXSArXjQjCEYJyhON4KFJ4noMgIIpCfh7fu+qHPP/kU3j9G/4KYB5gxWSSWDfHMq3EwdSlaCvk2lKpxLPPPstVV13Fddddx+7du9m1axe5XI4wDFFKkUqlGBkZIZPJEARBLApaLFKpFI7jUCqV8DyvIjqmB8d13Uqvju/7SOSSzTkEYQnHcZEwQ+iHeCkFUQCRwlEejhph3aYu/ueLF7Jy5Qko5a5QSj3Z5MM37diYiWXGIyIFYDVwbxiG9Pf3c9ttt3Heeefx1FNP0dfXVwmSTms78EAFKAIQDyUeEIEq4yAoyRKFKaKwSK51hLe/4w186MP/g+N4MAfG7lgxscxI4mDqCuC5wLdLpRJPPvkkn/nMZ7j55pspl8sVC6NYLJLJZKb9+pF4YLyS+C8BKB+UFhckQ8prwZGIUrCXdGaIz3zmHF756tMAtUwp9fS0NrDJ2ACsZUYhIguA44Ctvu+za9cubrjhBr7xjW+wZ88e+vr6yGQylEoloiiipaUFz/MIw7ChF3M9IqeME2VBUqBii0SVibNREFyGRgaBiHyugOdm2bplK6f+3mvI5VqeEpGjXlAahRUTy2EhIllgFXBsFEWX9/f388ADD3DxxRdzyy23MDQ0BEAYhpXuWvO3WCwSBAGp1BEoKaLC+B8XJWhBQQBXWyZAJpcCFRCEwsC+Mtdccz3Hbb6At53+13he9ikRySqlStPf2NmNdXMsEyYeYNcLbAauKZfL7Nmzh+9973tceeWVbNu2rZKZajJSTbDUjJVJpVL4vk8QBBQKBYrF6Q1JhG4JJ8yhogyKCEcNgfKJSCGSJZQIvACIIPDIplvwCHCdfZx3wVc4+YWvxnHcTUqp30xrQ5uEdXMsRxQRaQFWAveFYciBAwe45557uPDCC3nooYfYuXMnuVyOVCpFEASUy+VK12wYhhQKBUBbJL7vV8RmeHi4Yq1MX+NTRCrCdUraIMGDKA0q3q4XoFxBBPwoxIt8JAwYCX2uvOwyjjv+hRRaun8tIouUUs9Ob2NnN9YysdQlDqb2AhtF5LpyucyTTz7JJZdcUunS9X2fTCZDuVzG87w449Qhm80ShiG+71cExXTrJgfjmeWnk1B5oIq4+KjIw4nyEGWIlBC5I6h0GZ+AMFCknQIqiEi5gishrgp5z3vfxF+9/+OmdyejlCpPa4OPMNYysUwb8TD/VcCdYRgyMDDAT3/6Uy666CK2bdvGvn37EBEymUxlnIzrunieVxkzUywWKzkjrutWxs84jkMQBBWXx3XdSgbr9O1QCtQIqBIodJwkyoHjgxoipAQoosghnU1zYGg3rZkMCkUUeFz9g0s58aTf4fkvfDWumzoWuGt6Gzx7sZaJxQzzXwCcLCJXFotFnnjiCc477zxuu+02du/ePe29LjOWIM+r/+gYPvofZ9PRthLluOuUUo80u1mNwuaZWBqCiHShM1NvDsOQ3bt38+1vf5urrrqKffv2VSyNuYwSj5bWQd77vtM54/SP4KVyL1dK/bTZ7WoUVkwsh008PmYJcKqIfHV4eJi77rqLCy64gIceeoj9+/dX3BHf93Fdt9lNbirZLOw/sIuTTtjIF79wHkuXH4vrpZ+rlLqz2W1rBFZMLJMinvphHrAJ+Jnv+zz11FNccMEF3H333Wzbtq3SRWuCpa7rkslkKkP+5yrKOYCr2lCh8La3/RF//48fJZNtQynnqOjdsQFYy4SIx8csB34VRREDAwNs3bqVr3/962zfvp2hoSGGh4fJZrOVcTLJgXKm2NBcJoyKRFELaWnhh1f/hOc8dx1/8Ioz8FL5DcCsF5NGYi2To4w4mLoYeB5weblc5oEHHuCrX/0q99xzDzt37sTzvEqhoXQ6XbFGHMchlUoRhiGlUgkRqQzxn6ukc4OUBvNkvU5K5Z087+QFfOGcS5g3by2umzpeKXV/s9s4FaybYxlFPMy/C51YdlsQBOzatYvvfOc7XHvttTzzzDOVwkJJa8PERGrrpoJOfT8iXbcznCA4QNrrJJ/u5MD+XbR0DHPGu/6QD3zwS3heHiCvlBppdjsPFysmFgBEJA/0AE+ICCMjI9x22218+ctf5tlnn2Xbtm2jksSSYmH+mvhIskqZsUbK5fKcD8A6DkRhSOA75NMFlDNEvq2fL593Pied+BpcN32CUureZrfzcLFiMoeJx8csAI4BrgvDkO3bt3PllVeydetWnn76aYaGhhARUqnUtNcLOdpRpIEijvKIwhQpx0XcXZz60vV86Ytb8LzcrK4bawOwc5B4fMwK4P4oiujr6+P666/nwgsv5JFHHqFYLFaqkAGV6mOWqeIQiQICgkCB6xGUPO69+zGuuPJznPbmD6OU14Et82gtk5lMPMy/HXiZiFxqgqnnnnsuDz74IM888wzpdLpSMzVZYNmIiT0/U0RSoEJcVxEGCs/LEAU+jnuAVWtauPzy/6NQ6MF1vfVKqYeb3dzJYt2co5g4mNqOHh9zRxRF7N69my9/+cts3bqV/fv3V+qiJiu3m3hHKpUinU4DzPkckUYgIroAdcqhXC4ionAkC9EIuZY+3vKWV3LmRy5CKRdm4RSjVkyOQmI3ZiHwsJnWYevWrZx33nkcOHCAvXv3VoKhw8P6ek2n0wRBUOl5MbkiyddcD6BOFeWERKGH47pEUiSKwFN5PHFxvH247k6+esEVPP+Fr8Nx3I1KqQeb3ebJYMXkKCERTN0kIluDIODxxx/nU5/6FHfccQelki7uFQQBxWKxYnnEv62IiOmtMZ+b70xpRMvho5yQwI/AUbhefNwlC6ED0SDtbT7Hn7CSL5xzNflCJ66bmlWCYsVkFiMiCmhDT/9wfxRF9Pf385Of/IRLL72UJ598knK5XLE4isUirutWXBfTO2OmdUjOF5O0UJJp8ZYpoAKiKATHJPA5EIEjKQiErvYMytnJBz/8ft7wp/+E46QBPKXUrFBx25szC4kzU3uBJ0WEcrlcqVZ2++23MzAwMGr6SyPSuVxOz/8SWxymTggkp8CsJpiZTNZkTonl8HGdFC2tGdo62li5ah3zexbS2d7GymWr6GzvZumSxXgpn+6eVqLIR6kUSqk5OYmXtUymkTiY2gqcAPxvGIbs3LmTK6+8ku9///vs2LGDcvnghbsUU8sTmepzR6Q6VscIlQn61iN5PTiOg6hq/CY5F7CxpkyFNvNbY0lVqrM51fmEk3MKm0Q7k9Fr4kPmszAM8TwPL572Jrl9UxXfJOd1dHTQ1dVNz/z59PTMp3fhQpYvX8m8eT0sWrqM9rZ2CoUW0uk0rutVjgHwFeApIAUcAG4DHptNOSfWzZnhxFbIQuLMVDPA7jvf+Q4PPvggAwMDBEEwobEvaornWg7z9FStHyqWTvJmNj1I9c5/8sYXFSSm6pRRmbgm6zZpaZkYj+d5+tiE0ajtmiCzKUqdy+Uqv03WnXUcR88SGPTT1dXFokVL6Oqax/wFvfT0zGfp0mUsWrSEnp755PIF8rk8mUwGL5XCcdzkfr0fbWXsRAvGADAI9B0NFeutmMxAYitkHrpy+9YwDHn88cc599xzufvuu3niiScqN1Ly5pxJiWX1zqV5iid7i0z7jfVQW54x2VWtnDD5JK98njwWtSSDyUpGX/DmN6Y9ruvS0dFBd3cPvQsXsmD+Anp6eljQ20tv70IWL1pIoaWFfK5AOpOpDHKM9/U/gG1AH7AP2Avsj9+PAP5siX0cLlZMZhBxl24P8FsTTL3yyiu47LJvs3//fgYG9PwxxlUYGRnBcZwxAdVx18/UCi47MjU3yYwiNpgcF+MupNPpygBCIxJJ8XCUP0o0zXfme1N0OnGDV5YXEVra2pg3bx7zurvp6ppH97xuli1fzooVK5nfM5/Ork5yuTy5XI50OoPnuShVOWb/DTyOFoo9aLE4APQDQ0qpOV9jwYpJk4m7dBcDJ4jID8Iw5OGHH+a///vj3H//fZTLZYaGhuIbxhmTC2K6d4eHhw/p5jRDTJLXhBCOyaRNWhRmMKCJdSRdmjAMcfBHxVmMpWPWkcvn6Z43j+7u+WSzWVpaW1m3dh2bn3MiCxcuIpvNkS8UKBQKpFJpXNdJisU/AjvQArEnfu0HhoCyFYtDY3tzmkRcuf1Y4BdhGLJv3z6+9a1LuOyyy3nqqSfjAF11BK4WD4dMJlN5kpsnePLpfDCiKWq1c4hrJRkLqXUn9F9wnGoui8ldCYJy3GVdnXDLBD5NjovrKua1z6Ozs4t587pJZ9N0dc5j46Zj2HTMsfT0zCeby2o3JF/A81K1vVDvQ8cn9qFFYy86ZjEAlJSaakTJ0kismByCuGZqN3p8zNeHhga5/fbbueyKy7jhhhvwPI+BgQHyLQVSqZQesRsI6ayelDv0o0qXbzLeYG7MQ7k5zhR7cw71e4VC4j4foRocVUoLShAInutQDgMkivBSaZQSlBI62ztIpzN0d8+jUGihpa2VVavWsPmEzaxYsZKWllbS6RT5QoFsNlexzGL+E3gGLRa74//3ElsV6HiFHfI8i7BiUodEsaFjiLt0n376aS782oVct3UrO3fuBLT/PzIyRC5XoFwuE4aC56XjXBGHIAiJEvVCkrGGiU5ApaYY8+CQYqSodiALoFDKwXUd8vk8bW29LF66lCiKWLRoMb/zohfR0zOflpYWuru7yWZz5PJ5Muk0KhH3AP4GbVEU0T0h29HuyDBzILA5F7ExkwQikkHXTH3IdOn+6Ec/YsuWLfziF78AqMQ4TI+FoTamkPw/uWxtMlltGjxUu2KjKMKJ77lkV6yJT5jPjJthgp/J/6PAH9VmI2DpdJqurnnMm9dNW3sbhUKBZcuXs379BpYsWcq8ed0UWlrIZXOk0ulR5Q2As9FBTNMTshPthgwAJaxVMWuwAdgGIrpyewvwfGCL7/vccccdXHTRRWzduhXPS1MsFlFKkU6nK4WWTczDdI3W64kAIBrdNZrMtajtSgXGrMd1goTrMTrzNbmeyvKxK1EotJDP51i8ZInuFk2nWdDby6ZjjmHVqtX0zF9Aa0sLmWw27gXxkufyLKoxin1UA5z9wIgViqMHG4CdIvH4mAKwGrhHROjv7+fSSy/lsssuY/v27eRyucr0l0DFUjC9EeZ/YyXUWimVvIwwQt+j+uYXRHsWSscqRKgksDmOg6Ociki4roMSVXGLKj0guRydXV309Mwnk8nQ3tZOS2sry5cvZ+3a9SxeuoSuzi7tfmQycS9IJV7xTeBR4Gl0rGI/VdEYQPeC2MCmZdLMGcskFhAzf8x2EaFUKvHTn/6Ua665hi1btlQm3R4eHq4UW86k0qMEBRg1HYSxEGp7QioWiDPaEqktZGSqw5suY8/zSKVSdHZ20tHZyfzuLgqFFlasXMHatetZsKCX9o52WgotZLM5srHoJVyQzwNPoKdh2MdosbC5FZZRWDdnEoiIB2TQvQZEUcRTTz3F+eefz3XXXcfAwAAigu/7lZs6nU5X3Y1Iuxme52EG6CVjE8lxKvXEIptNV6wYY110dXXR29tLd3cPruuSzxdYvXo1a9aupXdBLx2dnXR0dOhErFQapyoWX0AHMndTFYpn0RZFP1C0gU3LZLBiMgFEpBU9PuYhgKGhIX70ox/x1a9+laeffpowDPF9P86H0K6K67qUy+WKtZFKpYj84VEjc5MZnskkLGN9ZDIZ2ts76O7upqW1FeV6rFi+nA0bN7F0yRI6u7ro7Oiita2VTCZDJpNNZn+eQ1UstlPN2hwC9isVj1qzWBqEFZNxiIOpC4EXiMjlQRBw22238bWvfY3rr7++knlqxMGkiZtkMjNozIiK4zjkM1AqlSojXD3PI18oML+nh9bWdnp6eli6dBmr16xl4aJFdHV1097eTj6fJ5VOx1ZORSw+iRaInVQHjm1Hd5cOHg0DxyyzCxuATRDnhHQCa4FbzDD/K664gksuuYT9+/eTSqVwXbcSBzFWSSaTIQgCMpnMqIFs8+fPp62tjc7OTlpyDvN7F3LcscezYuUq2ts7aG1to6W1lXTcZZpI7/4PtEjsotr7sZN4pKntBbEczcxKyySRVLYUuEtEGB4e5totP+ZLX/oi27dvo1Qq6e5O8SsioeMTedra2mlta6N7Xje9CxexYeMmlixZSkdHJx0dnbS3d5AvFEin0riju0z/De127EbHKsz/g+j0bisWllnFnHZz4lohJdMbc+89d/PFL36em266kSAI6OrqoqenB8dRrF6zltZ8jo3HHMuKFSvp6pqnLYu2VgqFFlKpdDJe8T9ocdiDzrHYgQ5y7kZncdqxIJajjjkrJiLSDhzQ88fcz8UXXUQmk6FQaOFFLzqFxYuX0N7REVfFSpFKpZPt+zRaGPag3ZBn4vcmEcvOC2GZc8xJMRGRdhE5UCqVKJfLZLPZSqGbmE+ha1ckg5s70UVuSja/wmIZy5wUE6hYJivQk1SZ2hV9aLGw+RUWyySZs2JisVgaSyPFZGplvCwWiyXGionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDaIqYiIgtHGuxHGUcUTEREVdE1gCRiBREJB1/ZsXFYpnlHLHq9CLiAoHv+4gI6XS6dpGXAo+ip6/wgdDOdWOxTC+zdeLytwwMDPC2t72Nxx57jJNPPplVq1Zx7LHHsmbNGhYvXnx9S0tLcrpOROSDwN3AU+hZ+EqAAJGdqtNimVkcEctERDpEZP9ZZ53FD3/4Q8rlMkEQ4LounucRRXq+70WLFrF48WKWL1/OSSedxLp161i6dCn5fB7PG6N7L0BP8bkfCIDATsRlsUyOWTUJVxwPie677z5e/epXV6b0HB4epq2tDd/3UUqhlMJ1XRzHoVwuIyLk83kymQzLli1j4cKFLF++nA0bNrBp0yYWL15MJpPBdd3k5t4NPIy2ZPaiLZkQEGvJWCxjmW1i8vxisfjLd77zndx5552ICEEQkEqlKJfLFTEIwxARqbyPogilFJ7nUS6XAUilUiiliKKI9vZ2li9fzvz581m3bh2bN29m5cqV9PT0kM/nR7lLMacAT6InKi+jrRkrMpY5zawRExFJA6XLL7+cj370o4yMjOB5HmEYViYcb9TOmHa0t7ezYsUKNmzYwLJly1i/fj1Llixh4cKF5HK5ihWU4IVod2kf2pKJlFJRQxplscxwZpOYvPmZZ5759pve9Ca2bdsGaOvCiElSVA6XKIpwHAfHcRCRioUD4DgOmUwGx3FoaWlh/vz5rF69ms2bN7Np0yZ6enpYsGAB2Wy21l16J3AfWmQG0FaMb0XGcrQxK8RERBZFUfTM2WefzcUXX4zrupTL5YplYt7X3MSTxoiJsThq9ycMQ5RSFcGJoqiyfcdxKBQKrFy5kmXLlrFs2TKOOeYY1q1bR29vb0WIErwDeAh4Gm3J+CboKyLKukyW2caMFxMTdP3lL3/JGWecge/7RFFEqVQil8sRRVFFVKa6fdd1iaKIKIpGWSRGXEyMRUQqL9NuIxTJZYyF09HRwcKFC1m7di1r1qzh2GOPZcWKFXR3d5PL5ZIi8zF00PebQDfakgmtyFhmA7NBTF4zMjLyw3e9613ceOONpFIpfN8HoFAoUCwW68UuDgvTfhEZtT4jDkk3yHXdcQXGCJwRuXqCl0qlWLRoEWvWrOG4445jyZIl9PT0sGzZMhYsWEAmk0m24cNoC+bbwEKqXdhi3SXLTGFGi4mI5IDhz3/+85xzzjmVp73JK3Fdl+HhYQDS6XQlx+RwSbY/aZGY71zXrcRSjDCY5ZK9R0mrJQx1ukomkyEIglEulHGTTLvb29srsZmWlhaWLl3KqaeeykknnURPTw/z5s0jnU7XukunAfcDO4E+IMIm4lmawIwVE+PePPTQQ7zxjW+kWCxi0udre28cxyGdTle6fQ8XIxa1ro4RABPkNe5LUmyMaJiEOPNbsx7TfW3aW+s+AZUeqmS3dRAElXWtWLGCRYsWsXz5clauXMmaNWtYt24dPT09ld8keD3wGLCdOPBrrRjLdDKTxWSt7/sPn3nmmVxzzTUMDg5WrA/Ti+P7PplMpmIpTDUAGwRB5Uav192cjIMY8UhaJuZ9Ms/FiBBUxap2nUnXJ2mp1Oa3mCQ80JZYKpUiiiLmz5/P4sWL2bBhA6tXr64k4rW3t5NOp5Pr+Dt0fsxv0JbMIDoRz1oylikzI8XEDOT7xS9+wemnn04YhhU3wTI+xlpyXZdMJkNvby/r16/nmGOOYdGiRaxatYr58+fT2dlZz5J5A3pw5E5Gx2SsyFgmxEwVk9fv3Lnzu+9///u56667RvWOWA5O0lIyVo7J/k2n03ieR3t7Ox0dHWzcuJGTTz6ZjRs30tXVRUdHB6lUqvY4vx74NXpw5DCJ3iWLJcmMExMR6RKRveeffz6f/OQnK+6L53nWMjkEtXEc85nB5OYYwTABYHOMN27cyPLlyyuWzNKlS1m6dCmdnZ14npdc798AvwUeB3ZQjclYK2YOM6PEBFBAdM899/ChD32IRx55pNJr4/v+lGMiRzvJXiLzvjYJL9nNDVSOrznGyW5xk+3b29vL4sWLOeWUU1i0aBHLli1j/vz5tLa21rpL70MLzENod6mIHVIwZ5hpYnJyuVy+9b3vfS8/+clPKrkWQRDgeV5DG3s0kxSEZO5LKpWqfJbECIzneRVBMQKTDAabwLJSis7OTnp7e9mwYQPHHXdcpZepp6eH9vb2SvA55t3owO824Fl04NcGfY8yZpKYpIDyZZddxr/9279RLpdJpVKICL7vk06nK70ilvrUdmXX5skk81ySWbtGcExinVkHMKrXysRfzO+Mm2S263ke3d3dlVoya9euZcOGDaxZs4auri4KhUKtu/RB4EHgEWA3MIRNxJu1zBQxcYBwx44dnHbaaWzbtq3S9VsoFBAR+vv7yefzDWvs0Ui9PJmkeJjiUcmMXagKh+kaN1aFWVcyzyYpRLVJhmZZM1YpCALK5XLFktmwYQPr169n7dq1zJ8/n0WLFrFw4ULa2tpqLZm/QbtKj6IDv0V04NdaMjOYmSImC6Mo2v7pT3+ac845p1L0yPM8isUiURSRz+ennJR2tFMrHrWCkRyomMyNMcslM3hrxcisozZPJvmZ67qVJDtzDmt74oyYOY5DKpWiUCiwZMkSVqxYwcqVK1m9ejVr166lt7eX1tbWeu7S42iRqRSssiIzM5gJYuIA4d1338Vpb/lT/HJY81Xyr7V+j2aMKCmlaG9vZ9myZRx//PFs3LiR3t5eVqxYwbx582hpaakNxp+OLvHwOLAH7S5ZkTnCzAQxOXloaOjW97zn3dx8y80gyRwHKyZzjdo4T9I6EhEWLlxYcZGOOeYYNm3axPLly2lvb6+ITMKSeT86JvMg2pIpY2My00azxSQDFK+44nL++Z//qc7XVkzmEsn4TD2XyoxxMssm3bG2tjaOP/54Vq9ezcqVK1myZAlLliwZ5S4lOAMdk9lGtSqetWSmSDPFRAHRo48+wgc/+EF+9asH6pQRqLVSrJgczSQFJDnWybxMzCcIAqIoqgR6TdC3VCpV1mECzoVCoZKAd8IJJ7B27VpWr15NT09Pvap4f4lOxnsYHfgtWytm4jRTTNYHQfDge97zF9xwww2VoNxobPr8XKQ2g9e8koFjIyCmyxuqAyVNr5QZCBmG4agMatd1aWtrY/HixZx44okce+yx9Pb20tvbS3d3d60l83HgV+gR2M+iYzIjVmTG0iwxcYBwy5Yt/NM/fZgDBw7UySOxQjKXqXV3kiQT64Ax/5sHk+u6Y4pT1ebJJMlkMvT09NDR0cGKFSvYuHEjxxxzDCtXrqSzs7PWkvkYulfpYXTgdz9z3JJplpi8bGBg4KdnnHEGt9/+S9LpdOVpUsWKyVzDZNwmE+HqFaiqt7yp/5L83nSHJ2v7HgyTaW3cpyiKKukInZ2dlZq+69atq8Rkuru7K9OhxPwX8ADakjExmTkhMs0QkzYR6Tv//PP5r//6r0o+QhCU4xNiRWSuUpu/UhuINWJiclWSVkuyuLgRkGTy3XiDH817s24z2tpYNkkXyrhUYRiSTqfJZrOkUimWLVvGmjVr2LRpE+vWrWPlypV0dXWRzWaTIvMv6FkKdqJFZg96HuyjpsxDM8TkHx5++OFP/tmf/Rm7du2q1HSt6dKrQ4QVGstMJjmOrL29nXXr1rFixQo2b97MqlWrWLlyJfl8vl7g90/QlswO9MRuZs6lWSUyR1pMThoZGbnjHe94G3fddU/FDE2OWh1L0jq0YmKZuSTHLxkLykxZayoCtre3s379ejZs2FDJ+F25ciUdHR3j1ff9LXqQ5AFmuCVzJMXEA/xrr72W973vr1FKm5CmOHQQBKTT6To/s2JimR3Uxnlqy0GEYUgqlRoVLE5aMqtXr+aEE07g2GOPZfHixZWqeDXuEuiqeKZ3qQ8dk2m6wBxJMTlt27Zt33rrW9/C9u3bGRkZwXVTZLPZSiPGFj+qjVlZMbHMXIybnhzXZIK55vsgCCr5MKY+r4nRJDshUil9b2QyGVasWMHatWt5znOew4oVK1i6dGmlvm9CZD6I7lV6Ch2TOcARLlh1pMRkQRRFO84++z+56KJvAsQTjmcAGB4eJpXKxIqdFBArJpbZRVJQDMlcl2SsxCTi1VYT9H2/kidTrzxEOp2mpaWF9evXV8puLliwgPnz59PW1lYrMh9CDyd4Ai0yw0yTyBwJMXGA8NZbb+Uv//LP6e/vRykVJxQlF0megNq/o5ezWGYiyeEAMLpb21grtd3ZJhnPdEQkq9yZOGKyNESypoyp9SMilelnu7u7Wbt2LZs2bWLt2rWsWrWKxYsX09raWjtTwd+gBcZMhzLIFIcUHAkxeWOxWLziTW96E/c/cDdI9aCUy0E86VSOUqlUmXOmrpioCMSrXbfFMmMwAdjaZLjk3NTJIK2pL2PKNhiXKFnRLjk0wCxXW/gKqi6UWYeJ15jpYBYvXsyaNWs48cQTWbduHQsXLqxUxaspvfkBtCXzJLp3aVApNaHiy9MtJh7gn3322VxyySUEYZEoJBFwzcbdwqmapDUrJpbZR/Lmrze/krFCkol0wKjPk+VJa9dhPoOqFW/cJh02SFU6NTzPq/QgmXsrGRROpVK0tLRUpkA57rjjWLx4McuXL2fhwoW0tLQkRebf0bGYh9ECs4M6QwqmU0z0QL5HHuaNf/wqhoaGQNobtrHDQur1Fk0CNbXiTKqBB/swW9Dk7c9uZIphhilHKeTgBdWTmcKjfmZ6mUiP+SyJSfhLpVIsXLSQDRvXs/mEYznm+E309i6gZ0E32Vw2Lr0JwJfQbtKj6J6lp9C9Sz4wpb2tFZNVQRA89pfvfDs33/J/Wl1Vk6vLH+JkHJJDTBdzqHTt6Y6rH+rJcKj2HZqDx6wa+WSaHqY4F/UUD59zCDGf6vE7WJavXiCou6whWQXPZAObhNKWlhZGSiVWrlzF8hUr2LB+E2vXb2D58pX0LlhIvtBCNptFqco18j70zJGmiPjApPYlsQMOEN515+38+Z+fTj6TZWhoiHLYj0g0yq80ZpvxJWt3rpapHfCpDo+YWgBYyZEbnlFfOKrtP5wJzeqP7E6sPfFd4qI65LLNxnEmphKhjH6Y1F6fIge/ficrJko5iESVY+m4VfdnPEw8ph6lcv+YWMzo7alRMZmki6WUwg+DSj0ZPaQgQz5XoKOzk57uBaxYuZr58xeydu0mNmzYzIL5i8jlWvC8FEqpNiYhKEkxaRGRgcD3EYxQhHWFYNTuTPnJaTC1TRu0ukOQbPVEN1n3hE+1wQ07frOIOodMpmZhj0vtOdM33jgL1zMMmiyg9duqRn1fvQQFUHWv0+p+j3arlFKEUQgCrufhVH5bWU8eGJlQW2vEogVYjH4c/ho4FnAZe5hlnP8nw3intPbzid5tk21HcvmJbqOemTJVNTiU6XMk1OZwt3GoYz6Z9U70epgstcdXMblrZapTUh7O/THe/SV1PgupHiM3fp/cZ7O/5j42y3vx3xDIxf/3oYXDiT8fYYJCAgdPWpvQQZ+6T2+xWJpFI2NmB7PhZnpkzmKxzCBmTkTNYrHMaqyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA1hyrOKz/zpJS0WSz0aPU2NtUwsFktDsGJisVgaghUTi8XSEKYcM7HUJQOsBrqBAnoO12eBp9BzuFomxmTnBbboYzYfWIK+/hxgN/AQMMg0Hs/JiIkDrALS8fsw/v2hJt52x/lcxS8/8ZnEnzlULySFnjw6jNflxC/znfk8ojppcyr+zR5gP0fugvSA/we8Enhn3M4UUAT6gW8BVwM/P4JtOhxcYCX64tvJ9LY1AywHsuhzF1E95+Y6aBSTmRw9ituwG/0wONR1PhPYCLwUOAM4geok5T7wK+B/gJ8AB6Zj4webuDxJHngXcCb6SZsGSkA5/m7Memv+1n6eJOlqCVWRUDWfm8+CeLugL3o3/sx8P4y+gV3gMeBs4DuH2L9G0Is+Pu9lfPcxRLf1I8BnmJmC0os+138Xv/8S8B9Mj0W1Gng/8Jfo8zWCPnfmgRLEy413PY3HZI9r7fUG1QfYNvQNeCvwXWamqCjgT4Avou/PloMsezPwPuDuhrdCRCbyerOIDMrsZK+IrJKJ7efhvtpE5EuTaFMgIu8QETXN7ZrsS4nIP9a0tSQifzsN23JE5JsiEk3iuDWTsoh8Q0R6pPnnqfb1ChEZmMS+PCoiXdLg62+iAdg/RyvebKQAvGwa1+8Ab0NbJBPFBb4OnDItLTp8cow9Ving92l8sL4DOI7GujHTSQp4B3A92nprNua4bQQu4+DWSC2rgW/T4JjpRC+Qn1J1LWYbPnDbNK1bAa9CuyzjMXSQ334OmNfgNk0FQcd2kijgURp/06eZmS7DoTgW+ATNF0FBi//fooW5HgdzTf8AeDcNfEhMdEXfAH6EjpME6Bu0HL83sZPalx+/TJwg+TLBrWEm5t+GiXWFiW2E6ODmMDpYWELfvAPx/7vQfuQDE9zPyaDQCv/v6CBiPT6P9k8vo/6NcyJwHjOnV03QgdAkEdA6DdsqMf5NMET1fPs1r+R1FDbolVxfkNjWeA/QtwCvPox9bhQmxnMm2muox2PouNf70Z0Q9fgUDdyPiV7EzwJvRfdSLELfuAH6IjM3idT8hdG9LgZV8/p7dDS/nrCNANcC9wBPxsvkqJ5oN94HP/7roC/SFrTI3AHcR+ODhwrtPp0PnFTnewE+iBayCC0mQ+genuSxcNAn8yx0kLPZAVnTCzZE1a01x7TRT2KF7iWp5RHgv+Pv2mq+M8fnUG2ZzHE0gd5kkN9BX08Z4M1odzS5TQ9tkf6Y5lhXgr4f/576980DwCuA7fH7u4GvAetqliugraxHgAen3qrmBPnM3x4RuadOgCgQkbtF5EQRSdX8RtWsq/ZVb7lGt98Rkf8QHZysZUhEPhsvk2xPVkQuqbO8iMhOEfnjaWrvZF4FEfmBiIQ17fuqiHgN3tY8GXvufdFBzuRxqz32jX7Vbqf2+wUi8hMZyw0i0joN52Air2NEZHudNomI/Db+vvZeOFiQ9koRycjY4z2pVzMyYJNPjVcAm+ss8yDwe8BdaKsjafUkfy91XvWWaySC7oY7k2rOTZJbgX9CP7GS7Smhu4TruVzz0TkAixvd2MOg1pKc7m0lcai6WfUS1uqd76m+attR+/1O4Lo6be+q074jQQfacltY57sB9LX3a0afR0Hvw+eo77r9MfDPU21YM9PpU8Bf1fl8CB0YOnBEWzMxFDoA9wXqJ+M9CpzG6ES8JE8C/4DOXailCx3IzU29mVOinpgcqZsmmaTWbJfPoNDnurY9JcZPyJyudij09fP/6nxfBj7N6Jyq2ofqv6Nd7lq330GL0OuZwoOkmWKyBlhf5/P/Q/e+NDtaXo9FwCeBnjrf7QH+Gp0xWQ9zUq8FPsbYXp4W4A3ooFm9JKq5gou2+JzEy028nAa/3HG248XtOA3dY1KbRLkHLShHArPt16IDqrWxzhI6fvNpxrfKBS0i7wXur/N9Bm25bOYwr71m9iJson7uyneYeV2G5uZ+G/ByxopwgO6V+Wn8frynqnnqX4DuyTkdbaEZPLS5+SRwSQPaPRv5A/S+e+jsatNVPZFs6lqS11FtJwDo85EMMJuUfh8tJK3A7zC2h0uhM0mPlJgIOj3+Yur3rO1CP8hGOLRFN4juAdrC2LSEpWgX6k8ZmyJwSJopJkn/2BCi4yQzkT9Bxzxqj1mIHnMz0d4YQYvPvwDHAyfXfF+It3Mb8PAU2jtbaUUf65nMr9E9eUeKFnSSYz0OAH8B7GDiruFdaHfpPMbG/V4K/CP6+pzUQ72Zbk69HTc5KzONFwAXUrWkkj7ng2jTcbw4yXjsRJ/Q7XW+Ww98k7Fi2wymM5g9Gymh8652HaHtOegUg03Ut0o+hXadJ3OOIvT1dQFj77cUWkzeczgNbRYdjD0AeaqDu2YCJjHtE4x2yUzgbRf6KTpeluuhuBH9BBip893z0YMUbc2ZmYVCuwpHojyCQndSvJmxFoSgr0sTJ5ksEfBh9Aj22gehi+6tfBmTiJ80+0KtbagpJzBTSAEfQN/YtW0dQSebPcThX1SCfkJ8EZ3Fm8RBJ7m9+zDXfbgcyZ6c8S7UEqNNbGMdRYlXoyymeus2mbD1zPw02g3d0IBtH4oXAP/KWAtV0CkIn2DyFnGSYXRy5WN1vlscr3/Cwz2aGTOpdyHNpB4MhR4afzr10+W/g84qnOoFHQL/BZzK2PhJJ/oJcRc6m3cusAO4lLF1a6aLMqNr9BgxMQHY16LzgJIsRJdp+AemT2x7gI/X2TZo1/jv0fGSqWxf0PGf/wa+wtjrfDO6h+ftTCB+MlPGhMw0FPAStBlYbzTm7WiLoVG9Tv1oc/YqdFGiJMuBL6N7OQ40aHszhXpW0DVon73ZFelMD945wBXA2prvVqNvvuI0bNtFX3un1vmujLaMbm3QtgS4CN0Z8F5GC4qHTle4F+1OmeXr0kw3p97TZiYE+xT6yfNFdOm7Wrahn1ZTMS9rEfT4iTPR+Qu1PAf99EjV+W46qM0IPVIkB+81G+P63IN+ate2aQ31M6CnioO2BP62zndl4Kvom9+0sREIWsC/X+e7LNrV+r1DraTZMZOZSAqdibqozne70N1w01XK8Ap0dm1tQNZDD+x6P81xBY+UyLvoJ+NMcnehvqvl1vlsKhhL6Lnoa6Deg8N06SaHajSKAB0/+XWd71oY/56oMBMtk2ai0CUDXouOVyQZRvfLb2H62hmhM2xvZazlU0CnPL9gmrY9HkdyrI6D3u+Z8JAzo4ePRV8TtW16ksamMQjaIv4a9V3rHegHynS4VcltvJ/6o7lPQMf2xu0gaXbM5EheqAfDtOFU4D+pX9f2f9Fp8NNNiWqG4pqa73rQT61XMX7a/lQZLxV7OqgXc3o5OgvYFEM29X0P9hs4eBsPdY0lb5Aw8VkLOpbVVec3T9DYDNg0+vo6puZzYxW+C/gt0/vAFeAGtGh8os73b0G7fZ9LLF/BiolG0F19X6C+kNyGDrgeqYS6x9A+7HmMHQd0LNp6eTfTE1uoN5J3OoY3mJhELcvRuTfpxHJJasfINBqz/ojxLaQD6DhKo46LQqcBvKPOd2V0zGI6LeIkgh7BvhCdFpHEi9vyCDpQPopmmpPZOtv3aE4GrAL+kPoDD/egR1s+eyQbBHwP/QSoffpl0WMnnjNN280xVuDTdT6bCsaFqJesZ7aXXDb5ql1Po1+Gg90bl6B7OBpFOzpBsZ4FtgV9cx/J8Wo+2kq6r853HcAb6/2omWIyRH3/rxlJa4K+iWqj8xF6qoefcOQtKEFbIJcxVmAjdDJRo9sUoHMYkgJWZPwErsNF0EJSb/T1TCZCd99/gMZZCQpdUa72HDvoGiR/RXOywvfE295X57vaim1Ac8XkbsY29Kdov7AZ3MXYcTJfpxonaUZwOEB3Fz9K9YIqo4/RTdPQphI6o3cX1Ys7QD8dG/1kNAHt8ayTmcYedCD2T2m8sG5DlxBIivgBdHrCZAbwNZpb0NefOUcRuo1XUydEMdFJuKaL16GTc1ai4xLvoHmJWQod2PwY2t36OTrLcDqj5xNlGdpK2YweSfxBpk90PXRE/y3oIO/XqRbcafTFkkcf41ehe6uSs/ll0A+bZOW1qVKv8FK9GEyy0tpT6N61b6FTAmp/3ygK6FjRm+PtnAX8bJq2NRkcdCb438VtOQdtrY+J1zVbTED3p2fQbo9Ru2Y2yswGWGpyO2px0W0rU73pprN9xuWb7hiWqSNiAr3mb3KmxmZiKuIfCcyxMNucruD3ZFHoe1TQ10Pd+3QmiInF0uwHiKUBzITkIIvFCslRQLPzTCwNQKnmeALWqrUksZaJxWJpCFZMLBZLQ5ismGSYfYlGeaZnrlyLxZJgMmLioosq/+40teVwUejYT73AgYuuGl9bwYyD/KaR7ZpJJSgtlmllomJiuu4WUk2lbXb/v6EX+CzVYdu17Xoe9S2T1wJnTGO7lqDH1tQbOGixHHVMVExM2H4XcH3NZ80mB7yGas9UbbtuRGcx1rIZWDF9zSIDvJj69WMtlqOOyXQNR+gZ7Uz2JcwMQUkBj1PfpYjQI37rDdU3maTThYOeSHomZDBOmUN0Azvoka+ghx8U0eelHZ3ZXFt533IUMtkArBk9OhNqtRqG0UJSm/ZtXDOf+jf0bhpbxzW5XdCDo4qJbc8Ut3C6eB26TMO16Cp1G4AH0AMH147/M8vRwmQsExNQNHOLGFJokTH5+gpd4OYU4Gm0YN2PHrRVO8CqVpAK8W8Xx9saQF+gT1KtvGW20xq3YxN6OoDl6OpXMHpmwBTVKQwy6DEnETqWMhSvR6haEhIfl7BO++phpmQw1k8ubvuKeD8WJto9zGgrSSX+LkPfdMPoSZ6eRJfPmymifTBM+yP0PpuJxwvomJGXWO5I7o9N0z+CTEZMHHQ5t7uBb8eftaPdiO+gh8QfR3VCqRvQ85YuQpeiM/Px1jN5s+hSfaehh9s/jK429rvx7wN0sPTxePkl6NKGHegbcA26FP/d6Jv4G+hh8w56tOMv0WUXj0eXZdwd/78QLTYL0CUIPosWotPQN8alEzguvegiv6Y62FlAd9yuXnT1tp/Fx+bb6NooUL3QN6FHYXroId/70dMovBj4UdzeejU5x0VEnHj7B4BdSqnpvqFS6GJO29FFg/aj9+fP0Bba06Zp09wO4u1uQM9Gl0Zfm08coW3PbURkoi9XRG4QkffE75WIdIjIAyLyXBF5jYjcKCJvEJFOEUmLSFZEuuPXjSJyl4ikEutUIuKIyNUicrOIbBaRrvg78/v5IvLnIrJdRNbVbLtDRF4uIo+LyMki0i4ivfE2zLpviNuGiBREpEdEWkTkoyKyJd5eV9xGFb++LCL/NcHjslxEdsS/N9toF5E/EJHfisgLRaQ1/ixds+9/KiLb4n2YH7fbidexSkT+VUQeEpGlB2tDzfl0RETCsCwiocR0iUhOtMg0+jpRE/zsSL1SIvL2eL99EXlfE9sy2dfzRd9f5lpsdnsm9ZqoZZI000uJ9yG6t+QitDvyWsYWPCqizf6Xop9QL2Z0j9CbgFXoXBAzZ6+i6qYU0XOF9KELyKyPt3sg/v5BtHvyULxMX2IdUHVdVNz+4fh9P/CDRHvNPpnpFkztikORibdttmfW/xt00Zv7qT837fFoS+g16CK9yTYMo+uVnI12d65GH5+JlAPwAIrlvaCG8aSddLZrr2meiOSVUo0sSGT2aT5wItq18dDWyiDa6gzQ18mOBm73YO0x15FLtbDPbHB5/gQ9X86n0PfUv1F/HqUZyUTFxNxkPYyeAiKFvrk3outC7qP+STMxlveiJ5N6fuK7AvqGS07+bX6fXNetaLdqCfoGM98V0OZssjcnGVvpi9uZXCfovJSemuVNWyfTCzMSt8vEQsx6WtGlFd3E52Y7GbQr9j7q1xJNtuUS4IXook3/MIH2+AD57AL6B39Ded9/E+Seh5d7AenCBlCpYRFZBOxWSjWiTkcGeBF6KstNVOeT2Y12UYvoB9AbaZyYHEwYTC0aUxvl8XGWm2mk0OcuHb/ejQ5qvwHtpte7J2pJxuBqr+lpZ7KWyTaqQUFTEcsF7kRP3jPejppK5Hegpw3IUC1RN0I1SJlcPikIoIXqPVQL5kji90bsqPO7Lqpxmtr2DSa2l/x+MgWBnHgbtcdS0DdPquYz0MWgDwBba9pT79hFaL//TPRT/lCV35T5k8ssQ/yvEgWXURr8XcKR03DzzyNdWLMdHERkGfDsFETFRVuVpuSgQvfmPIiOZ/0O+mGTi1+mfbX7mQGWoo9lP+OLTh4tUFl0ztOuOssI1cJCMHbS7xQ6SNyBjvEcQN/Ek7Va2tHX7V709TLIxGcLSF6jZrs+Oqb2DFp4T0HH3G4AXgH83yHW5wInAS9BP3AzcRvvRedZXcPU5yY+KJMJwAraKklmk0bop7spL3eo7s8ALUZpqmLyC/S0AX8NnEtVRZMl9BT6ZH2XsQejAx3wq70hTA/NENULK3kChxh//9PUnwhpPLYx+qlg2tiWWCZ5Ab0N7WIlrbHxEHRJS0EHdrcdbGGlVCQiLiCpVEs0MtRLLu/hyAOUD/wnbnQy/f3HI6kTae/+3afARUQ6gX6l1GSfYDm0KW6ehP+NDoSPoG/aD6CnRhCq7ka9Hr0XoV05B7gY/dCo15aXoS21FrQr8M+MvR7M8AqzLeOu9qBvytPRPXlZ9DVYRFt9F1G/ePJ4/D56wqwSWjyH0J0MVzDx+XSy6DmkF8ftvho9ReeF6Hlr/hp93K5BW//m3CevsTS6zOa70dnpxXi9plrdn1ItCv4NdPnPZ5gGa2UyGbBmB5IXhYt+ktye+OxQ6zBdiMTvt6HjKP+Odl9OQz+lMjW/Ta4jKVpD6CdErQiZ9uYZ/cQw33tUn5bJ7xz0gZ9oopUb71OtNVNGn2iTy5JslzH/TXwmjb4ATAnLbPxdPv6ugL5Aj51Ig5RSUdyD88nc6isohy8l7WQgeIjBfReRis7E738XA7svxS/uAGQ/EIoOdE8mUJtFn7thtG//RarXR0BVlE1Kgdl/gzkm7VSnBa2dhCpJctDmUupfbxHVY96HPm4r0VXlv4Z+4reiz3Me/TD6OPpGbuPQJBM281QznV+E7pW7Ad0bl1y2Hm9Fu+7fQddV/Rw6TnYO+nh9IP6sk+oEXeaaTt4Hr0EL+AaqD88fo++nT8br34k+D29Gx/fedoj2eWjr5qVoS6ebCWjFZPNMyoy2AMyNN1FFN9Wtay+Ce9FzrP4+Wkw+gw7o7kOX+9+CdqNM7kfyICh0jKN2btakMNRLThth7LgZFbcxZKx5XIs5sUZ4arOCjQ9ce8LM7/4ZfQG68baM9ZVDX+C/Roukj77IFzD5TNJPUXjxh+nNM7DzUpxO6Auhz8uSDw6Q2fcOUgegmIG92dNZOP+c8x3y5wOIiHsIS8WN2zoSt+076AeLwUGf6yDer/Em1SJebijez0cZ/6kZxetqZfyq9g76ug7Q5+BD6BjdRvR5/SbaZdgfb+8jaLF5IXr61X/l4DVfTfvvQc+1lEZfu29Fn6+T0Dfzixg766I597+Htsbz6Hvq6bjNC9FxkqvQIvEVdJH1tvjzjrjdScukI7H+Mjrw/1ZGpxO8Hd2JYe73R2r2Jdm+56CF6AXo82sSVa9Fp1mMG/earJtjXsnPJlr0VsXLHmBs6rug/boL45dJQjsWfXK2oBX13YyeItEkjNW2y2wvRB/UjprPBX2CapOpjFClOLQLkhSNEmPFzIhvKvHerL8HPWPflsT3JfTJMxaOOZEmhuNS7cGaUM+EUmqPiKxMt21+PD30Wdj+d6xsa4FiDrwCRAUYOkB2OMPi1gFCriIovBCvsApwwkMIiomD5dCCV3tzC/q4++ibPyms9QL0xkpMJifWMkzVMhlPTCReh0ma+zP0ef41+ub/LaPPxU/RlvVK9EwJlwK/GmfdJH77ONUZAn6IFtPz0KK0Gi1Kf8voe8NYM38c728J/eD8d3Re0MfR06O+Am2h3Y92e94Z/34NVS/AtOVOtPv/YvRD6TjgZrT79Ey8X29FH68M2i29uWZfzP8b0Q9v0yNn7oERdE/Ty9FiU9fVnkzegRGDer85VKwEqievg7GD8motjV3og/Z1dI/HC9BTT9yFDkqZ5SKq2aq1wS/jhmVq2pwM3Na2zfzfzcQrknvxNkxg1KzHQZuWtf5ziLa4IrQ47Ea7CP1oc3R//NqOPg674++SGcSTCaI9jfIgeC7cthC+1wc3DjJ43TNw4xDcloefRvCDn+Fe9lcUL/kd9t15PmF5GPQTdzwEbSXsQz+dX4/2/Q154JVU3Y6DCWBI9fh1U9+ag6rrGzH++XEY3XsH+oY7mdFCYvahH52MORB/dso46zXUOweCnv3uzVTHib2O6nil5P50om9MQd/sZ6OPz2/Q8SJBXzMnxf9vpfrA3lSzLrPdP0Ef/++jhWAdugfoTHQaxsvRx+5/0LGm2t4e4jZ/kursDB9HC+Mp6PuwjD43H2Qc3ZhsElO9k1wbwzjYb81Tw7gdncCrqVoXZhmn5ndPo7NnL0MHm5KUqB//AH0CTL5DLR76wBuSgrKT6oVwKPJoM7/2qWuesNmaz0ErfvJGrSeqyfcO+jjVWj8TQQE8eNU32PaTbQz/EvZfO0jppiJ9W/fT/6NdDP98P/03HuDJO4UdpReR7n4+bjoP9aeHTLZ5GB1ANwP9Po32tVein4Cb0cfYQd/M4wnKYqpP8HsZa+ma3yxCn2+H8SduN93CZh2/Rfv9tT13hhAdL8nH353C4ddG3oO2UiJ0D5+JwSSvrX70k91Yc8nrI4i37aKvkRVotw/0uU9T3zsQdAb6O9CxljL6mP5nvI5d6IfyWTW/SZJGWzVtaPH9JDqw/Jv4//vjNv8/RqeHVJism1PP5A2YWJdY0oUwYpFBB+3+ED0ozJCcFyYpMleiI/210ewRxndLTFSbmt956G692jYSr6/uAavDqVS7p5MYcavtyhW0SfwvVH375L7Wa8+p6MmqfjbBNukfi3iAP7LzMVrLu8njkC+3kPcVYXmIMAxId8KB9HzKv/c+lrzur3E75oF2wc5VSh2qGzpCP+lejhbUP0ALUAF94RWpPmz+Ce2q/l/8uXmALELfBEa8W9HWq3kam/O/CO37m/lbcugnpU81tyVE38QnxJ+FwOXxejqoXsNpquO18mjTXcW/WUz1vNSi4naawlo+1bFfxoUxLlaEDmCaCb6N9dyKflhtRrsV/4wOtK5FWylptEgHaCFYR/XaWIbOXTLHxNyPXvxahg5MG8vNi9v4LNql7oiPn3GbzfGI0CLRFn93J9V5rEAL9B1oD2E5OgxRO9Zu0gpcTzj8Op8dDJNMBPqgXoU2nd5V2zjGuj9tVN0GcwP66MBV7TzB5iIOqXbzJtefYWzRJLPO+9GBuX9FH9zaJ6p5vwhtSu5hbN2SAfRNVUA/jZLr+DH6wnkLukvyUG7LK9GR/8nOLtgKwrPXnEf6lu/TKcBQP6Q6Yf4m9ndmKC/fwII3foSORetAm8o/mMRYnhDtjv0h2hTejL4R+tDn9RfomEBvvOy30E/lEH0eu9HXw6J4fRF6DNbL0WLRinb5TK9QR7ycQl8z70GfH4eqeARUxQh0ouRrE+vbE7dne7zOfegnslk+z/hB8854v1ZQvReMCJmg/bz4Ny3oB+W/xL/ri7ffj77pzfb+AZ1X0klVUFNUp05ti9fvxvt8RrwvIdWH5Eh8PBdTvQ6HEssch7b4clRjliPxcRqIt9sfvzfBcrNPRsiWxOstM96sm5PIvXdF5FoR+YvEZ52ix9S8egK/V6LHsTwpemyMGQdzoojsEz1+xjnIb5eKHqfy3MRniB6f88u4DfXGM1wtIq+v8/k7ReTBeL+oWWeriDwqIm8eZ52IHgN0nYh8QkR+LiILa77vEj3u5hXj/P5lIrJbRH7/EMfsH0Xk+viY1V1unPOqRET8+26VZ995ksifZURei0SvQA584FTZduUXxO/bIzF/ITo3ZVwmcG57RI+PepXocUauiGRE5PR4PysDhepQFpFIRALR42nGw3wXHGQZQ1TzfiK/ERG5RUTyUv+8LxSRoQmsoziBZXypHpPksUkeg6DOZxPZ9k4Rea+I/L2IDCfWc7BzIFI9ZoOix9iZe/SMxDI/EX1eD3tsDlQVOM1oszzDxHx5c9WbPmtjqt2NHgF8PnAB1VGeYbzuAjqK/Vl0t95dNevrQ/vn56EDaTei81WG4+0ku1ST1sG1wN/E6zwX3eX1ePzdIPAX6IDYAvSI2L64TW1os/gstPXyBNVRxkn2o586X0P7rr9Ej4Y27tjP0ElJX0O7ClvjdnvoJ9wadD7AGrR5Xy/uUxcRqZjA3jfPpmvbI5BKMdTSSekNf07LK06nfcEqs3haKTXVui6CjmHUxjFC9LF7FG29LKWaMm7KNmTRZniGqsuSzE3ZTzXPpoC2Zj30U95kUrdQtdrM9JpDVC1T0NdoG/o8L40/T3bpl9FWyS2MnxFbBj6PNvXTcVuy6Cd+hA6oLoh/b/bPuD3tcVtNMNq4R6YL27RVUXVRJN6PLqqJZ2ZdEn9mev364tf96GvJjOm5g6oLsyhuc4h2VczYqWG0pbIcPVK/DZ3g9tJ4P98Ur2t//HndDPHJTA/qoLPpdqCH8yv0wf8TdE/LExNYRzvad/53Rg8YVGhX5e/RJu59aFPqOKo30UfQwgNj3Q6AP0J3jxXQ2aVb0Bfk69En+ZbEb4l/vw7tahyD7i77EqN7HjbE212PNov7qBb9+Ti6v74H3fX2FUbngZjtvAYd5e9H31jX1bR7Sbzfp6CP4R50ALMFHSO6kEPkl9SbhEtE3rL/2W2Xdn5gKfvownnz2bS85E/wOiuTC2QAf6IuzSSuk7pNTPyt7Y0wn+eodoObWEmKar0c426a7M6JxuoO1qZkYDQZm6sXvzI4Nb9pNrW9U8m/yWXq9eAkOzxM+OEf0K57iL6/HfQ9sR+dlfuFOuvXK5vkRXKw7r2J/h5Gn7haMmjrxfRzm4Smiab/JgNLE8WY+OYiTuKglbqDaldiX836D5VrkzyZ490ABfQTKIv2Y3cfZNnRK68RExHJAiPbH3yA/d96Nxs/cAVO1zLQwnYdEEy2xskUxeRQHIuOB+TRovzz+PN3oQe6PYxOmOoBPoq+Pj5G1Uq1TJ3kRfR8dNDf9HANo8/J7Rzk/j/cicunKiqWBlJHTF45sHfPj/MtLbiZSiJvbgK9M5OiQQKjgFehzeccumvzwvi7C9EmdoS24JajE7Ra0IHILdjrcDo4mNU2/o+m+YljaR6vRLt3yWEIM5XFaBfVRQ9qNLkVz0NbLTvQMa4ser8yVGNjluljUu6cFRPLTKbeE9IyQzncTD+L5UgwXkDRMgM57JqgFovFksSKicViaQj/H0ecN7mC/VekAAAAAElFTkSuQmCC" class="header-logo"/></div></div>""", unsafe_allow_html=True)


with st.sidebar:
    st.markdown("""<div style="text-align:center;padding:12px 0;"><div style="display:inline-block;background:linear-gradient(135deg,#004d40,#00695c);border-radius:14px;padding:12px 16px;box-shadow:0 2px 8px rgba(0,0,0,0.2);"><img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAARMAAAEqCAYAAAA23LTdAAAACXBIWXMAAA7EAAAOxAGVKw4bAAAF8WlUWHRYTUw6Y29tLmFkb2JlLnhtcAAAAAAAPD94cGFja2V0IGJlZ2luPSLvu78iIGlkPSJXNU0wTXBDZWhpSHpyZVN6TlRjemtjOWQiPz4gPHg6eG1wbWV0YSB4bWxuczp4PSJhZG9iZTpuczptZXRhLyIgeDp4bXB0az0iQWRvYmUgWE1QIENvcmUgNi4wLWMwMDIgNzkuMTY0NDYwLCAyMDIwLzA1LzEyLTE2OjA0OjE3ICAgICAgICAiPiA8cmRmOlJERiB4bWxuczpyZGY9Imh0dHA6Ly93d3cudzMub3JnLzE5OTkvMDIvMjItcmRmLXN5bnRheC1ucyMiPiA8cmRmOkRlc2NyaXB0aW9uIHJkZjphYm91dD0iIiB4bWxuczp4bXA9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC8iIHhtbG5zOmRjPSJodHRwOi8vcHVybC5vcmcvZGMvZWxlbWVudHMvMS4xLyIgeG1sbnM6cGhvdG9zaG9wPSJodHRwOi8vbnMuYWRvYmUuY29tL3Bob3Rvc2hvcC8xLjAvIiB4bWxuczp4bXBNTT0iaHR0cDovL25zLmFkb2JlLmNvbS94YXAvMS4wL21tLyIgeG1sbnM6c3RFdnQ9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC9zVHlwZS9SZXNvdXJjZUV2ZW50IyIgeG1wOkNyZWF0b3JUb29sPSJBZG9iZSBQaG90b3Nob3AgMjEuMiAoV2luZG93cykiIHhtcDpDcmVhdGVEYXRlPSIyMDIxLTAyLTA4VDExOjI2OjE2KzAyOjAwIiB4bXA6TW9kaWZ5RGF0ZT0iMjAyMS0wMi0wOVQwMDozOTo0NyswMjowMCIgeG1wOk1ldGFkYXRhRGF0ZT0iMjAyMS0wMi0wOVQwMDozOTo0NyswMjowMCIgZGM6Zm9ybWF0PSJpbWFnZS9wbmciIHBob3Rvc2hvcDpDb2xvck1vZGU9IjMiIHBob3Rvc2hvcDpJQ0NQcm9maWxlPSJzUkdCIElFQzYxOTY2LTIuMSIgeG1wTU06SW5zdGFuY2VJRD0ieG1wLmlpZDozMWVjODIwMS1jODdlLTJmNGQtYTliYi0wNWJhMzAxZjlhZmYiIHhtcE1NOkRvY3VtZW50SUQ9ImFkb2JlOmRvY2lkOnBob3Rvc2hvcDpmYTE5YzI5Zi0wMDQyLTMwNGEtOGQ4ZC01MjkzZmI1ZmQzOWIiIHhtcE1NOk9yaWdpbmFsRG9jdW1lbnRJRD0ieG1wLmRpZDowNjBiNDU0Yi0yMDQ3LTIyNDgtYWMxZC1mMjE5MjcwYmVkNDkiPiA8eG1wTU06SGlzdG9yeT4gPHJkZjpTZXE+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJjcmVhdGVkIiBzdEV2dDppbnN0YW5jZUlEPSJ4bXAuaWlkOjA2MGI0NTRiLTIwNDctMjI0OC1hYzFkLWYyMTkyNzBiZWQ0OSIgc3RFdnQ6d2hlbj0iMjAyMS0wMi0wOFQxMToyNjoxNiswMjowMCIgc3RFdnQ6c29mdHdhcmVBZ2VudD0iQWRvYmUgUGhvdG9zaG9wIDIxLjIgKFdpbmRvd3MpIi8+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJzYXZlZCIgc3RFdnQ6aW5zdGFuY2VJRD0ieG1wLmlpZDozMWVjODIwMS1jODdlLTJmNGQtYTliYi0wNWJhMzAxZjlhZmYiIHN0RXZ0OndoZW49IjIwMjEtMDItMDlUMDA6Mzk6NDcrMDI6MDAiIHN0RXZ0OnNvZnR3YXJlQWdlbnQ9IkFkb2JlIFBob3Rvc2hvcCAyMS4yIChXaW5kb3dzKSIgc3RFdnQ6Y2hhbmdlZD0iLyIvPiA8L3JkZjpTZXE+IDwveG1wTU06SGlzdG9yeT4gPC9yZGY6RGVzY3JpcHRpb24+IDwvcmRmOlJERj4gPC94OnhtcG1ldGE+IDw/eHBhY2tldCBlbmQ9InIiPz725RFTAABaG0lEQVR4nO2dd5gkV3W331tVnSfvzO5szlFhhQQIZLAwGPMRjMEYgzBBAmxjAzbYGNsSDtjCmGiSkBASQQEUEEEgWGlBMkYJ5QQoorRabd6d3N0VzvfHrdtd09OzO7PTsz2zc9/n6Wemu6urbqVfnXPuuecqEcFisVimitPsBlgslqMDKyaWWlqA3wdWNbshltmFN9UVKKUa0Q7LDEBEWkAG/GIfXqoF5XqrlVK/bXa7LNNHI8Mc1jKxACAiDjAQlvbyxM2fYO+vbwZ4V5ObZZlFWDGxGDpA2PfQefTf/Tn23XgZ+P6ZInJcsxtmmR1YMbEgIjlgL8XfEv32HLL7irj3/Jjgf68FuE9ErC9rOSRWTCwAzyMaQJ45iw4O0FmG/K5tRF/9EFEYAKxocvssswArJnMcEcmD/LzU/y2UXEsmDOj0PVpLIengMfp3PwtwcrPbaZn5WDGxnCLiU+z/OkQHQClkQNHiODAS0PfsEwSlkW+LyEua3VDLzMaKyRxGRBZBuJV9X6Fl6B5w00Qpn6DXZyQdQeCy/EO/S+narwHcICJuk5tsmcFYMZnbnCDBfkYOXIEQQJRCIgcRQADS0AbPXP1Vyv17AV7W1NZaZjRWTOYoItIBcs1I/3X4/p2IEwJZkFRiqQzFMswb2c7wfTeAyLUi0t6kJltmOFZM5iCxu7I/CnczvOd/cJwR3EoutIcCIgFEUSy00+L3Mfz5dzD82/sA3t6kZltmOFZM5iZrkDLF/ZeTlgdJu6BHRYwAAUo5IA44IUMiqGKZBcPDyK1XI1H4BRFZ2OT2W2YgVkzmGCKSAh6Myr/B3/UfpKWMF2WIQiAaBqeEIxEiEThl8mmHdAbcMvC9zzHyxEMAr2/qTlhmJFZM5h4rkTL+vvPIO/tJhw4qbEUkDQ44HuCAKMARvFIfhEAGckP7Gdn6dSK/dI6IHNPk/bDMMKyYzCHiWMlDkf8k/p6rSDkKJwJHPJTkQWVRLkRuHDNRLnkRSmXAyeAoIXvdFxm492cAD8SDAy0WwIrJXOMYpMjQngvJOSGMOKBAomGiyIEoRxRRtVCiDG4A6WyBkBQ4UCiVSN9+FcFwP8Apzd0dy0zCiskcQUQKIPeODF8PwxchfhnIASHKGQRVhigDEVRKXEQOvq//9cUHBYQwsuVrDD90O8Av7CBAi8GKydzhOImG6N/9GTxnL64KtPkhIUIEDIESIkkjERABEqEyWRTgUIIU0AJdaYf9F/0bft9egMgKigWsmMwJYqvkFn/4l2S4C8IyKhNpa0RpVwdHwBlCSEEIEgLKp0yWUmmItAsIFAcB5dK7/T52X32u2USuWftmmTlYMTnKia2GQQn3MrL7HAqUcCKInDJIEXBBwFGAO4w4QiRZVARIgHJc/AAdS/EcigLRiE8mHCF3+w8o7XgSbFexBSsmc4EOEEr7f0i6dB2OP4LnQhBB6DmgvHgcDkAEygcnipPYHFx88rkUUdxd3NGZx3E9/IEA98lfceDHF0AYXCIiK5u2h5YZgRWTo5/jJOzHP/Adcu4QTgQIRMoDN02oFKgUCkC54PrglLXrE2UIRvoh5SKpNCOlCH9kmCiKSOXStKVdMjdcSP+dWwF+a7uK5zb25B/FiMgHIfr5YP89eDwGaI0IfIAMQRTHTSU2O8QDpTUFBUiafBZGimVK4uKlPVQIIJDKMTI4hDuwm/KdP0aCMsCaI7+XlpmCmmqpezvVxcxFRB5m92lr6bsBPxgiUoKL4IYhSiJwhYCAyNPdwZmoAx7zKf5yiNK90F726M+nyAcBXimKM2FD+gpQ9qCzBN4AkG9j4EOX0HrKHwI4Sik7TeQswU51YTkkIpIG1tJ/H2F5hDAMEQlBRfoCiiIIw0Oux5VI/+MAngAOmQByPoQKwjSQchi+6SeEw4NgSzzOWayYHL2sjaISFJ8higKUAsfRlqRSisoHh8CLdNJJmBJd6kRBtgS5MgRAMQ9DwwNkbr6W8sP3AtxiYydzE3vSj0JEZCkSPFAauAEcwXUdrRsqqlgnuAkxqXiq0Zh1uYSIEnw3ouzG34cObuCAQOAAUUiHGsb//lco7dsB8Lrp3kfLzMOKyVFGnFfyVFh6gvKud0DkI/igAkAQU14giiAKJrDGiNCJiJKhMaUAFy9I4Q1Doa2FoH8H0V1XEzx0G8BVIpKdht2zzGCsmBx9dENEce/3KKR2EUY+IiHasxEcV/+FKB4afHCCuGfHjSAdOoAi9BwC1yETOhTCNOHeQcppKBSK9J33YYLtjwO8aXp30zLTsGJyFBEXPtoVlbcR7v88XgiOIyjj4hACke76dZjQ2Q8cUAKZ0EFFitCBYioicOLendDBLTik28Evl+jZ/xTBz64CkW/oOrOWuYIVk6OLY5ARhnd/lqy3A8qgUg5KCVEUEoZCEERINDY2Mh5RbJUQKIg8fAdCR3DxQQIQn6AcEQSQEUipkKHvfInhx38D8Nbp2lHLzMOKyVGCiHjA3cP7fwojnyedCiFohdBHCCudN5UOnDi/wHF0+KTyeRhSSR1SSvcIR+jENuUQOg6uRPozJ4R0hGQgq8AtAVGA4xUp3rEVyiNfFJHfO4KHwdJErJgcPWySqIj0fZG0MrHVjuq3B8ktrOYtOZX35jMvTr9HOVpPBNKh/osHpbRQVkpnsYUQeQ6BGuHA9z5HsOUigOttiYK5gRWTowARyQH3lvb9nEx4EyqOrYYcOsBqRENPvCVjvswEaCFytcuTDmOBccB3IXSAUCBSRMolctO051P0+H143/4YoV8GWN/A3bXMUKyYHB38buTvQw78B54M4/rgOS34auDgv5Kav3Eym1K6zAlEqEhnuvpuAER4YYQS8GMxSQMpgcjzKKZciiUfd3CYQjAEpZ0ceOQegD+elr22zCismMxydOEjtpT2XUPWu0vnnUUpBBfxEmIyxkjRp97ER5RyKmIyKjFWIHLAd0Mc/LgCm0Oo4nX6+m/oCiqXJo2LWyzj+GXKUZnwlu8R7NvxMRGxNU+OcqyYzH5eHpV3Eg1+iSgoEoWgHI8gLOF64/Ta1Ml2rwz4MpZJJQir3RtREDqRDpZEOmHNiXSVtggI/CJuGJJ2XfA8yIJKQ27rBQzeeBXAd+Oua8tRihWTWYyI9EL0veH9V+CphwmDOFbqFkEVJ7CGQ5/+MCEqgQu4QOTh+SncMEWYdSEHqRBSQyNQKlGOfEYiKPvQemAPg1d8ksFnHgN42+Huq2XmY8VkdvOGyN9DOHwRod+Pq9IoFJEjeGmI/EOvwPTcSKILp/KvCJFycCQuTO/E4iIKAg8lLv3lkMgDzwUlDrgeUS6FZCClE2bpjIoMPvEgwIUisnjajoalqdh6JrMUEVmGhE/u3f52MuG3yPlxnkfQEncDO0CKA7kRMl5IUCqRVaCiHEo6GPIVTk4RZfsplgZo9yBTbIVtbYQPtvPkvQEdmSV0PXU95BVDRUGcNDkvi1v2IShBMdIRWAeKafDTHpFSpEPIhIJTCtjjddFdHoFUjqF/u4rCiS/5mFLqI009eJYKjaxn4jVsTZYjRjwz35NB8TG8oevJOOAGgOTiKvMlQjfQxdMA10mRUlm8MAVOgBSfpS0LZFyGilkysgpfXozb+UroXIr3e89hVRCy947/Zc+lz1De/ihdmRRZAb+vHxHwCh5kMuA6lEaGUGVoiQLCEMoRlFIKJ58l6w6ACghDBU8+gGx+0VkicoFS6ommHkRLw7GWySxERE4hGrhp6PH3kC1/D1d8EIU4EDo+oVONsboeyAikfT2FBU7s+7RBeQBG1PHkl34BJ3scbqbLbOIEtM1xG8N99N9+LSPXfY3Cb/6PQlhCZT3KQchIpAjFIaMcCgoIyzqdNpMmUNBXLJPNQ6oMKvLY7fTQ+i/fotVaJzOGRlomVkxmGfHQ/pHy3stI73wPhAOAA15AkNID8yIBJ/JwIxenWMJNoQOncTcubhtBtJyg673QeirZ1vXEvlEBKCulgnhbHnoK0J9TGmbgtusY2PpN8g9fT8ofIKUER3S2red4OCpNUCoTRgGplIOT8iBVpjgC2VQLhFn8416M988XoXItpyilbmnGMbRUsWIyhxGR5xEN3bb3V69mXuY2KJZ0hlkqoJTSAVI38PCCjC65KBF4vp7bwvEoqeNg/sdwCs8lle0xq20DhpRSdfuS4xKQxwJ3SmmYvluvZviK/6Rn54OkoggCLWC+SiGOS9Z1tZVS9KEVigGUBj3axUMyGQ584Et0vvytYOvFNh0rJnOUOG1+eGjfpUQ7/5xWZ0SPiVEuoRvhewKSIhU4uDKkfxN1Eqos5fQx0HEGbsvvkMkvB3gtsBUoTfSGjmM1xwN3URxgcOtFjFz+Cdr7d+gZ/4iIghARcF0ARRQJYUsHSrJ4fQMgEUOnvJr0351PqqXzdUqpHzT+SFkmihWTOYqI/JmE2y7Z8dgJ9Ob3ovrRwREHfMchJAtRjlTk43IAFAz2/Bchx5JvewGpTA/AO4ArlVIjU2hHFng+8HPKIxRvvY6+r/4T3f2P46ZCKAUMR+BkIVvupE+VcLIemZJP2o8YaWtn52v+lhVvPRPAHc8iskw/VkzmICIyD2TPnqc/RC74LAVABhXKE8TTCWW+gARdqHA+XvqFON4iosUfIp3qMKvJKqVKDWxTGtgA3Eu5SOmXP6bvik+Te+YespmIoFQiV5wH6SJ7gyFy7RmUD+VSxMiyzbR9+Fvkl659j1LqK41qk2VyWDGZg4jIa8Xf9YPh3/4+Ofd+ZAjcdI5IjRCmoCSQyTsMDm8m2/Up0i3Pw023gbZELgOC6bIA4mr0xwD3URph8OYf0n/5x2k/8ASFfsDth1TEkAe+DzlXEeTa6X/NWSx804cAWpRSQ9PRNsvBsWIyxxCR5UjwhL/3uzi7/xrX3Uvo64zUwMvgploYKS9CFd5Kpv21ZFvWA+olwK2NtEQm0E4XOA64m+IQw7+6lYGrP036gf+lMyrqivhlIfTBz8GejafS9hdfpm35pvcrpb50pNppqWLFZA4RP/VDSo8y/NgfkePXiAt+CspuB1H6fWRaTiOTX4VyswDvB77ezCd9LCrPAW6X8iDFh29j+Nufpf3u6/EcXeoRF/blXPzTPsb8P/w7lJtap5R6pFltnqs0UkwQkSm9LNOLiGyUqCRDT7xVBu9Hosc8KT22QPY88ceyb88V4vt9EnN8HMOYMYiIKyLPFRGR0rCEN/9Q+v/uJVL8oxaRP8pK8Gpk35tbpP+2n4iIfLbZ7Z2LTPX+H6UFVkxmLiKSExEp7vtfOfDwPCk/XZC+R18jA7uvl8AfNCJyqszwGfREJCsim0VEpDggw3dfK/s/8XaRP2wXeXVBhs5+u4QjgyIiL2x2W+caVkzmCCLymijYJwe2nSv7n/5L2f/s58UvPmtE5AWiM1RnDbE4Pr9iqdz4Q+l76wul/9/fJfvuu1VE5KPNbuNcw4rJHEFEPlws/kYG+h6QMCxKghnlzkyW2FJ5jhaVEQkP7JHBHU9LFIYitkTBEaWRYmIDsDMU0RXdI2EERQ7gROBBoHi0pKBLMk9Fs0Qp9UwTmzTnaKRBYMVkBhMLynOBXymlhpvdnuki3k9HKRU2uy1zDSsmFoulITRSTGZ0L4DFYpk9WDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyYWi6UhWDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyYWi6UhWDGxWCwNwYqJxWJpCFZMLBZLQ7BiYrFYGoIVE4vF0hCsmFgsloZgxcRisTQEKyaWOYGIuCKyWUQ2NLstRytWTCxHNSLSISIvAYJisXiP75d/IyLvFpFUs9t2tGGnB7UcdcRCsQH4HRE5d2BggN/85jd885sXUihk+chZH6W9oxMgp5QqNre1zaWR04N6DVuTxdJERMQBFgAnA98rlUo89dRTXHrJZfzwRz+ir283mbTL0OBeOtpaOPMjH0MpNSIic15QGoW1TCyzGhFpBVYA90VRxJ49e7jxxhv57ne/y6233krKy1Mul3DdgCAcobWQxi8P87nPn88f/L/XARSUUsPN3Idm0kjLxIqJZdYhIi6wHNggIteUy2UeeughLr74Yq677jqGh4dRShEEAX7ZIZdLUy4P4KoA5QRk0g5LFy/k0suvZV73ApjDgmLFxDInEZEO4ATghiAI6Ovr48c//jFbtmzhpptuwnEc0uk0vu8TRRGpVArHd8EtE6gA33fxUu2U/SL5wiBv+tOX869nnYvnZWCOCoqNmVjmDCKSBpaig6nfHBgY4JFHHuHiiy9my5YtAJTLZQBc1yWKIkQEESEMQ5AIoghRDkoplCM4TsTQwDA/+tGPeMHzT+VVr3wHQBaYc2LSSKxlYplxxMHUhcBaYivkmWee4brrrmPLli38+te/JggCPM9jcHAQz/NwXRcRIYoiQF+XSilcCYjwiMQliCKclIvnRJRLg6RSIae++GS+fO5lZDJ5lFLzlFL7mrnvRxrr5liOSuJg6jrgjjAMGRoa4u677+aiiy7ipptuolQq4bouQRAQhiGZTAbH0RZHGIYEQQBoC0UpRRRFeKpIJDlQWUIJiGQEz3NwVRpPXHy/jzPP+gCnv/tDKOUAOEqpxt1hMxwrJpajBhHx0G7McSLyg2KxyPbt2/nud7/L97//ffbv38/Q0BDZbBbHcfB9H4AoiirXnrmGlVI4jlP5PooiHKcIksFx8jiegx8MxCJTQAUpstmIQn6IL5z7ZZ538itRSi1QSu1qxrFoBjZmYpn1iMh8YD3wf0EQsH//fm666SbOP/98nn76aUZGRgAdD3EcBxGhWCwiIuRyOaIoIgxDHRcBHMfBcZyKiCil8DwPRYYgBD8okvZSuG4KCUMQvWypGFAcGeS7V13K5ue8mHS6sFNE5pSgNAorJpYjhojkgDXASVEUfX1oaIjHHnuMc889lwceeIC9e/dSKpUq4qGUIpPJVH5vLJFSqUQYhqTTaTzPM+smDMOKkJiXSArXjQjCEYJyhON4KFJ4noMgIIpCfh7fu+qHPP/kU3j9G/4KYB5gxWSSWDfHMq3EwdSlaCvk2lKpxLPPPstVV13Fddddx+7du9m1axe5XI4wDFFKkUqlGBkZIZPJEARBLApaLFKpFI7jUCqV8DyvIjqmB8d13Uqvju/7SOSSzTkEYQnHcZEwQ+iHeCkFUQCRwlEejhph3aYu/ueLF7Jy5Qko5a5QSj3Z5MM37diYiWXGIyIFYDVwbxiG9Pf3c9ttt3Heeefx1FNP0dfXVwmSTms78EAFKAIQDyUeEIEq4yAoyRKFKaKwSK51hLe/4w186MP/g+N4MAfG7lgxscxI4mDqCuC5wLdLpRJPPvkkn/nMZ7j55pspl8sVC6NYLJLJZKb9+pF4YLyS+C8BKB+UFhckQ8prwZGIUrCXdGaIz3zmHF756tMAtUwp9fS0NrDJ2ACsZUYhIguA44Ctvu+za9cubrjhBr7xjW+wZ88e+vr6yGQylEoloiiipaUFz/MIw7ChF3M9IqeME2VBUqBii0SVibNREFyGRgaBiHyugOdm2bplK6f+3mvI5VqeEpGjXlAahRUTy2EhIllgFXBsFEWX9/f388ADD3DxxRdzyy23MDQ0BEAYhpXuWvO3WCwSBAGp1BEoKaLC+B8XJWhBQQBXWyZAJpcCFRCEwsC+Mtdccz3Hbb6At53+13he9ikRySqlStPf2NmNdXMsEyYeYNcLbAauKZfL7Nmzh+9973tceeWVbNu2rZKZajJSTbDUjJVJpVL4vk8QBBQKBYrF6Q1JhG4JJ8yhogyKCEcNgfKJSCGSJZQIvACIIPDIplvwCHCdfZx3wVc4+YWvxnHcTUqp30xrQ5uEdXMsRxQRaQFWAveFYciBAwe45557uPDCC3nooYfYuXMnuVyOVCpFEASUy+VK12wYhhQKBUBbJL7vV8RmeHi4Yq1MX+NTRCrCdUraIMGDKA0q3q4XoFxBBPwoxIt8JAwYCX2uvOwyjjv+hRRaun8tIouUUs9Ob2NnN9YysdQlDqb2AhtF5LpyucyTTz7JJZdcUunS9X2fTCZDuVzG87w449Qhm80ShiG+71cExXTrJgfjmeWnk1B5oIq4+KjIw4nyEGWIlBC5I6h0GZ+AMFCknQIqiEi5gishrgp5z3vfxF+9/+OmdyejlCpPa4OPMNYysUwb8TD/VcCdYRgyMDDAT3/6Uy666CK2bdvGvn37EBEymUxlnIzrunieVxkzUywWKzkjrutWxs84jkMQBBWXx3XdSgbr9O1QCtQIqBIodJwkyoHjgxoipAQoosghnU1zYGg3rZkMCkUUeFz9g0s58aTf4fkvfDWumzoWuGt6Gzx7sZaJxQzzXwCcLCJXFotFnnjiCc477zxuu+02du/ePe29LjOWIM+r/+gYPvofZ9PRthLluOuUUo80u1mNwuaZWBqCiHShM1NvDsOQ3bt38+1vf5urrrqKffv2VSyNuYwSj5bWQd77vtM54/SP4KVyL1dK/bTZ7WoUVkwsh008PmYJcKqIfHV4eJi77rqLCy64gIceeoj9+/dX3BHf93Fdt9lNbirZLOw/sIuTTtjIF79wHkuXH4vrpZ+rlLqz2W1rBFZMLJMinvphHrAJ+Jnv+zz11FNccMEF3H333Wzbtq3SRWuCpa7rkslkKkP+5yrKOYCr2lCh8La3/RF//48fJZNtQynnqOjdsQFYy4SIx8csB34VRREDAwNs3bqVr3/962zfvp2hoSGGh4fJZrOVcTLJgXKm2NBcJoyKRFELaWnhh1f/hOc8dx1/8Ioz8FL5DcCsF5NGYi2To4w4mLoYeB5weblc5oEHHuCrX/0q99xzDzt37sTzvEqhoXQ6XbFGHMchlUoRhiGlUgkRqQzxn6ukc4OUBvNkvU5K5Z087+QFfOGcS5g3by2umzpeKXV/s9s4FaybYxlFPMy/C51YdlsQBOzatYvvfOc7XHvttTzzzDOVwkJJa8PERGrrpoJOfT8iXbcznCA4QNrrJJ/u5MD+XbR0DHPGu/6QD3zwS3heHiCvlBppdjsPFysmFgBEJA/0AE+ICCMjI9x22218+ctf5tlnn2Xbtm2jksSSYmH+mvhIskqZsUbK5fKcD8A6DkRhSOA75NMFlDNEvq2fL593Pied+BpcN32CUureZrfzcLFiMoeJx8csAI4BrgvDkO3bt3PllVeydetWnn76aYaGhhARUqnUtNcLOdpRpIEijvKIwhQpx0XcXZz60vV86Ytb8LzcrK4bawOwc5B4fMwK4P4oiujr6+P666/nwgsv5JFHHqFYLFaqkAGV6mOWqeIQiQICgkCB6xGUPO69+zGuuPJznPbmD6OU14Et82gtk5lMPMy/HXiZiFxqgqnnnnsuDz74IM888wzpdLpSMzVZYNmIiT0/U0RSoEJcVxEGCs/LEAU+jnuAVWtauPzy/6NQ6MF1vfVKqYeb3dzJYt2co5g4mNqOHh9zRxRF7N69my9/+cts3bqV/fv3V+qiJiu3m3hHKpUinU4DzPkckUYgIroAdcqhXC4ionAkC9EIuZY+3vKWV3LmRy5CKRdm4RSjVkyOQmI3ZiHwsJnWYevWrZx33nkcOHCAvXv3VoKhw8P6ek2n0wRBUOl5MbkiyddcD6BOFeWERKGH47pEUiSKwFN5PHFxvH247k6+esEVPP+Fr8Nx3I1KqQeb3ebJYMXkKCERTN0kIluDIODxxx/nU5/6FHfccQelki7uFQQBxWKxYnnEv62IiOmtMZ+b70xpRMvho5yQwI/AUbhefNwlC6ED0SDtbT7Hn7CSL5xzNflCJ66bmlWCYsVkFiMiCmhDT/9wfxRF9Pf385Of/IRLL72UJ598knK5XLE4isUirutWXBfTO2OmdUjOF5O0UJJp8ZYpoAKiKATHJPA5EIEjKQiErvYMytnJBz/8ft7wp/+E46QBPKXUrFBx25szC4kzU3uBJ0WEcrlcqVZ2++23MzAwMGr6SyPSuVxOz/8SWxymTggkp8CsJpiZTNZkTonl8HGdFC2tGdo62li5ah3zexbS2d7GymWr6GzvZumSxXgpn+6eVqLIR6kUSqk5OYmXtUymkTiY2gqcAPxvGIbs3LmTK6+8ku9///vs2LGDcvnghbsUU8sTmepzR6Q6VscIlQn61iN5PTiOg6hq/CY5F7CxpkyFNvNbY0lVqrM51fmEk3MKm0Q7k9Fr4kPmszAM8TwPL572Jrl9UxXfJOd1dHTQ1dVNz/z59PTMp3fhQpYvX8m8eT0sWrqM9rZ2CoUW0uk0rutVjgHwFeApIAUcAG4DHptNOSfWzZnhxFbIQuLMVDPA7jvf+Q4PPvggAwMDBEEwobEvaornWg7z9FStHyqWTvJmNj1I9c5/8sYXFSSm6pRRmbgm6zZpaZkYj+d5+tiE0ajtmiCzKUqdy+Uqv03WnXUcR88SGPTT1dXFokVL6Oqax/wFvfT0zGfp0mUsWrSEnp755PIF8rk8mUwGL5XCcdzkfr0fbWXsRAvGADAI9B0NFeutmMxAYitkHrpy+9YwDHn88cc599xzufvuu3niiScqN1Ly5pxJiWX1zqV5iid7i0z7jfVQW54x2VWtnDD5JK98njwWtSSDyUpGX/DmN6Y9ruvS0dFBd3cPvQsXsmD+Anp6eljQ20tv70IWL1pIoaWFfK5AOpOpDHKM9/U/gG1AH7AP2Avsj9+PAP5siX0cLlZMZhBxl24P8FsTTL3yyiu47LJvs3//fgYG9PwxxlUYGRnBcZwxAdVx18/UCi47MjU3yYwiNpgcF+MupNPpygBCIxJJ8XCUP0o0zXfme1N0OnGDV5YXEVra2pg3bx7zurvp6ppH97xuli1fzooVK5nfM5/Ork5yuTy5XI50OoPnuShVOWb/DTyOFoo9aLE4APQDQ0qpOV9jwYpJk4m7dBcDJ4jID8Iw5OGHH+a///vj3H//fZTLZYaGhuIbxhmTC2K6d4eHhw/p5jRDTJLXhBCOyaRNWhRmMKCJdSRdmjAMcfBHxVmMpWPWkcvn6Z43j+7u+WSzWVpaW1m3dh2bn3MiCxcuIpvNkS8UKBQKpFJpXNdJisU/AjvQArEnfu0HhoCyFYtDY3tzmkRcuf1Y4BdhGLJv3z6+9a1LuOyyy3nqqSfjAF11BK4WD4dMJlN5kpsnePLpfDCiKWq1c4hrJRkLqXUn9F9wnGoui8ldCYJy3GVdnXDLBD5NjovrKua1z6Ozs4t587pJZ9N0dc5j46Zj2HTMsfT0zCeby2o3JF/A81K1vVDvQ8cn9qFFYy86ZjEAlJSaakTJ0kismByCuGZqN3p8zNeHhga5/fbbueyKy7jhhhvwPI+BgQHyLQVSqZQesRsI6ayelDv0o0qXbzLeYG7MQ7k5zhR7cw71e4VC4j4foRocVUoLShAInutQDgMkivBSaZQSlBI62ztIpzN0d8+jUGihpa2VVavWsPmEzaxYsZKWllbS6RT5QoFsNlexzGL+E3gGLRa74//3ElsV6HiFHfI8i7BiUodEsaFjiLt0n376aS782oVct3UrO3fuBLT/PzIyRC5XoFwuE4aC56XjXBGHIAiJEvVCkrGGiU5ApaYY8+CQYqSodiALoFDKwXUd8vk8bW29LF66lCiKWLRoMb/zohfR0zOflpYWuru7yWZz5PJ5Muk0KhH3AP4GbVEU0T0h29HuyDBzILA5F7ExkwQikkHXTH3IdOn+6Ec/YsuWLfziF78AqMQ4TI+FoTamkPw/uWxtMlltGjxUu2KjKMKJ77lkV6yJT5jPjJthgp/J/6PAH9VmI2DpdJqurnnMm9dNW3sbhUKBZcuXs379BpYsWcq8ed0UWlrIZXOk0ulR5Q2As9FBTNMTshPthgwAJaxVMWuwAdgGIrpyewvwfGCL7/vccccdXHTRRWzduhXPS1MsFlFKkU6nK4WWTczDdI3W64kAIBrdNZrMtajtSgXGrMd1goTrMTrzNbmeyvKxK1EotJDP51i8ZInuFk2nWdDby6ZjjmHVqtX0zF9Aa0sLmWw27gXxkufyLKoxin1UA5z9wIgViqMHG4CdIvH4mAKwGrhHROjv7+fSSy/lsssuY/v27eRyucr0l0DFUjC9EeZ/YyXUWimVvIwwQt+j+uYXRHsWSscqRKgksDmOg6Ociki4roMSVXGLKj0guRydXV309Mwnk8nQ3tZOS2sry5cvZ+3a9SxeuoSuzi7tfmQycS9IJV7xTeBR4Gl0rGI/VdEYQPeC2MCmZdLMGcskFhAzf8x2EaFUKvHTn/6Ua665hi1btlQm3R4eHq4UW86k0qMEBRg1HYSxEGp7QioWiDPaEqktZGSqw5suY8/zSKVSdHZ20tHZyfzuLgqFFlasXMHatetZsKCX9o52WgotZLM5srHoJVyQzwNPoKdh2MdosbC5FZZRWDdnEoiIB2TQvQZEUcRTTz3F+eefz3XXXcfAwAAigu/7lZs6nU5X3Y1Iuxme52EG6CVjE8lxKvXEIptNV6wYY110dXXR29tLd3cPruuSzxdYvXo1a9aupXdBLx2dnXR0dOhErFQapyoWX0AHMndTFYpn0RZFP1C0gU3LZLBiMgFEpBU9PuYhgKGhIX70ox/x1a9+laeffpowDPF9P86H0K6K67qUy+WKtZFKpYj84VEjc5MZnskkLGN9ZDIZ2ts76O7upqW1FeV6rFi+nA0bN7F0yRI6u7ro7Oiita2VTCZDJpNNZn+eQ1UstlPN2hwC9isVj1qzWBqEFZNxiIOpC4EXiMjlQRBw22238bWvfY3rr7++knlqxMGkiZtkMjNozIiK4zjkM1AqlSojXD3PI18oML+nh9bWdnp6eli6dBmr16xl4aJFdHV1097eTj6fJ5VOx1ZORSw+iRaInVQHjm1Hd5cOHg0DxyyzCxuATRDnhHQCa4FbzDD/K664gksuuYT9+/eTSqVwXbcSBzFWSSaTIQgCMpnMqIFs8+fPp62tjc7OTlpyDvN7F3LcscezYuUq2ts7aG1to6W1lXTcZZpI7/4PtEjsotr7sZN4pKntBbEczcxKyySRVLYUuEtEGB4e5totP+ZLX/oi27dvo1Qq6e5O8SsioeMTedra2mlta6N7Xje9CxexYeMmlixZSkdHJx0dnbS3d5AvFEin0riju0z/De127EbHKsz/g+j0bisWllnFnHZz4lohJdMbc+89d/PFL36em266kSAI6OrqoqenB8dRrF6zltZ8jo3HHMuKFSvp6pqnLYu2VgqFFlKpdDJe8T9ocdiDzrHYgQ5y7kZncdqxIJajjjkrJiLSDhzQ88fcz8UXXUQmk6FQaOFFLzqFxYuX0N7REVfFSpFKpZPt+zRaGPag3ZBn4vcmEcvOC2GZc8xJMRGRdhE5UCqVKJfLZLPZSqGbmE+ha1ckg5s70UVuSja/wmIZy5wUE6hYJivQk1SZ2hV9aLGw+RUWyySZs2JisVgaSyPFZGplvCwWiyXGionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDaIqYiIgtHGuxHGUcUTEREVdE1gCRiBREJB1/ZsXFYpnlHLHq9CLiAoHv+4gI6XS6dpGXAo+ip6/wgdDOdWOxTC+zdeLytwwMDPC2t72Nxx57jJNPPplVq1Zx7LHHsmbNGhYvXnx9S0tLcrpOROSDwN3AU+hZ+EqAAJGdqtNimVkcEctERDpEZP9ZZ53FD3/4Q8rlMkEQ4LounucRRXq+70WLFrF48WKWL1/OSSedxLp161i6dCn5fB7PG6N7L0BP8bkfCIDATsRlsUyOWTUJVxwPie677z5e/epXV6b0HB4epq2tDd/3UUqhlMJ1XRzHoVwuIyLk83kymQzLli1j4cKFLF++nA0bNrBp0yYWL15MJpPBdd3k5t4NPIy2ZPaiLZkQEGvJWCxjmW1i8vxisfjLd77zndx5552ICEEQkEqlKJfLFTEIwxARqbyPogilFJ7nUS6XAUilUiiliKKI9vZ2li9fzvz581m3bh2bN29m5cqV9PT0kM/nR7lLMacAT6InKi+jrRkrMpY5zawRExFJA6XLL7+cj370o4yMjOB5HmEYViYcb9TOmHa0t7ezYsUKNmzYwLJly1i/fj1Llixh4cKF5HK5ihWU4IVod2kf2pKJlFJRQxplscxwZpOYvPmZZ5759pve9Ca2bdsGaOvCiElSVA6XKIpwHAfHcRCRioUD4DgOmUwGx3FoaWlh/vz5rF69ms2bN7Np0yZ6enpYsGAB2Wy21l16J3AfWmQG0FaMb0XGcrQxK8RERBZFUfTM2WefzcUXX4zrupTL5YplYt7X3MSTxoiJsThq9ycMQ5RSFcGJoqiyfcdxKBQKrFy5kmXLlrFs2TKOOeYY1q1bR29vb0WIErwDeAh4Gm3J+CboKyLKukyW2caMFxMTdP3lL3/JGWecge/7RFFEqVQil8sRRVFFVKa6fdd1iaKIKIpGWSRGXEyMRUQqL9NuIxTJZYyF09HRwcKFC1m7di1r1qzh2GOPZcWKFXR3d5PL5ZIi8zF00PebQDfakgmtyFhmA7NBTF4zMjLyw3e9613ceOONpFIpfN8HoFAoUCwW68UuDgvTfhEZtT4jDkk3yHXdcQXGCJwRuXqCl0qlWLRoEWvWrOG4445jyZIl9PT0sGzZMhYsWEAmk0m24cNoC+bbwEKqXdhi3SXLTGFGi4mI5IDhz3/+85xzzjmVp73JK3Fdl+HhYQDS6XQlx+RwSbY/aZGY71zXrcRSjDCY5ZK9R0mrJQx1ukomkyEIglEulHGTTLvb29srsZmWlhaWLl3KqaeeykknnURPTw/z5s0jnU7XukunAfcDO4E+IMIm4lmawIwVE+PePPTQQ7zxjW+kWCxi0udre28cxyGdTle6fQ8XIxa1ro4RABPkNe5LUmyMaJiEOPNbsx7TfW3aW+s+AZUeqmS3dRAElXWtWLGCRYsWsXz5clauXMmaNWtYt24dPT09ld8keD3wGLCdOPBrrRjLdDKTxWSt7/sPn3nmmVxzzTUMDg5WrA/Ti+P7PplMpmIpTDUAGwRB5Uav192cjIMY8UhaJuZ9Ms/FiBBUxap2nUnXJ2mp1Oa3mCQ80JZYKpUiiiLmz5/P4sWL2bBhA6tXr64k4rW3t5NOp5Pr+Dt0fsxv0JbMIDoRz1oylikzI8XEDOT7xS9+wemnn04YhhU3wTI+xlpyXZdMJkNvby/r16/nmGOOYdGiRaxatYr58+fT2dlZz5J5A3pw5E5Gx2SsyFgmxEwVk9fv3Lnzu+9///u56667RvWOWA5O0lIyVo7J/k2n03ieR3t7Ox0dHWzcuJGTTz6ZjRs30tXVRUdHB6lUqvY4vx74NXpw5DCJ3iWLJcmMExMR6RKRveeffz6f/OQnK+6L53nWMjkEtXEc85nB5OYYwTABYHOMN27cyPLlyyuWzNKlS1m6dCmdnZ14npdc798AvwUeB3ZQjclYK2YOM6PEBFBAdM899/ChD32IRx55pNJr4/v+lGMiRzvJXiLzvjYJL9nNDVSOrznGyW5xk+3b29vL4sWLOeWUU1i0aBHLli1j/vz5tLa21rpL70MLzENod6mIHVIwZ5hpYnJyuVy+9b3vfS8/+clPKrkWQRDgeV5DG3s0kxSEZO5LKpWqfJbECIzneRVBMQKTDAabwLJSis7OTnp7e9mwYQPHHXdcpZepp6eH9vb2SvA55t3owO824Fl04NcGfY8yZpKYpIDyZZddxr/9279RLpdJpVKICL7vk06nK70ilvrUdmXX5skk81ySWbtGcExinVkHMKrXysRfzO+Mm2S263ke3d3dlVoya9euZcOGDaxZs4auri4KhUKtu/RB4EHgEWA3MIRNxJu1zBQxcYBwx44dnHbaaWzbtq3S9VsoFBAR+vv7yefzDWvs0Ui9PJmkeJjiUcmMXagKh+kaN1aFWVcyzyYpRLVJhmZZM1YpCALK5XLFktmwYQPr169n7dq1zJ8/n0WLFrFw4ULa2tpqLZm/QbtKj6IDv0V04NdaMjOYmSImC6Mo2v7pT3+ac845p1L0yPM8isUiURSRz+ennJR2tFMrHrWCkRyomMyNMcslM3hrxcisozZPJvmZ67qVJDtzDmt74oyYOY5DKpWiUCiwZMkSVqxYwcqVK1m9ejVr166lt7eX1tbWeu7S42iRqRSssiIzM5gJYuIA4d1338Vpb/lT/HJY81Xyr7V+j2aMKCmlaG9vZ9myZRx//PFs3LiR3t5eVqxYwbx582hpaakNxp+OLvHwOLAH7S5ZkTnCzAQxOXloaOjW97zn3dx8y80gyRwHKyZzjdo4T9I6EhEWLlxYcZGOOeYYNm3axPLly2lvb6+ITMKSeT86JvMg2pIpY2My00azxSQDFK+44nL++Z//qc7XVkzmEsn4TD2XyoxxMssm3bG2tjaOP/54Vq9ezcqVK1myZAlLliwZ5S4lOAMdk9lGtSqetWSmSDPFRAHRo48+wgc/+EF+9asH6pQRqLVSrJgczSQFJDnWybxMzCcIAqIoqgR6TdC3VCpV1mECzoVCoZKAd8IJJ7B27VpWr15NT09Pvap4f4lOxnsYHfgtWytm4jRTTNYHQfDge97zF9xwww2VoNxobPr8XKQ2g9e8koFjIyCmyxuqAyVNr5QZCBmG4agMatd1aWtrY/HixZx44okce+yx9Pb20tvbS3d3d60l83HgV+gR2M+iYzIjVmTG0iwxcYBwy5Yt/NM/fZgDBw7UySOxQjKXqXV3kiQT64Ax/5sHk+u6Y4pT1ebJJMlkMvT09NDR0cGKFSvYuHEjxxxzDCtXrqSzs7PWkvkYulfpYXTgdz9z3JJplpi8bGBg4KdnnHEGt9/+S9LpdOVpUsWKyVzDZNwmE+HqFaiqt7yp/5L83nSHJ2v7HgyTaW3cpyiKKukInZ2dlZq+69atq8Rkuru7K9OhxPwX8ADakjExmTkhMs0QkzYR6Tv//PP5r//6r0o+QhCU4xNiRWSuUpu/UhuINWJiclWSVkuyuLgRkGTy3XiDH817s24z2tpYNkkXyrhUYRiSTqfJZrOkUimWLVvGmjVr2LRpE+vWrWPlypV0dXWRzWaTIvMv6FkKdqJFZg96HuyjpsxDM8TkHx5++OFP/tmf/Rm7du2q1HSt6dKrQ4QVGstMJjmOrL29nXXr1rFixQo2b97MqlWrWLlyJfl8vl7g90/QlswO9MRuZs6lWSUyR1pMThoZGbnjHe94G3fddU/FDE2OWh1L0jq0YmKZuSTHLxkLykxZayoCtre3s379ejZs2FDJ+F25ciUdHR3j1ff9LXqQ5AFmuCVzJMXEA/xrr72W973vr1FKm5CmOHQQBKTT6To/s2JimR3Uxnlqy0GEYUgqlRoVLE5aMqtXr+aEE07g2GOPZfHixZWqeDXuEuiqeKZ3qQ8dk2m6wBxJMTlt27Zt33rrW9/C9u3bGRkZwXVTZLPZSiPGFj+qjVlZMbHMXIybnhzXZIK55vsgCCr5MKY+r4nRJDshUil9b2QyGVasWMHatWt5znOew4oVK1i6dGmlvm9CZD6I7lV6Ch2TOcARLlh1pMRkQRRFO84++z+56KJvAsQTjmcAGB4eJpXKxIqdFBArJpbZRVJQDMlcl2SsxCTi1VYT9H2/kidTrzxEOp2mpaWF9evXV8puLliwgPnz59PW1lYrMh9CDyd4Ai0yw0yTyBwJMXGA8NZbb+Uv//LP6e/vRykVJxQlF0megNq/o5ezWGYiyeEAMLpb21grtd3ZJhnPdEQkq9yZOGKyNESypoyp9SMilelnu7u7Wbt2LZs2bWLt2rWsWrWKxYsX09raWjtTwd+gBcZMhzLIFIcUHAkxeWOxWLziTW96E/c/cDdI9aCUy0E86VSOUqlUmXOmrpioCMSrXbfFMmMwAdjaZLjk3NTJIK2pL2PKNhiXKFnRLjk0wCxXW/gKqi6UWYeJ15jpYBYvXsyaNWs48cQTWbduHQsXLqxUxaspvfkBtCXzJLp3aVApNaHiy9MtJh7gn3322VxyySUEYZEoJBFwzcbdwqmapDUrJpbZR/Lmrze/krFCkol0wKjPk+VJa9dhPoOqFW/cJh02SFU6NTzPq/QgmXsrGRROpVK0tLRUpkA57rjjWLx4McuXL2fhwoW0tLQkRebf0bGYh9ECs4M6QwqmU0z0QL5HHuaNf/wqhoaGQNobtrHDQur1Fk0CNbXiTKqBB/swW9Dk7c9uZIphhilHKeTgBdWTmcKjfmZ6mUiP+SyJSfhLpVIsXLSQDRvXs/mEYznm+E309i6gZ0E32Vw2Lr0JwJfQbtKj6J6lp9C9Sz4wpb2tFZNVQRA89pfvfDs33/J/Wl1Vk6vLH+JkHJJDTBdzqHTt6Y6rH+rJcKj2HZqDx6wa+WSaHqY4F/UUD59zCDGf6vE7WJavXiCou6whWQXPZAObhNKWlhZGSiVWrlzF8hUr2LB+E2vXb2D58pX0LlhIvtBCNptFqco18j70zJGmiPjApPYlsQMOEN515+38+Z+fTj6TZWhoiHLYj0g0yq80ZpvxJWt3rpapHfCpDo+YWgBYyZEbnlFfOKrtP5wJzeqP7E6sPfFd4qI65LLNxnEmphKhjH6Y1F6fIge/ficrJko5iESVY+m4VfdnPEw8ph6lcv+YWMzo7alRMZmki6WUwg+DSj0ZPaQgQz5XoKOzk57uBaxYuZr58xeydu0mNmzYzIL5i8jlWvC8FEqpNiYhKEkxaRGRgcD3EYxQhHWFYNTuTPnJaTC1TRu0ukOQbPVEN1n3hE+1wQ07frOIOodMpmZhj0vtOdM33jgL1zMMmiyg9duqRn1fvQQFUHWv0+p+j3arlFKEUQgCrufhVH5bWU8eGJlQW2vEogVYjH4c/ho4FnAZe5hlnP8nw3intPbzid5tk21HcvmJbqOemTJVNTiU6XMk1OZwt3GoYz6Z9U70epgstcdXMblrZapTUh7O/THe/SV1PgupHiM3fp/cZ7O/5j42y3vx3xDIxf/3oYXDiT8fYYJCAgdPWpvQQZ+6T2+xWJpFI2NmB7PhZnpkzmKxzCBmTkTNYrHMaqyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA3BionFYmkIVkwsFktDsGJisVgaghUTi8XSEKyYWCyWhmDFxGKxNAQrJhaLpSFYMbFYLA1hyrOKz/zpJS0WSz0aPU2NtUwsFktDsGJisVgaghUTi8XSEKYcM7HUJQOsBrqBAnoO12eBp9BzuFomxmTnBbboYzYfWIK+/hxgN/AQMMg0Hs/JiIkDrALS8fsw/v2hJt52x/lcxS8/8ZnEnzlULySFnjw6jNflxC/znfk8ojppcyr+zR5gP0fugvSA/we8Enhn3M4UUAT6gW8BVwM/P4JtOhxcYCX64tvJ9LY1AywHsuhzF1E95+Y6aBSTmRw9ituwG/0wONR1PhPYCLwUOAM4geok5T7wK+B/gJ8AB6Zj4webuDxJHngXcCb6SZsGSkA5/m7Memv+1n6eJOlqCVWRUDWfm8+CeLugL3o3/sx8P4y+gV3gMeBs4DuH2L9G0Is+Pu9lfPcxRLf1I8BnmJmC0os+138Xv/8S8B9Mj0W1Gng/8Jfo8zWCPnfmgRLEy413PY3HZI9r7fUG1QfYNvQNeCvwXWamqCjgT4Avou/PloMsezPwPuDuhrdCRCbyerOIDMrsZK+IrJKJ7efhvtpE5EuTaFMgIu8QETXN7ZrsS4nIP9a0tSQifzsN23JE5JsiEk3iuDWTsoh8Q0R6pPnnqfb1ChEZmMS+PCoiXdLg62+iAdg/RyvebKQAvGwa1+8Ab0NbJBPFBb4OnDItLTp8cow9Ving92l8sL4DOI7GujHTSQp4B3A92nprNua4bQQu4+DWSC2rgW/T4JjpRC+Qn1J1LWYbPnDbNK1bAa9CuyzjMXSQ334OmNfgNk0FQcd2kijgURp/06eZmS7DoTgW+ATNF0FBi//fooW5HgdzTf8AeDcNfEhMdEXfAH6EjpME6Bu0HL83sZPalx+/TJwg+TLBrWEm5t+GiXWFiW2E6ODmMDpYWELfvAPx/7vQfuQDE9zPyaDQCv/v6CBiPT6P9k8vo/6NcyJwHjOnV03QgdAkEdA6DdsqMf5NMET1fPs1r+R1FDbolVxfkNjWeA/QtwCvPox9bhQmxnMm2muox2PouNf70Z0Q9fgUDdyPiV7EzwJvRfdSLELfuAH6IjM3idT8hdG9LgZV8/p7dDS/nrCNANcC9wBPxsvkqJ5oN94HP/7roC/SFrTI3AHcR+ODhwrtPp0PnFTnewE+iBayCC0mQ+genuSxcNAn8yx0kLPZAVnTCzZE1a01x7TRT2KF7iWp5RHgv+Pv2mq+M8fnUG2ZzHE0gd5kkN9BX08Z4M1odzS5TQ9tkf6Y5lhXgr4f/576980DwCuA7fH7u4GvAetqliugraxHgAen3qrmBPnM3x4RuadOgCgQkbtF5EQRSdX8RtWsq/ZVb7lGt98Rkf8QHZysZUhEPhsvk2xPVkQuqbO8iMhOEfnjaWrvZF4FEfmBiIQ17fuqiHgN3tY8GXvufdFBzuRxqz32jX7Vbqf2+wUi8hMZyw0i0joN52Air2NEZHudNomI/Db+vvZeOFiQ9koRycjY4z2pVzMyYJNPjVcAm+ss8yDwe8BdaKsjafUkfy91XvWWaySC7oY7k2rOTZJbgX9CP7GS7Smhu4TruVzz0TkAixvd2MOg1pKc7m0lcai6WfUS1uqd76m+attR+/1O4Lo6be+q074jQQfacltY57sB9LX3a0afR0Hvw+eo77r9MfDPU21YM9PpU8Bf1fl8CB0YOnBEWzMxFDoA9wXqJ+M9CpzG6ES8JE8C/4DOXailCx3IzU29mVOinpgcqZsmmaTWbJfPoNDnurY9JcZPyJyudij09fP/6nxfBj7N6Jyq2ofqv6Nd7lq330GL0OuZwoOkmWKyBlhf5/P/Q/e+NDtaXo9FwCeBnjrf7QH+Gp0xWQ9zUq8FPsbYXp4W4A3ooFm9JKq5gou2+JzEy028nAa/3HG248XtOA3dY1KbRLkHLShHArPt16IDqrWxzhI6fvNpxrfKBS0i7wXur/N9Bm25bOYwr71m9iJson7uyneYeV2G5uZ+G/ByxopwgO6V+Wn8frynqnnqX4DuyTkdbaEZPLS5+SRwSQPaPRv5A/S+e+jsatNVPZFs6lqS11FtJwDo85EMMJuUfh8tJK3A7zC2h0uhM0mPlJgIOj3+Yur3rO1CP8hGOLRFN4juAdrC2LSEpWgX6k8ZmyJwSJopJkn/2BCi4yQzkT9Bxzxqj1mIHnMz0d4YQYvPvwDHAyfXfF+It3Mb8PAU2jtbaUUf65nMr9E9eUeKFnSSYz0OAH8B7GDiruFdaHfpPMbG/V4K/CP6+pzUQ72Zbk69HTc5KzONFwAXUrWkkj7ng2jTcbw4yXjsRJ/Q7XW+Ww98k7Fi2wymM5g9Gymh8652HaHtOegUg03Ut0o+hXadJ3OOIvT1dQFj77cUWkzeczgNbRYdjD0AeaqDu2YCJjHtE4x2yUzgbRf6KTpeluuhuBH9BBip893z0YMUbc2ZmYVCuwpHojyCQndSvJmxFoSgr0sTJ5ksEfBh9Aj22gehi+6tfBmTiJ80+0KtbagpJzBTSAEfQN/YtW0dQSebPcThX1SCfkJ8EZ3Fm8RBJ7m9+zDXfbgcyZ6c8S7UEqNNbGMdRYlXoyymeus2mbD1zPw02g3d0IBtH4oXAP/KWAtV0CkIn2DyFnGSYXRy5WN1vlscr3/Cwz2aGTOpdyHNpB4MhR4afzr10+W/g84qnOoFHQL/BZzK2PhJJ/oJcRc6m3cusAO4lLF1a6aLMqNr9BgxMQHY16LzgJIsRJdp+AemT2x7gI/X2TZo1/jv0fGSqWxf0PGf/wa+wtjrfDO6h+ftTCB+MlPGhMw0FPAStBlYbzTm7WiLoVG9Tv1oc/YqdFGiJMuBL6N7OQ40aHszhXpW0DVon73ZFelMD945wBXA2prvVqNvvuI0bNtFX3un1vmujLaMbm3QtgS4CN0Z8F5GC4qHTle4F+1OmeXr0kw3p97TZiYE+xT6yfNFdOm7Wrahn1ZTMS9rEfT4iTPR+Qu1PAf99EjV+W46qM0IPVIkB+81G+P63IN+ate2aQ31M6CnioO2BP62zndl4Kvom9+0sREIWsC/X+e7LNrV+r1DraTZMZOZSAqdibqozne70N1w01XK8Ap0dm1tQNZDD+x6P81xBY+UyLvoJ+NMcnehvqvl1vlsKhhL6Lnoa6Deg8N06SaHajSKAB0/+XWd71oY/56oMBMtk2ai0CUDXouOVyQZRvfLb2H62hmhM2xvZazlU0CnPL9gmrY9HkdyrI6D3u+Z8JAzo4ePRV8TtW16ksamMQjaIv4a9V3rHegHynS4VcltvJ/6o7lPQMf2xu0gaXbM5EheqAfDtOFU4D+pX9f2f9Fp8NNNiWqG4pqa73rQT61XMX7a/lQZLxV7OqgXc3o5OgvYFEM29X0P9hs4eBsPdY0lb5Aw8VkLOpbVVec3T9DYDNg0+vo6puZzYxW+C/gt0/vAFeAGtGh8os73b0G7fZ9LLF/BiolG0F19X6C+kNyGDrgeqYS6x9A+7HmMHQd0LNp6eTfTE1uoN5J3OoY3mJhELcvRuTfpxHJJasfINBqz/ojxLaQD6DhKo46LQqcBvKPOd2V0zGI6LeIkgh7BvhCdFpHEi9vyCDpQPopmmpPZOtv3aE4GrAL+kPoDD/egR1s+eyQbBHwP/QSoffpl0WMnnjNN280xVuDTdT6bCsaFqJesZ7aXXDb5ql1Po1+Gg90bl6B7OBpFOzpBsZ4FtgV9cx/J8Wo+2kq6r853HcAb6/2omWIyRH3/rxlJa4K+iWqj8xF6qoefcOQtKEFbIJcxVmAjdDJRo9sUoHMYkgJWZPwErsNF0EJSb/T1TCZCd99/gMZZCQpdUa72HDvoGiR/RXOywvfE295X57vaim1Ac8XkbsY29Kdov7AZ3MXYcTJfpxonaUZwOEB3Fz9K9YIqo4/RTdPQphI6o3cX1Ys7QD8dG/1kNAHt8ayTmcYedCD2T2m8sG5DlxBIivgBdHrCZAbwNZpb0NefOUcRuo1XUydEMdFJuKaL16GTc1ai4xLvoHmJWQod2PwY2t36OTrLcDqj5xNlGdpK2YweSfxBpk90PXRE/y3oIO/XqRbcafTFkkcf41ehe6uSs/ll0A+bZOW1qVKv8FK9GEyy0tpT6N61b6FTAmp/3ygK6FjRm+PtnAX8bJq2NRkcdCb438VtOQdtrY+J1zVbTED3p2fQbo9Ru2Y2yswGWGpyO2px0W0rU73pprN9xuWb7hiWqSNiAr3mb3KmxmZiKuIfCcyxMNucruD3ZFHoe1TQ10Pd+3QmiInF0uwHiKUBzITkIIvFCslRQLPzTCwNQKnmeALWqrUksZaJxWJpCFZMLBZLQ5ismGSYfYlGeaZnrlyLxZJgMmLioosq/+40teVwUejYT73AgYuuGl9bwYyD/KaR7ZpJJSgtlmllomJiuu4WUk2lbXb/v6EX+CzVYdu17Xoe9S2T1wJnTGO7lqDH1tQbOGixHHVMVExM2H4XcH3NZ80mB7yGas9UbbtuRGcx1rIZWDF9zSIDvJj69WMtlqOOyXQNR+gZ7Uz2JcwMQUkBj1PfpYjQI37rDdU3maTThYOeSHomZDBOmUN0Azvoka+ghx8U0eelHZ3ZXFt533IUMtkArBk9OhNqtRqG0UJSm/ZtXDOf+jf0bhpbxzW5XdCDo4qJbc8Ut3C6eB26TMO16Cp1G4AH0AMH147/M8vRwmQsExNQNHOLGFJokTH5+gpd4OYU4Gm0YN2PHrRVO8CqVpAK8W8Xx9saQF+gT1KtvGW20xq3YxN6OoDl6OpXMHpmwBTVKQwy6DEnETqWMhSvR6haEhIfl7BO++phpmQw1k8ubvuKeD8WJto9zGgrSSX+LkPfdMPoSZ6eRJfPmymifTBM+yP0PpuJxwvomJGXWO5I7o9N0z+CTEZMHHQ5t7uBb8eftaPdiO+gh8QfR3VCqRvQ85YuQpeiM/Px1jN5s+hSfaehh9s/jK429rvx7wN0sPTxePkl6NKGHegbcA26FP/d6Jv4G+hh8w56tOMv0WUXj0eXZdwd/78QLTYL0CUIPosWotPQN8alEzguvegiv6Y62FlAd9yuXnT1tp/Fx+bb6NooUL3QN6FHYXroId/70dMovBj4UdzeejU5x0VEnHj7B4BdSqnpvqFS6GJO29FFg/aj9+fP0Bba06Zp09wO4u1uQM9Gl0Zfm08coW3PbURkoi9XRG4QkffE75WIdIjIAyLyXBF5jYjcKCJvEJFOEUmLSFZEuuPXjSJyl4ikEutUIuKIyNUicrOIbBaRrvg78/v5IvLnIrJdRNbVbLtDRF4uIo+LyMki0i4ivfE2zLpviNuGiBREpEdEWkTkoyKyJd5eV9xGFb++LCL/NcHjslxEdsS/N9toF5E/EJHfisgLRaQ1/ixds+9/KiLb4n2YH7fbidexSkT+VUQeEpGlB2tDzfl0RETCsCwiocR0iUhOtMg0+jpRE/zsSL1SIvL2eL99EXlfE9sy2dfzRd9f5lpsdnsm9ZqoZZI000uJ9yG6t+QitDvyWsYWPCqizf6Xop9QL2Z0j9CbgFXoXBAzZ6+i6qYU0XOF9KELyKyPt3sg/v5BtHvyULxMX2IdUHVdVNz+4fh9P/CDRHvNPpnpFkztikORibdttmfW/xt00Zv7qT837fFoS+g16CK9yTYMo+uVnI12d65GH5+JlAPwAIrlvaCG8aSddLZrr2meiOSVUo0sSGT2aT5wItq18dDWyiDa6gzQ18mOBm73YO0x15FLtbDPbHB5/gQ9X86n0PfUv1F/HqUZyUTFxNxkPYyeAiKFvrk3outC7qP+STMxlveiJ5N6fuK7AvqGS07+bX6fXNetaLdqCfoGM98V0OZssjcnGVvpi9uZXCfovJSemuVNWyfTCzMSt8vEQsx6WtGlFd3E52Y7GbQr9j7q1xJNtuUS4IXook3/MIH2+AD57AL6B39Ded9/E+Seh5d7AenCBlCpYRFZBOxWSjWiTkcGeBF6KstNVOeT2Y12UYvoB9AbaZyYHEwYTC0aUxvl8XGWm2mk0OcuHb/ejQ5qvwHtpte7J2pJxuBqr+lpZ7KWyTaqQUFTEcsF7kRP3jPejppK5Hegpw3IUC1RN0I1SJlcPikIoIXqPVQL5kji90bsqPO7Lqpxmtr2DSa2l/x+MgWBnHgbtcdS0DdPquYz0MWgDwBba9pT79hFaL//TPRT/lCV35T5k8ssQ/yvEgWXURr8XcKR03DzzyNdWLMdHERkGfDsFETFRVuVpuSgQvfmPIiOZ/0O+mGTi1+mfbX7mQGWoo9lP+OLTh4tUFl0ztOuOssI1cJCMHbS7xQ6SNyBjvEcQN/Ek7Va2tHX7V709TLIxGcLSF6jZrs+Oqb2DFp4T0HH3G4AXgH83yHW5wInAS9BP3AzcRvvRedZXcPU5yY+KJMJwAraKklmk0bop7spL3eo7s8ALUZpqmLyC/S0AX8NnEtVRZMl9BT6ZH2XsQejAx3wq70hTA/NENULK3kChxh//9PUnwhpPLYx+qlg2tiWWCZ5Ab0N7WIlrbHxEHRJS0EHdrcdbGGlVCQiLiCpVEs0MtRLLu/hyAOUD/wnbnQy/f3HI6kTae/+3afARUQ6gX6l1GSfYDm0KW6ehP+NDoSPoG/aD6CnRhCq7ka9Hr0XoV05B7gY/dCo15aXoS21FrQr8M+MvR7M8AqzLeOu9qBvytPRPXlZ9DVYRFt9F1G/ePJ4/D56wqwSWjyH0J0MVzDx+XSy6DmkF8ftvho9ReeF6Hlr/hp93K5BW//m3CevsTS6zOa70dnpxXi9plrdn1ItCv4NdPnPZ5gGa2UyGbBmB5IXhYt+ktye+OxQ6zBdiMTvt6HjKP+Odl9OQz+lMjW/Ta4jKVpD6CdErQiZ9uYZ/cQw33tUn5bJ7xz0gZ9oopUb71OtNVNGn2iTy5JslzH/TXwmjb4ATAnLbPxdPv6ugL5Aj51Ig5RSUdyD88nc6isohy8l7WQgeIjBfReRis7E738XA7svxS/uAGQ/EIoOdE8mUJtFn7thtG//RarXR0BVlE1Kgdl/gzkm7VSnBa2dhCpJctDmUupfbxHVY96HPm4r0VXlv4Z+4reiz3Me/TD6OPpGbuPQJBM281QznV+E7pW7Ad0bl1y2Hm9Fu+7fQddV/Rw6TnYO+nh9IP6sk+oEXeaaTt4Hr0EL+AaqD88fo++nT8br34k+D29Gx/fedoj2eWjr5qVoS6ebCWjFZPNMyoy2AMyNN1FFN9Wtay+Ce9FzrP4+Wkw+gw7o7kOX+9+CdqNM7kfyICh0jKN2btakMNRLThth7LgZFbcxZKx5XIs5sUZ4arOCjQ9ce8LM7/4ZfQG68baM9ZVDX+C/Roukj77IFzD5TNJPUXjxh+nNM7DzUpxO6Auhz8uSDw6Q2fcOUgegmIG92dNZOP+c8x3y5wOIiHsIS8WN2zoSt+076AeLwUGf6yDer/Em1SJebijez0cZ/6kZxetqZfyq9g76ug7Q5+BD6BjdRvR5/SbaZdgfb+8jaLF5IXr61X/l4DVfTfvvQc+1lEZfu29Fn6+T0Dfzixg766I597+Htsbz6Hvq6bjNC9FxkqvQIvEVdJH1tvjzjrjdScukI7H+Mjrw/1ZGpxO8Hd2JYe73R2r2Jdm+56CF6AXo82sSVa9Fp1mMG/earJtjXsnPJlr0VsXLHmBs6rug/boL45dJQjsWfXK2oBX13YyeItEkjNW2y2wvRB/UjprPBX2CapOpjFClOLQLkhSNEmPFzIhvKvHerL8HPWPflsT3JfTJMxaOOZEmhuNS7cGaUM+EUmqPiKxMt21+PD30Wdj+d6xsa4FiDrwCRAUYOkB2OMPi1gFCriIovBCvsApwwkMIiomD5dCCV3tzC/q4++ibPyms9QL0xkpMJifWMkzVMhlPTCReh0ma+zP0ef41+ub/LaPPxU/RlvVK9EwJlwK/GmfdJH77ONUZAn6IFtPz0KK0Gi1Kf8voe8NYM38c728J/eD8d3Re0MfR06O+Am2h3Y92e94Z/34NVS/AtOVOtPv/YvRD6TjgZrT79Ey8X29FH68M2i29uWZfzP8b0Q9v0yNn7oERdE/Ty9FiU9fVnkzegRGDer85VKwEqievg7GD8motjV3og/Z1dI/HC9BTT9yFDkqZ5SKq2aq1wS/jhmVq2pwM3Na2zfzfzcQrknvxNkxg1KzHQZuWtf5ziLa4IrQ47Ea7CP1oc3R//NqOPg674++SGcSTCaI9jfIgeC7cthC+1wc3DjJ43TNw4xDcloefRvCDn+Fe9lcUL/kd9t15PmF5GPQTdzwEbSXsQz+dX4/2/Q154JVU3Y6DCWBI9fh1U9+ag6rrGzH++XEY3XsH+oY7mdFCYvahH52MORB/dso46zXUOweCnv3uzVTHib2O6nil5P50om9MQd/sZ6OPz2/Q8SJBXzMnxf9vpfrA3lSzLrPdP0Ef/++jhWAdugfoTHQaxsvRx+5/0LGm2t4e4jZ/kursDB9HC+Mp6PuwjD43H2Qc3ZhsElO9k1wbwzjYb81Tw7gdncCrqVoXZhmn5ndPo7NnL0MHm5KUqB//AH0CTL5DLR76wBuSgrKT6oVwKPJoM7/2qWuesNmaz0ErfvJGrSeqyfcO+jjVWj8TQQE8eNU32PaTbQz/EvZfO0jppiJ9W/fT/6NdDP98P/03HuDJO4UdpReR7n4+bjoP9aeHTLZ5GB1ANwP9Po32tVein4Cb0cfYQd/M4wnKYqpP8HsZa+ma3yxCn2+H8SduN93CZh2/Rfv9tT13hhAdL8nH353C4ddG3oO2UiJ0D5+JwSSvrX70k91Yc8nrI4i37aKvkRVotw/0uU9T3zsQdAb6O9CxljL6mP5nvI5d6IfyWTW/SZJGWzVtaPH9JDqw/Jv4//vjNv8/RqeHVJism1PP5A2YWJdY0oUwYpFBB+3+ED0ozJCcFyYpMleiI/210ewRxndLTFSbmt956G692jYSr6/uAavDqVS7p5MYcavtyhW0SfwvVH375L7Wa8+p6MmqfjbBNukfi3iAP7LzMVrLu8njkC+3kPcVYXmIMAxId8KB9HzKv/c+lrzur3E75oF2wc5VSh2qGzpCP+lejhbUP0ALUAF94RWpPmz+Ce2q/l/8uXmALELfBEa8W9HWq3kam/O/CO37m/lbcugnpU81tyVE38QnxJ+FwOXxejqoXsNpquO18mjTXcW/WUz1vNSi4naawlo+1bFfxoUxLlaEDmCaCb6N9dyKflhtRrsV/4wOtK5FWylptEgHaCFYR/XaWIbOXTLHxNyPXvxahg5MG8vNi9v4LNql7oiPn3GbzfGI0CLRFn93J9V5rEAL9B1oD2E5OgxRO9Zu0gpcTzj8Op8dDJNMBPqgXoU2nd5V2zjGuj9tVN0GcwP66MBV7TzB5iIOqXbzJtefYWzRJLPO+9GBuX9FH9zaJ6p5vwhtSu5hbN2SAfRNVUA/jZLr+DH6wnkLukvyUG7LK9GR/8nOLtgKwrPXnEf6lu/TKcBQP6Q6Yf4m9ndmKC/fwII3foSORetAm8o/mMRYnhDtjv0h2hTejL4R+tDn9RfomEBvvOy30E/lEH0eu9HXw6J4fRF6DNbL0WLRinb5TK9QR7ycQl8z70GfH4eqeARUxQh0ouRrE+vbE7dne7zOfegnslk+z/hB8854v1ZQvReMCJmg/bz4Ny3oB+W/xL/ri7ffj77pzfb+AZ1X0klVUFNUp05ti9fvxvt8RrwvIdWH5Eh8PBdTvQ6HEssch7b4clRjliPxcRqIt9sfvzfBcrNPRsiWxOstM96sm5PIvXdF5FoR+YvEZ52ix9S8egK/V6LHsTwpemyMGQdzoojsEz1+xjnIb5eKHqfy3MRniB6f88u4DfXGM1wtIq+v8/k7ReTBeL+oWWeriDwqIm8eZ52IHgN0nYh8QkR+LiILa77vEj3u5hXj/P5lIrJbRH7/EMfsH0Xk+viY1V1unPOqRET8+26VZ995ksifZURei0SvQA584FTZduUXxO/bIzF/ITo3ZVwmcG57RI+PepXocUauiGRE5PR4PysDhepQFpFIRALR42nGw3wXHGQZQ1TzfiK/ERG5RUTyUv+8LxSRoQmsoziBZXypHpPksUkeg6DOZxPZ9k4Rea+I/L2IDCfWc7BzIFI9ZoOix9iZe/SMxDI/EX1eD3tsDlQVOM1oszzDxHx5c9WbPmtjqt2NHgF8PnAB1VGeYbzuAjqK/Vl0t95dNevrQ/vn56EDaTei81WG4+0ku1ST1sG1wN/E6zwX3eX1ePzdIPAX6IDYAvSI2L64TW1os/gstPXyBNVRxkn2o586X0P7rr9Ej4Y27tjP0ElJX0O7ClvjdnvoJ9wadD7AGrR5Xy/uUxcRqZjA3jfPpmvbI5BKMdTSSekNf07LK06nfcEqs3haKTXVui6CjmHUxjFC9LF7FG29LKWaMm7KNmTRZniGqsuSzE3ZTzXPpoC2Zj30U95kUrdQtdrM9JpDVC1T0NdoG/o8L40/T3bpl9FWyS2MnxFbBj6PNvXTcVuy6Cd+hA6oLoh/b/bPuD3tcVtNMNq4R6YL27RVUXVRJN6PLqqJZ2ZdEn9mev364tf96GvJjOm5g6oLsyhuc4h2VczYqWG0pbIcPVK/DZ3g9tJ4P98Ur2t//HndDPHJTA/qoLPpdqCH8yv0wf8TdE/LExNYRzvad/53Rg8YVGhX5e/RJu59aFPqOKo30UfQwgNj3Q6AP0J3jxXQ2aVb0Bfk69En+ZbEb4l/vw7tahyD7i77EqN7HjbE212PNov7qBb9+Ti6v74H3fX2FUbngZjtvAYd5e9H31jX1bR7Sbzfp6CP4R50ALMFHSO6kEPkl9SbhEtE3rL/2W2Xdn5gKfvownnz2bS85E/wOiuTC2QAf6IuzSSuk7pNTPyt7Y0wn+eodoObWEmKar0c426a7M6JxuoO1qZkYDQZm6sXvzI4Nb9pNrW9U8m/yWXq9eAkOzxM+OEf0K57iL6/HfQ9sR+dlfuFOuvXK5vkRXKw7r2J/h5Gn7haMmjrxfRzm4Smiab/JgNLE8WY+OYiTuKglbqDaldiX836D5VrkzyZ490ABfQTKIv2Y3cfZNnRK68RExHJAiPbH3yA/d96Nxs/cAVO1zLQwnYdEEy2xskUxeRQHIuOB+TRovzz+PN3oQe6PYxOmOoBPoq+Pj5G1Uq1TJ3kRfR8dNDf9HANo8/J7Rzk/j/cicunKiqWBlJHTF45sHfPj/MtLbiZSiJvbgK9M5OiQQKjgFehzeccumvzwvi7C9EmdoS24JajE7Ra0IHILdjrcDo4mNU2/o+m+YljaR6vRLt3yWEIM5XFaBfVRQ9qNLkVz0NbLTvQMa4ser8yVGNjluljUu6cFRPLTKbeE9IyQzncTD+L5UgwXkDRMgM57JqgFovFksSKicViaQj/H0ecN7mC/VekAAAAAElFTkSuQmCC" width="130" style="display:block;"/></div></div>""", unsafe_allow_html=True)
    st.title("لوحة التحكم")
    st.markdown("---")
    st.subheader(">_ سجل العمليات")
    log_placeholder = st.empty()
    if 'logs' not in st.session_state: st.session_state.logs = ["System Ready..."]

    def log(msg):
        st.session_state.logs.append(f"> {msg}")
        if len(st.session_state.logs) > 15: st.session_state.logs.pop(0)
        log_txt = "\n".join(st.session_state.logs)
        log_placeholder.markdown(f'<div class="log-container"><pre>{log_txt}</pre></div>', unsafe_allow_html=True)
        time.sleep(0.05)

    st.info("نظام ذكي لتوليد الاختبارات المتوازنة من ملفات الإكسل مباشرة.")


tab_exam, tab_bank = st.tabs(["🎓 توليد الاختبارات", "📝 بناء بنك الأسئلة"])

# =====================================================================
# تبويب بناء بنك الأسئلة
# =====================================================================
with tab_bank:
    bank_sub1, bank_sub2, bank_sub3 = st.tabs(["✍️ إدخال يدوي", "🔢 توليد مسائل رياضية", "🤖 توليد بالذكاء الاصطناعي"])

    # ── إدخال يدوي ──────────────────────────────────────────────────
    with bank_sub1:
        st.markdown("### إضافة سؤال جديد")
        if 'manual_bank' not in st.session_state:
            st.session_state.manual_bank = []

        with st.form("manual_q_form", clear_on_submit=True):
            mq_text    = st.text_area("نص السؤال *", height=80)
            mc1, mc2   = st.columns(2)
            with mc1:
                mq_correct = st.text_input("الإجابة الصحيحة *")
                mq_w1      = st.text_input("خطأ 1 *")
            with mc2:
                mq_w2      = st.text_input("خطأ 2 *")
                mq_w3      = st.text_input("خطأ 3")
            mc3, mc4   = st.columns(2)
            with mc3:
                mq_unit = st.text_input("الوحدة الدراسية *")
            with mc4:
                mq_diff = st.selectbox("الصعوبة", ["سهل", "متوسط", "صعب"])
            add_btn = st.form_submit_button("➕ إضافة السؤال", use_container_width=True)

        if add_btn:
            if mq_text and mq_correct and mq_w1 and mq_w2 and mq_unit:
                st.session_state.manual_bank.append({
                    "السؤال": mq_text.strip(),
                    "الإجابة الصحيحة": mq_correct.strip(),
                    "خطأ1": mq_w1.strip(),
                    "خطأ2": mq_w2.strip(),
                    "خطأ3": mq_w3.strip(),
                    "الوحدة": mq_unit.strip(),
                    "الصعوبة": mq_diff,
                })
                st.success(f"✅ تمت الإضافة! إجمالي الأسئلة: {len(st.session_state.manual_bank)}")
            else:
                st.warning("⚠️ يرجى ملء جميع الحقول الإلزامية (*)")

        if st.session_state.manual_bank:
            st.markdown(f"### البنك الحالي ({len(st.session_state.manual_bank)} سؤال)")
            bank_df = pd.DataFrame(st.session_state.manual_bank)
            st.dataframe(bank_df[["السؤال","الإجابة الصحيحة","الوحدة","الصعوبة"]], use_container_width=True)

            bc1, bc2 = st.columns(2)
            with bc1:
                if st.button("🗑️ مسح الكل", use_container_width=True):
                    st.session_state.manual_bank = []
                    st.rerun()
            with bc2:
                excel_bytes = questions_to_excel(st.session_state.manual_bank)
                st.download_button(
                    "💾 تنزيل البنك (Excel)",
                    data=excel_bytes,
                    file_name="بنك_الاسئلة_اليدوي.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

    # ── توليد مسائل رياضية برمجياً ──────────────────────────────────
    with bank_sub2:
        st.markdown("### توليد مسائل رياضية تلقائياً بدون ذكاء اصطناعي")
        st.info("يولّد التطبيق أسئلة اختيار من متعدد رياضية عشوائية جاهزة للاستخدام مباشرةً في بنك الأسئلة.")

        mc1, mc2, mc3 = st.columns(3)
        with mc1:
            math_topic = st.selectbox("الموضوع الرياضي", [
                "عمليات حسابية (+−×÷)",
                "كسور",
                "مساحات أشكال هندسية",
                "قوى وجذور",
                "نسب مئوية",
                "إحصاء (متوسط ومدى)",
            ])
        with mc2:
            math_diff = st.selectbox("مستوى الصعوبة", ["سهل", "متوسط", "صعب"], key="math_diff")
        with mc3:
            math_num = st.number_input("عدد الأسئلة المطلوبة", min_value=5, max_value=100, value=20, step=5)

        math_unit = st.text_input("اسم الوحدة الدراسية", value=math_topic)

        if st.button("⚙️ توليد المسائل", use_container_width=True):
            with st.spinner("جاري التوليد..."):
                math_qs = generate_math_questions(math_topic, math_num, math_diff, math_unit)
            if math_qs:
                st.session_state.math_questions = math_qs
                st.success(f"✅ تم توليد {len(math_qs)} سؤال!")
            else:
                st.error("حدث خطأ أثناء التوليد.")

        if 'math_questions' in st.session_state and st.session_state.math_questions:
            mqs = st.session_state.math_questions
            st.markdown(f"### الأسئلة المولّدة ({len(mqs)} سؤال)")
            st.dataframe(pd.DataFrame(mqs)[["السؤال","الإجابة الصحيحة","الصعوبة"]], use_container_width=True)

            mec1, mec2 = st.columns(2)
            with mec1:
                excel_math = questions_to_excel(mqs)
                st.download_button(
                    "💾 تنزيل البنك (Excel)",
                    data=excel_math,
                    file_name=f"بنك_{math_unit}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            with mec2:
                if st.button("➕ إضافة للبنك اليدوي", use_container_width=True, key="add_math"):
                    if 'manual_bank' not in st.session_state:
                        st.session_state.manual_bank = []
                    st.session_state.manual_bank.extend(mqs)
                    st.success(f"✅ تمت الإضافة! إجمالي البنك: {len(st.session_state.manual_bank)} سؤال")

    # ── توليد بالذكاء الاصطناعي ─────────────────────────────────────
    with bank_sub3:
        st.markdown("### توليد أسئلة من محتوى تعليمي")

        # تثبيت المكتبات المطلوبة
        missing = []
        if not PYMUPDF_OK:   missing.append("`py -m pip install PyMuPDF`")
        if not GEMINI_OK:    missing.append("`py -m pip install google-generativeai`")
        if not CLAUDE_OK:    missing.append("`py -m pip install anthropic`")
        if missing:
            st.info("لتفعيل جميع الميزات، شغّل في CMD:\n\n" + "\n\n".join(missing))

        # اختيار مزود الذكاء الاصطناعي
        ai_provider = st.radio("مزود الذكاء الاصطناعي:", ["🟦 Google Gemini", "🟧 Anthropic Claude", "🆓 OpenRouter (مجاني)"], horizontal=True)

        if "Gemini" in ai_provider:
            ai_key = st.text_input("مفتاح Gemini API (من aistudio.google.com)", type="password", key="gemini_key")
        elif "Claude" in ai_provider:
            ai_key = st.text_input("مفتاح Claude API (من console.anthropic.com)", type="password", key="claude_key")
        else:
            ai_key = st.text_input("مفتاح OpenRouter (من openrouter.ai/settings/keys)", type="password", key="or_key")
            or_model = st.selectbox("النموذج المجاني:", [
                "qwen/qwen3.6-plus:free",
                "meta-llama/llama-3.3-70b-instruct:free",
                "deepseek/deepseek-r1:free",
                "microsoft/phi-4:free",
                "stepfun/step-3-5-flash:free",
            ], key="or_model")
            st.info("✅ النماذج المجانية في OpenRouter لا تحتاج رصيداً")

        st.markdown("---")

        # رفع المحتوى
        ai_content_type = st.radio("نوع المحتوى:", ["📄 ملف PDF (كتاب/محاضرة)", "📝 نص مكتوب مباشرة"], horizontal=True)

        text_content = None
        img_content  = None

        if "PDF" in ai_content_type:
            pdf_file = st.file_uploader("ارفع ملف PDF", type=["pdf"], key="ai_pdf")
            if pdf_file:
                if not PYMUPDF_OK:
                    st.error("⚠️ PyMuPDF غير مثبت. شغّل: py -m pip install PyMuPDF")
                else:
                    with st.spinner("جاري قراءة الملف..."):
                        txt, imgs, err = extract_pdf_content(pdf_file.read())
                    if err:
                        st.error(f"خطأ: {err}")
                    elif txt:
                        text_content = txt
                        st.success(f"✅ PDF نصي — تم استخراج {len(txt)} حرف")
                    elif imgs:
                        img_content = imgs
                        st.success(f"✅ PDF مصوّر — تم استخراج {len(imgs)} صفحة")
        else:
            text_content = st.text_area("الصق المحتوى التعليمي هنا:", height=200, key="ai_text_input")

        st.markdown("---")

        # إعدادات التوليد
        ag1, ag2, ag3 = st.columns(3)
        with ag1:
            ai_unit    = st.text_input("اسم الوحدة/الفصل *", key="ai_unit")
        with ag2:
            ai_num_q   = st.number_input("عدد الأسئلة المطلوبة", min_value=5, max_value=300, value=20, key="ai_num")
        with ag3:
            ai_diff    = st.selectbox("مستوى الصعوبة", ["سهل", "متوسط", "صعب", "مختلطة"], key="ai_diff")

        gen_btn = st.button("🚀 توليد الأسئلة", use_container_width=True, type="primary")

        if gen_btn:
            if not ai_key:
                st.error("⚠️ أدخل مفتاح API أولاً.")
            elif not ai_unit:
                st.error("⚠️ أدخل اسم الوحدة.")
            elif not text_content and not img_content:
                st.error("⚠️ ارفع ملف PDF أو أدخل نصاً.")
            else:
                with st.spinner("جاري التوليد... قد يستغرق دقيقة"):
                    prompt = build_ai_prompt(ai_unit, ai_num_q, ai_diff)
                    if "Gemini" in ai_provider:
                        raw, err = call_gemini(ai_key, prompt, text_content, img_content)
                    elif "Claude" in ai_provider:
                        raw, err = call_claude(ai_key, prompt, text_content, img_content)
                    else:
                        _or_model = st.session_state.get("or_model", "qwen/qwen3.6-plus:free")
                        raw, err = call_openrouter(ai_key, prompt, text_content, img_content, model=_or_model)

                    if err:
                        st.error(f"❌ {err}")
                    elif raw:
                        questions, parse_err = parse_ai_questions(raw)
                        if parse_err or not questions:
                            st.error(f"❌ {parse_err}")
                            with st.expander("الاستجابة الخام"):
                                st.text(raw)
                        else:
                            st.success(f"✅ تم توليد {len(questions)} سؤال!")
                            st.session_state.ai_questions = questions

        # عرض وتصدير الأسئلة المولّدة
        if 'ai_questions' in st.session_state and st.session_state.ai_questions:
            qs = st.session_state.ai_questions
            st.markdown(f"### الأسئلة المولّدة ({len(qs)} سؤال)")
            st.dataframe(pd.DataFrame(qs)[["السؤال","الإجابة الصحيحة","الوحدة","الصعوبة"]], use_container_width=True)

            ec1, ec2 = st.columns(2)
            with ec1:
                excel_ai = questions_to_excel(qs)
                st.download_button(
                    "💾 تنزيل البنك (Excel)",
                    data=excel_ai,
                    file_name=f"بنك_{ai_unit}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            with ec2:
                if st.button("➕ إضافة للبنك اليدوي", use_container_width=True):
                    if 'manual_bank' not in st.session_state:
                        st.session_state.manual_bank = []
                    st.session_state.manual_bank.extend(qs)
                    st.success(f"✅ تمت الإضافة! إجمالي البنك: {len(st.session_state.manual_bank)} سؤال")

# =====================================================================
# تبويب توليد الاختبارات
# =====================================================================
with tab_exam:
    st.subheader("1️⃣ رفع بنوك الأسئلة")
    uploaded_banks = st.file_uploader(
        "ارفع ملفات البنك (xls أو xlsx)",
        accept_multiple_files=True,
        type=['xls', 'xlsx'],
        key="banks_uploader"
    )

    st.markdown("---")

    st.subheader("2️⃣ عدد النماذج المطلوبة")

    ALL_MODELS_CONFIG = [
        {"name": "ا_صباحي",   "folder": "صباحي",    "key": "am1"},
        {"name": "ب_صباحي",   "folder": "صباحي",    "key": "am2"},
        {"name": "ا_مسائي",   "folder": "مسائي",    "key": "pm1"},
        {"name": "ب_مسائي",   "folder": "مسائي",    "key": "pm2"},
        {"name": "ا_دور_ثاني","folder": "دور_ثاني", "key": "sec1"},
        {"name": "ب_دور_ثاني","folder": "دور_ثاني", "key": "sec2"},
    ]

    _nc1, _nc2 = st.columns([1, 3])
    with _nc1:
        num_models_choice = st.selectbox(
            "عدد النماذج (2 – 6):",
            options=[2, 3, 4, 5, 6],
            index=4,          # افتراضي = 6
            key="num_models_select"
        )
    with _nc2:
        _preview_names = " | ".join(m["name"] for m in ALL_MODELS_CONFIG[:num_models_choice])
        st.info(f"✅ النماذج التي ستُولَّد ({num_models_choice}): **{_preview_names}**")

    models_config = ALL_MODELS_CONFIG[:num_models_choice]

    st.subheader("3️⃣ رفع التمبلتات ومفاتيح الإجابة")

    _upload_c1, _upload_c2 = st.columns(2)
    with _upload_c1:
        uploaded_templates = st.file_uploader(
            "📄 ارفع ملفات التمبلت (Word) — حتى 6 ملفات",
            type=['docx'], accept_multiple_files=True, key="templates_uploader"
        )
    with _upload_c2:
        uploaded_keys = st.file_uploader(
            "🗝️ ارفع مفاتيح الإجابة (Excel) — حتى 6 ملفات",
            type=['xlsx'], accept_multiple_files=True, key="keys_uploader"
        )

    def _match_files_to_models(files, cfg):
        """تعرف تلقائي على الملفات بحسب اسم الملف ثم بالترتيب كاحتياطي."""
        result = {m['name']: None for m in cfg}
        kw_map = {
            'ا_صباحي':   ['صباحي_أ', 'صباحي_ا', 'ا_صباحي', 'اصباحي', 'am1', 'a_sabahi'],
            'ب_صباحي':   ['صباحي_ب', 'ب_صباحي', 'بصباحي',  'am2', 'b_sabahi'],
            'ا_مسائي':   ['مسائي_أ', 'مسائي_ا', 'ا_مسائي', 'امسائي', 'pm1', 'a_masai'],
            'ب_مسائي':   ['مسائي_ب', 'ب_مسائي', 'بمسائي',  'pm2', 'b_masai'],
            'ا_دور_ثاني':['دور_ثاني_أ', 'دور_ثاني_ا', 'ا_دور', 'ادور', 'sec1', 'a_door'],
            'ب_دور_ثاني':['دور_ثاني_ب', 'ب_دور',      'بدور',  'sec2', 'b_door'],
        }
        used_names = set()
        for f in (files or []):
            fname = f.name.replace(' ', '').replace('-', '_').lower()
            for m in cfg:
                if result[m['name']] is not None:
                    continue
                if any(k.lower() in fname for k in kw_map.get(m['name'], [])):
                    result[m['name']] = f
                    used_names.add(f.name)
                    break
        # احتياطي: توزيع الملفات غير المعرَّفة بالترتيب
        unmatched = [f for f in (files or []) if f.name not in used_names]
        for m_name in [m['name'] for m in cfg if result[m['name']] is None]:
            if unmatched:
                result[m_name] = unmatched.pop(0)
        return result

    tmpl_map = _match_files_to_models(uploaded_templates, models_config)
    key_map  = _match_files_to_models(uploaded_keys,      models_config)

    # عرض جدول خريطة التعرف
    if uploaded_templates or uploaded_keys:
        st.markdown("**📋 نتيجة التعرف التلقائي على الملفات:**")
        _rows = []
        for m in models_config:
            t = tmpl_map[m['name']]
            k = key_map[m['name']]
            _rows.append({
                "النموذج": m['name'],
                "التمبلت": f"✅ {t.name}" if t else "⬜ لا يوجد",
                "مفتاح الإجابة": f"✅ {k.name}" if k else "⬜ لا يوجد",
            })
        st.dataframe(_rows, use_container_width=True, hide_index=True)
        st.caption(
            "💡 إذا لم يكن التعرف صحيحاً، سمِّ ملفاتك بهذا الشكل: "
            "صباحي_أ / صباحي_ب / مسائي_أ / مسائي_ب / دور_ثاني_أ / دور_ثاني_ب"
        )

    # بناء model_inputs النهائي
    model_inputs = {
        m['name']: {"folder": m['folder'], "template": tmpl_map[m['name']], "key": key_map[m['name']]}
        for m in models_config
    }

    st.markdown("---")

    # ===================== فلتر الوحدات =====================
    st.subheader("4️⃣ تصفية الوحدات الدراسية")

    col_prev1, col_prev2 = st.columns([1, 3])
    with col_prev1:
        preview_btn = st.button("🔍 اعرض الوحدات المتاحة في البنك", use_container_width=True)

    if preview_btn:
        if not uploaded_banks:
            st.warning("⚠️ ارفع بنك الأسئلة أولاً ثم اضغط هنا.")
        else:
            with st.spinner("جاري قراءة البنك..."):
                _prev_dfs = []
                for _f in uploaded_banks:
                    _f.seek(0)
                    _df = fetch_smart_questions(_f)
                    if not _df.empty:
                        _prev_dfs.append(_df)
                if _prev_dfs:
                    _all = pd.concat(_prev_dfs)
                    _units = sorted([u for u in _all['unit'].unique().tolist() if u])
                    st.session_state['available_units'] = _units
                    st.success(f"✅ تم العثور على {len(_units)} وحدة و {len(_all)} سؤال في البنك.")
                else:
                    st.error("لم يتم العثور على أسئلة في البنك.")

    # عرض قائمة الوحدات للاختيار
    selected_units = None
    if 'available_units' in st.session_state and st.session_state['available_units']:
        st.markdown("**اختر الوحدات التي تريد الأسئلة منها:**")
        all_units = st.session_state['available_units']

        # أزرار تحديد الكل / إلغاء الكل
        cb1, cb2, _ = st.columns([1, 1, 4])
        with cb1:
            if st.button("✅ تحديد الكل", use_container_width=True):
                st.session_state['selected_units'] = all_units
        with cb2:
            if st.button("❌ إلغاء الكل", use_container_width=True):
                st.session_state['selected_units'] = []

        if 'selected_units' not in st.session_state:
            st.session_state['selected_units'] = all_units

        selected_units = st.multiselect(
            "الوحدات المتاحة:",
            options=all_units,
            default=st.session_state['selected_units'],
            key="units_multiselect"
        )
        st.session_state['selected_units'] = selected_units

        if selected_units:
            st.info(f"📚 الوحدات المحددة: **{len(selected_units)}** من أصل {len(all_units)}")
        else:
            st.warning("⚠️ لم تحدد أي وحدة — سيتم استخدام جميع الوحدات.")

    st.markdown("---")

    # ===================== خيارات متقدمة =====================
    st.subheader("5️⃣ الخيارات المتقدمة")

    col_adv1, col_adv2 = st.columns(2)

    with col_adv1:
        no_repeat = st.checkbox(
            f"عدم تكرار الأسئلة بين النماذج ({num_models_choice})",
            value=True,
            help=f"عند التفعيل: لن يتكرر أي سؤال في أكثر من نموذج واحد. يتطلب أن يكون البنك يحتوي على ما يكفي من الأسئلة (عدد الأسئلة × {num_models_choice} على الأقل)."
        )

    with col_adv2:
        difficulty_choice = st.selectbox(
            "مستوى صعوبة الأسئلة",
            options=["الكل (بدون تصفية)", "سهل فقط", "متوسط فقط", "صعب فقط", "توزيع مخصص (نسب %)"],
            index=0,
            help="اختر مستوى الصعوبة أو حدد نسبة مخصصة من كل مستوى."
        )

    # خريطة اختيار الصعوبة
    DIFF_MAP = {
        "الكل (بدون تصفية)": None,
        "سهل فقط": "سهل",
        "متوسط فقط": "متوسط",
        "صعب فقط": "صعب",
        "توزيع مخصص (نسب %)": "custom",
    }

    # حقول النسب المخصصة (تظهر فقط عند اختيار "توزيع مخصص")
    easy_pct = mid_pct = hard_pct = 0
    if difficulty_choice == "توزيع مخصص (نسب %)":
        st.markdown("##### حدد نسبة كل مستوى (المجموع يجب أن يساوي 100%)")
        c1, c2, c3 = st.columns(3)
        with c1:
            easy_pct = st.number_input("سهل %", min_value=0, max_value=100, value=50, step=5)
        with c2:
            mid_pct  = st.number_input("متوسط %", min_value=0, max_value=100, value=30, step=5)
        with c3:
            hard_pct = st.number_input("صعب %", min_value=0, max_value=100, value=20, step=5)

        total_pct = easy_pct + mid_pct + hard_pct
        if total_pct != 100:
            st.warning(f"⚠️ مجموع النسب = {total_pct}% — يجب أن يكون 100% بالضبط.")
        else:
            st.success(f"✅ التوزيع: سهل {easy_pct}% | متوسط {mid_pct}% | صعب {hard_pct}%")

    st.markdown("---")

    c_num, c_btn = st.columns([1, 3])
    with c_num:
        max_q_count = st.number_input("عدد الأسئلة لكل نموذج", min_value=1, value=30)
    with c_btn:
        st.text("")
        st.text("")
        start_btn = st.button("🚀 بدء توليد الاختبارات", use_container_width=True)

    if start_btn:
        if not uploaded_banks:
            st.error("⚠️ خطأ: يجب رفع بنك أسئلة واحد على الأقل.")
        elif difficulty_choice == "توزيع مخصص (نسب %)" and (easy_pct + mid_pct + hard_pct) != 100:
            st.error("⚠️ مجموع النسب يجب أن يكون 100% بالضبط قبل البدء.")
        else:
            try:
                log("بدء المعالجة...")
                progress_bar = st.progress(0)
                all_dfs = []

                # قراءة البنوك
                total_files = len(uploaded_banks)
                for i, f in enumerate(uploaded_banks):
                    log(f"جاري قراءة الملف: {f.name}")
                    df = fetch_smart_questions(f)
                    if not df.empty: all_dfs.append(df)
                    progress_bar.progress((i + 1) / (total_files * 2))
                    gc.collect()

                if not all_dfs:
                    st.error("لم يتم العثور على أسئلة صالحة في الملفات!")
                else:
                    BIG_DF = pd.concat(all_dfs).reset_index(drop=True)
                    log(f"تم استخراج {len(BIG_DF)} سؤال بنجاح.")

                    # ===== تصفية الوحدات =====
                    _sel_units = st.session_state.get('selected_units', None)
                    if _sel_units and len(_sel_units) > 0:
                        _before = len(BIG_DF)
                        BIG_DF = BIG_DF[BIG_DF['unit'].isin(_sel_units)].reset_index(drop=True)
                        log(f"تصفية الوحدات: {len(_sel_units)} وحدة — {len(BIG_DF)} سؤال متاح (من {_before}).")
                        if BIG_DF.empty:
                            st.error("⚠️ لا توجد أسئلة في الوحدات المحددة!")
                            st.stop()
                    else:
                        log("الوحدات: جميع الوحدات مستخدمة.")

                    # ===== تصفية/توزيع الصعوبة =====
                    selected_diff = DIFF_MAP[difficulty_choice]
                    use_custom_mix = (selected_diff == "custom")

                    if use_custom_mix:
                        # وضع التوزيع المخصص: لا نصفي، الدالة ستتولى التوزيع
                        filtered_df = BIG_DF
                        log(f"توزيع مخصص: سهل {easy_pct}% | متوسط {mid_pct}% | صعب {hard_pct}%")
                    elif selected_diff is not None:
                        if 'difficulty' in BIG_DF.columns:
                            filtered_df = BIG_DF[BIG_DF['difficulty'] == selected_diff].reset_index(drop=True)
                            if filtered_df.empty:
                                st.warning(f"⚠️ لا توجد أسئلة بمستوى '{selected_diff}'. سيتم استخدام جميع الأسئلة.")
                                log(f"تحذير: لم يُعثر على أسئلة بمستوى {selected_diff}، يُستخدم الكل.")
                                filtered_df = BIG_DF
                            else:
                                log(f"تصفية الصعوبة: {selected_diff} - {len(filtered_df)} سؤال متاح.")
                        else:
                            st.warning("⚠️ البنك لا يحتوي على عمود صعوبة. سيتم استخدام جميع الأسئلة.")
                            log("تحذير: لا يوجد عمود صعوبة في البنك.")
                            filtered_df = BIG_DF
                    else:
                        filtered_df = BIG_DF
                        log("مستوى الصعوبة: الكل.")

                    # ===== التحقق من الكفاية عند عدم التكرار =====
                    total_models = len(model_inputs)
                    if no_repeat:
                        min_needed = max_q_count * total_models
                        if len(filtered_df) < min_needed:
                            st.warning(
                                f"⚠️ البنك يحتوي على {len(filtered_df)} سؤال فقط، "
                                f"بينما يلزم {min_needed} سؤال لضمان عدم التكرار الكامل "
                                f"({max_q_count} × {total_models} نماذج). "
                                f"سيتم توزيع المتاح بأفضل شكل ممكن مع تقليل التكرار."
                            )
                            log(f"تحذير: أسئلة غير كافية للتوزيع الكامل ({len(filtered_df)}/{min_needed}).")

                    # ===== توليد النماذج =====
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:

                        if no_repeat:
                            log("توليد النماذج بدون تكرار...")
                            if use_custom_mix:
                                all_exams = generate_all_unique_exams(
                                    filtered_df, max_q_count, num_models=total_models,
                                    easy_pct=easy_pct, mid_pct=mid_pct, hard_pct=hard_pct)
                            else:
                                all_exams = generate_all_unique_exams(filtered_df, max_q_count, num_models=total_models)
                        else:
                            all_exams = None  # سيتم التوليد المنفرد داخل الحلقة

                        for idx, (m_name, inputs) in enumerate(model_inputs.items()):
                            log(f"جاري بناء النموذج: {m_name}")

                            pattern = get_master_pattern_from_file(inputs['key'], limit=max_q_count)
                            gc.collect()

                            if no_repeat:
                                exam_df = all_exams[idx]
                            elif use_custom_mix:
                                exam_df = generate_balanced_exam_with_mix(filtered_df, max_q_count, easy_pct, mid_pct, hard_pct)
                            else:
                                exam_df = generate_balanced_exam(filtered_df, max_q_count)

                            if exam_df.empty:
                                log(f"تخطي النموذج {m_name}: لا توجد أسئلة كافية.")
                                continue

                            if len(exam_df) > max_q_count: exam_df = exam_df.iloc[:max_q_count]
                            exam_df = exam_df.reset_index(drop=True)

                            if inputs['template']:
                                inputs['template'].seek(0)
                                doc = Document(inputs['template'])
                                # حذف كل محتوى الجسم (جداول + فقرات) مع الحفاظ على الرأس والتذييل وإعدادات الصفحة
                                body = doc.element.body
                                for child in list(body):
                                    if child.tag != qn('w:sectPr'):
                                        body.remove(child)
                            else:
                                doc = Document()

                            set_section_rtl_and_margins(doc.sections[0])
                            # إضافة عنوان "اختر الإجابة الصحيحة" دائماً
                            doc.add_paragraph("")
                            sub = doc.add_paragraph(); fix_paragraph_alignment(sub)
                            run_st = sub.add_run("اختر الإجابة الصحيحة :")
                            force_font(run_st, size=14, is_bold=True)

                            for i, row in enumerate(exam_df.to_dict('records')):
                                target_idx = pattern[i % len(pattern)]
                                final_opts = force_align_options(row['options'], row['correct_text'], target_idx)
                                add_question_block(doc, i+1, row['question'], final_opts)

                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            zip_path = f"{inputs['folder']}/{m_name}.docx"
                            zf.writestr(zip_path, doc_io.getvalue())

                            del doc; del doc_io; gc.collect()
                            progress_bar.progress(0.5 + ((idx + 1) / (total_models * 2)))

                    progress_bar.progress(100)
                    log("✅ تمت العملية بنجاح!")
                    st.balloons()
                    st.success("تم إنشاء الملفات وجاهزة للتحميل.")

                    # ── حفظ تلقائي في Downloads دائماً ──────────────
                    import os as _os, subprocess as _sp
                    _dl_dir = _os.path.join(_os.path.expanduser("~"), "Downloads")
                    _os.makedirs(_dl_dir, exist_ok=True)
                    _save = _os.path.join(_dl_dir, "Taibah_Exams_Generated.zip")
                    _i = 1
                    _base = _save[:-4]
                    while _os.path.exists(_save):
                        _save = f"{_base}_{_i}.zip"
                        _i += 1
                    with open(_save, "wb") as _f:
                        _f.write(zip_buffer.getvalue())

                    st.info(f"📁 تم حفظ الملف تلقائياً في:\n**{_save}**")

                    col_a, col_b = st.columns(2)
                    with col_a:
                        if st.button("📂 فتح مجلد التنزيلات", use_container_width=True):
                            _sp.Popen(f'explorer /select,"{_save}"')
                    with col_b:
                        st.download_button(
                            label="📥 تحميل مباشر (ZIP)",
                            data=zip_buffer.getvalue(),
                            file_name="Taibah_Exams_Generated.zip",
                            mime="application/zip",
                            use_container_width=True
                        )

            except Exception as e:
                st.error(f"حدث خطأ غير متوقع: {e}")
                log(f"Error: {str(e)}")
